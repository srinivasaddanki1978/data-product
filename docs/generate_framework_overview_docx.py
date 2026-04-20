"""Generate Word document from the Snowflake Cost Optimisation Framework Overview."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


# --- Colour palette ---
DARK_BLUE = RGBColor(0x1B, 0x2A, 0x4A)
ACCENT_BLUE = RGBColor(0x2E, 0x40, 0x57)
LIGHT_BLUE = RGBColor(0x5A, 0x7D, 0x9A)
HEADER_BG = "1B2A4A"
ALT_ROW_BG = "F0F4F8"
ACCENT_BG = "E8EDF3"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY = RGBColor(0x99, 0x99, 0x99)
DARK_GREY = RGBColor(0x44, 0x44, 0x44)
GREEN = RGBColor(0x27, 0x7B, 0x4A)


def set_cell_shading(cell, color):
    """Set background color for a table cell."""
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), color)
    shading_elm.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_table_borders(table, color="BFBFBF"):
    """Set borders for the entire table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        borders.append(border)
    tblPr.append(borders)


def set_cell_margins(cell, top=50, bottom=50, left=80, right=80):
    """Set cell margins in twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("start", left), ("end", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a professionally formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table, "D0D5DD")

    # Header row
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        run.font.color.rgb = WHITE
        set_cell_shading(cell, HEADER_BG)
        set_cell_margins(cell)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            run.font.name = "Calibri"
            run.font.color.rgb = DARK_GREY
            if row_idx % 2 == 1:
                set_cell_shading(cell, ALT_ROW_BG)
            set_cell_margins(cell)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def add_callout_box(doc, text, bg_color=ACCENT_BG):
    """Add a highlighted callout/info box."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.font.color.rgb = ACCENT_BLUE
    set_cell_shading(cell, bg_color)
    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
    set_table_borders(table, "B0BEC5")
    return table


def add_bullet(doc, text, bold_prefix=None, indent_level=0):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style="List Bullet")
    if indent_level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * (indent_level + 1))
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = "Calibri"
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = "Calibri"
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = "Calibri"
    return p


def add_para(doc, text, bold=False, italic=False, size=10.5, color=None, align=None):
    """Add a paragraph with formatting."""
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    return p


def create_framework_overview():
    """Create the Framework Overview Word document."""
    doc = Document()

    # ---- Set default font ----
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Configure heading styles
    for level, size, spacing in [(1, 20, 18), (2, 15, 12), (3, 12, 8)]:
        h = doc.styles[f"Heading {level}"]
        h.font.color.rgb = DARK_BLUE
        h.font.size = Pt(size)
        h.font.name = "Calibri"
        h.paragraph_format.space_before = Pt(spacing)
        h.paragraph_format.space_after = Pt(6)

    # Configure List Bullet style
    lb = doc.styles["List Bullet"]
    lb.font.name = "Calibri"
    lb.font.size = Pt(10.5)
    lb.paragraph_format.space_after = Pt(3)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # =====================================================================
    # TITLE PAGE
    # =====================================================================
    for _ in range(5):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Snowflake Cost Optimisation\nFramework")
    run.font.size = Pt(34)
    run.bold = True
    run.font.color.rgb = DARK_BLUE
    run.font.name = "Calibri"

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("A Complete Solution for Cost Visibility,\nAttribution, and Savings")
    run.font.size = Pt(16)
    run.font.color.rgb = LIGHT_BLUE
    run.font.name = "Calibri"

    doc.add_paragraph()
    doc.add_paragraph()

    # Metadata
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Prepared by", "Bilvantis Technologies"),
        ("Date", "April 2026"),
        ("Version", "1.0"),
        ("Classification", "Confidential"),
    ]
    for i, (label, value) in enumerate(meta_data):
        for j, text in enumerate([label, value]):
            cell = meta_table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(11)
            run.font.name = "Calibri"
            if j == 0:
                run.bold = True
                run.font.color.rgb = ACCENT_BLUE
            else:
                run.font.color.rgb = DARK_GREY

    doc.add_page_break()

    # =====================================================================
    # TABLE OF CONTENTS
    # =====================================================================
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. What This Framework Is",
        "2. The Problem We Solve",
        "3. What We Deliver — Ten Core Capabilities",
        "    3.1  Complete Cost Visibility",
        "    3.2  Team-Level Cost Attribution",
        "    3.3  Query Anti-Pattern Detection",
        "    3.4  Proactive Alerting with Microsoft Teams",
        "    3.5  Prioritised Savings Recommendations",
        "    3.6  Cost Forecasting",
        "    3.7  Recommendation Tracking and ROI Verification",
        "    3.8  Seasonality-Aware Anomaly Detection",
        "    3.9  Scheduled Executive Reports",
        "    3.10 Data Freshness Transparency",
        "4. Architecture — How It Works",
        "5. Demo Approach — How We Prove It Works",
        "6. What Makes This Different",
        "7. Delivery Timeline",
        "8. What the Customer Gets",
        "9. Technical Foundation",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = "Calibri"

    doc.add_page_break()

    # =====================================================================
    # 1. WHAT THIS FRAMEWORK IS
    # =====================================================================
    doc.add_heading("1. What This Framework Is", level=1)

    add_para(doc,
        "The Snowflake Cost Optimisation Framework is a fully self-contained analytics system that "
        "lives entirely inside a customer's Snowflake account. It takes the raw metadata that Snowflake "
        "already captures about every query, every warehouse, and every storage object, and transforms it "
        "into clear answers about where money is being spent, why costs are growing, and exactly what to "
        "do about it."
    )

    add_para(doc,
        "There are no external tools, no third-party SaaS subscriptions, and no data leaving the "
        "customer's environment. The entire framework — data models, business logic, dashboards, "
        "alerting, and recommendations — runs natively on Snowflake using dbt for data transformation, "
        "Streamlit for interactive visualisation, and Snowflake Tasks for automation."
    )

    # =====================================================================
    # 2. THE PROBLEM WE SOLVE
    # =====================================================================
    doc.add_heading("2. The Problem We Solve", level=1)

    add_para(doc,
        "Organisations running multiple teams on Snowflake face a recurring set of challenges that "
        "grow more painful as usage scales:"
    )

    problems = [
        ("Nobody knows who is spending what. ",
         "When the monthly Snowflake invoice arrives, finance teams cannot break it down by business "
         "unit, team, or data product. The bill says \"you spent $50,000 on compute\" — but it does "
         "not say which team drove $30,000 of that, or which single query cost $800 to run."),
        ("Warehouses are oversized and idle. ",
         "Teams provision large warehouses \"just in case\" and leave them running long after their "
         "workloads finish. A medium warehouse sitting idle for four hours costs the same as running "
         "queries for four hours — but delivers zero value."),
        ("Inefficient queries run unchecked. ",
         "Without governance, engineers write queries that scan entire tables when they only need a "
         "handful of rows, join tables without proper conditions (creating explosive result sets), or "
         "run the same expensive query dozens of times a day when the results could be cached."),
        ("Storage accumulates silently. ",
         "Tables that no one has read in months continue to occupy storage. Time Travel retention is "
         "set to maximum on tables that do not need it. Clones diverge from their source and quietly "
         "consume dedicated storage."),
        ("Cost management is reactive, not proactive. ",
         "Most organisations review Snowflake costs from monthly invoices — weeks after the spending "
         "has already happened. By the time anyone notices a spike, thousands of dollars may have been wasted."),
    ]
    for bold_part, rest in problems:
        add_bullet(doc, rest, bold_prefix=bold_part)

    doc.add_paragraph()
    add_callout_box(doc,
        "The result: Snowflake costs typically grow 15 to 30 percent quarter-over-quarter without "
        "intervention, with potential savings of 20 to 40 percent left unrealised."
    )

    # =====================================================================
    # 3. WHAT WE DELIVER
    # =====================================================================
    doc.add_heading("3. What We Deliver — Ten Core Capabilities", level=1)

    # --- 3.1 Complete Cost Visibility ---
    doc.add_heading("3.1 Complete Cost Visibility", level=2)
    add_para(doc,
        "We build a comprehensive picture of every dollar spent on Snowflake, broken down across "
        "three cost categories:"
    )

    cost_items = [
        ("Compute costs — ",
         "the credits consumed by warehouses executing queries. We track this at every level of "
         "detail: by warehouse, by team, by user, by individual query, by hour of day, and by day "
         "of week. We calculate the estimated dollar cost of each query based on how long it ran, "
         "how large the warehouse was, and what the customer pays per credit."),
        ("Storage costs — ",
         "the ongoing cost of data at rest. We break this down by database, schema, and individual "
         "table, distinguishing between active data (what is actually in use), Time Travel data "
         "(historical snapshots kept for recovery), and Fail-safe data (Snowflake's built-in "
         "seven-day protection layer). We identify tables that are consuming storage but have not "
         "been read in 90 or more days."),
        ("Serverless costs — ",
         "the credits consumed by Snowflake's automated features: Snowpipe (continuous data loading), "
         "automatic clustering, materialised view refreshes, serverless tasks, and search optimisation. "
         "These costs often go unnoticed because no warehouse is visibly running, but they can add up "
         "significantly."),
    ]
    for bold_part, rest in cost_items:
        add_bullet(doc, rest, bold_prefix=bold_part)

    # --- 3.2 Team-Level Cost Attribution ---
    doc.add_heading("3.2 Team-Level Cost Attribution", level=2)
    add_para(doc,
        "One of the most valuable capabilities is answering \"which team is responsible for which "
        "portion of the bill.\" We derive team ownership from three sources, applied in priority order:"
    )

    attribution_items = [
        ("Query tags — ",
         "structured labels attached to queries by applications or workload generators. When a query "
         "carries a tag like team:analytics, we know exactly which team ran it."),
        ("Role names — ",
         "Snowflake roles often follow naming conventions that map to teams. An ANALYST_ROLE maps to "
         "the Analytics team, a SYSADMIN role maps to the Platform team, and so on."),
        ("Warehouse names — ",
         "when a warehouse is dedicated to a team (such as ETL_WH for Data Engineering or ANALYTICS_WH "
         "for the Analytics team), any query running on that warehouse can be attributed accordingly."),
    ]
    for bold_part, rest in attribution_items:
        add_bullet(doc, rest, bold_prefix=bold_part)

    add_para(doc,
        "This approach works without requiring manual mapping files. The framework dynamically reads "
        "the actual roles, users, and warehouses from Snowflake's own metadata, so it always reflects "
        "the current state of the environment."
    )

    # --- 3.3 Query Anti-Pattern Detection ---
    doc.add_heading("3.3 Query Anti-Pattern Detection", level=2)
    add_para(doc,
        "The framework automatically identifies six categories of wasteful query behaviour:"
    )

    antipatterns = [
        ("Full table scans", "Description", "Queries that read more than 80% of a table's partitions when filters should limit the read. On billion-row tables, this means scanning terabytes unnecessarily."),
        ("SELECT * queries", "Description", "Queries that retrieve every column when only a few are needed. Wastes I/O bandwidth and memory."),
        ("Spill to storage", "Description", "Queries that exhaust the warehouse's in-memory capacity and overflow to disk or remote storage — a signal the query or warehouse size needs attention."),
        ("Repeated identical queries", "Description", "The same query (by parameterised hash) running dozens of times a day. Candidates for result caching or consolidation."),
        ("Cartesian joins", "Description", "Queries where rows produced vastly exceeds rows scanned, indicating a missing or incorrect join condition."),
        ("Large sorts without limits", "Description", "Queries that sort millions of rows with ORDER BY but no LIMIT clause. If only top results are needed, the sort does far more work than necessary."),
    ]

    add_styled_table(doc,
        ["Anti-Pattern", "What It Catches"],
        [[ap[0], ap[2]] for ap in antipatterns],
        col_widths=[4.5, 12.5],
    )

    doc.add_paragraph()
    add_para(doc,
        "Each detected anti-pattern includes the specific query, the user who ran it, the warehouse "
        "it ran on, the estimated cost wasted, and a plain-language recommendation for how to fix it."
    )

    # --- 3.4 Proactive Alerting ---
    doc.add_heading("3.4 Proactive Alerting with Microsoft Teams Integration", level=2)
    add_para(doc,
        "Rather than waiting for someone to open a dashboard and notice a problem, the framework "
        "actively monitors for cost anomalies and sends notifications directly to Microsoft Teams "
        "channels. Six alert types are built in:"
    )

    alerts = [
        ("Daily cost spike — ", "today's spend exceeds the seasonality-adjusted threshold (or 2x the 30-day rolling average as fallback). Something unusual is happening."),
        ("Warehouse idle — ", "a warehouse has been running for 30 or more minutes with zero queries. It is consuming credits for nothing."),
        ("Budget threshold — ", "monthly credit usage has crossed 80 percent of the budget (warning) or 100 percent (critical)."),
        ("Heavy query spill — ", "a query spilled more than one gigabyte to remote storage. This specific query needs attention."),
        ("Storage growth anomaly — ", "a database's storage grew more than 20 percent in a single week. Something unexpected may have been loaded."),
        ("Repeated expensive query — ", "the same query has run more than 20 times today, costing more than a dollar each time. The results should be cached or the pattern should change."),
    ]
    for bold_part, rest in alerts:
        add_bullet(doc, rest, bold_prefix=bold_part)

    doc.add_paragraph()
    add_para(doc,
        "Each alert type can be independently enabled or disabled. The system tracks alert episodes "
        "— if the same warehouse is idle for three consecutive check cycles, it sends one notification "
        "for the episode, not three. When the condition resolves, it logs the resolution. This prevents "
        "alert fatigue while ensuring nothing is missed."
    )
    add_para(doc,
        "The alerting runs on a Snowflake Task that executes every 15 minutes, calling a Python stored "
        "procedure that posts formatted Adaptive Cards to Teams via webhook. The entire pipeline is "
        "native to Snowflake — no external scheduler or middleware is needed."
    )

    # --- 3.5 Prioritised Savings Recommendations ---
    doc.add_heading("3.5 Prioritised Savings Recommendations", level=2)
    add_para(doc,
        "The framework produces a unified, ranked list of every savings opportunity identified across "
        "all three categories (warehouse optimisation, query improvement, storage cleanup). Each "
        "recommendation includes:"
    )

    rec_items = [
        "A plain-language description of what is happening and why it matters",
        "The current monthly cost of the inefficiency",
        "The estimated monthly savings if the recommendation is applied",
        "The effort level required (low for a configuration change, medium for a query rewrite, high for an architecture change)",
        "A confidence level based on the strength of the signal",
        "The actual SQL command to apply the fix, where applicable (such as an ALTER WAREHOUSE statement)",
        "A priority score that ranks recommendations by return on investment — highest savings relative to lowest effort appear first",
    ]
    for item in rec_items:
        add_bullet(doc, item)

    doc.add_paragraph()
    add_callout_box(doc,
        "The Recommendations Hub is the centrepiece of the client presentation. It answers the question "
        "every executive asks: \"How much can we save, and what do we do first?\""
    )

    # --- 3.6 Cost Forecasting ---
    doc.add_heading("3.6 Cost Forecasting", level=2)
    add_para(doc,
        "The framework projects future Snowflake costs using linear regression on the last 90 days of "
        "historical spending. Forecasts are generated at three levels:"
    )

    forecast_items = [
        ("Daily total cost forecast — ",
         "90 days of forward projections with 95% confidence intervals, broken down by compute, "
         "storage, and serverless categories. The regression uses Snowflake's built-in REGR_SLOPE "
         "and REGR_INTERCEPT functions — no external packages required."),
        ("Monthly aggregated forecast — ",
         "daily projections rolled up into monthly totals, presented alongside the last three months "
         "of actuals for visual comparison. A \"projected annual spend\" figure combines year-to-date "
         "actuals with the remaining forecast for the current calendar year."),
        ("Per-team forecast — ",
         "each team's monthly cost trajectory is projected independently, giving finance teams early "
         "warning when a specific team's spend is accelerating. Teams require at least two months of "
         "historical data before projections are generated."),
    ]
    for bold_part, rest in forecast_items:
        add_bullet(doc, rest, bold_prefix=bold_part)

    add_para(doc,
        "The forecasting page shows KPI cards (next month, next quarter, projected annual), an "
        "actuals-vs-forecast line chart with shaded confidence bands, a stacked bar chart of forecast "
        "by cost category, and a filterable team projections table."
    )

    # --- 3.7 Recommendation Tracking and ROI ---
    doc.add_heading("3.7 Recommendation Tracking and ROI Verification", level=2)
    add_para(doc,
        "Recommendations alone are not enough — organisations need to track whether savings were "
        "actually realised. The framework adds a full recommendation lifecycle:"
    )

    roi_items = [
        ("Status tracking — ",
         "each recommendation progresses through states: OPEN, ACCEPTED, IMPLEMENTED, REJECTED, "
         "or DEFERRED. Status is tracked via a seed file that operations teams update as they action "
         "recommendations."),
        ("Actual savings measurement — ",
         "for implemented warehouse recommendations, the framework compares the pre-recommendation "
         "cost against the current 30-day cost of the target warehouse, calculating the real dollar "
         "savings achieved."),
        ("ROI computation — ",
         "actual savings divided by estimated savings gives a concrete ROI percentage for each "
         "implemented recommendation. This proves the framework's value to stakeholders with hard "
         "numbers."),
        ("Conversion funnel — ",
         "an ROI dashboard shows how many recommendations are open, accepted, implemented, rejected, "
         "and deferred, with total estimated and actual savings across all categories."),
    ]
    for bold_part, rest in roi_items:
        add_bullet(doc, rest, bold_prefix=bold_part)

    # --- 3.8 Seasonality-Aware Anomaly Detection ---
    doc.add_heading("3.8 Seasonality-Aware Anomaly Detection", level=2)
    add_para(doc,
        "Simple threshold-based alerting generates false positives. The framework uses a "
        "seasonality-aware approach for cost spike detection:"
    )

    season_items = [
        ("Day-of-week baselines — ",
         "costs vary predictably by weekday (Monday ETL runs cost more than weekend idle time). "
         "The system computes 90-day averages per day of week."),
        ("Month-end adjustment — ",
         "batch processing at month-end naturally spikes costs. The system identifies the last three "
         "days of each month and applies a higher standard deviation threshold to avoid false alarms."),
        ("Trend adjustment — ",
         "if costs are growing steadily (as they do in most organisations), the baseline shifts upward "
         "with the trend. A linear regression on weekly totals adjusts the day-of-week average so that "
         "organic growth does not trigger alerts."),
        ("Z-score detection — ",
         "instead of a fixed \"2x rolling average\" rule, alerts fire when costs exceed the adjusted "
         "baseline by more than 2.0 standard deviations (configurable). This adapts to each "
         "organisation's natural cost patterns."),
        ("Graceful fallback — ",
         "when fewer than 30 days of history exist (new deployments), the system falls back to the "
         "simple multiplier approach until sufficient data accumulates."),
    ]
    for bold_part, rest in season_items:
        add_bullet(doc, rest, bold_prefix=bold_part)

    # --- 3.9 Scheduled Executive Reports ---
    doc.add_heading("3.9 Scheduled Executive Reports", level=2)
    add_para(doc,
        "Executives rarely log into dashboards. The framework delivers cost intelligence proactively "
        "via weekly email summaries:"
    )

    report_items = [
        ("Weekly comparison — ", "this week's cost vs last week's cost, with percentage change and direction indicator."),
        ("Top cost drivers — ", "the three highest-cost warehouses, surfacing where the money is going."),
        ("Top savings opportunities — ", "the three largest unrealised recommendations, showing what could be saved."),
        ("Alert summary — ", "count of new alert episodes in the past seven days."),
        ("Native delivery — ", "reports are sent via Snowflake's built-in SYSTEM$SEND_EMAIL function on a Monday 8 AM UTC schedule, requiring no external email infrastructure."),
    ]
    for bold_part, rest in report_items:
        add_bullet(doc, rest, bold_prefix=bold_part)

    # --- 3.10 Data Freshness Transparency ---
    doc.add_heading("3.10 Data Freshness Transparency", level=2)
    add_para(doc,
        "Snowflake's ACCOUNT_USAGE views have up to 45 minutes of latency. Users viewing a dashboard "
        "need to know whether they are looking at current data or stale data. The framework:"
    )

    freshness_items = [
        ("Monitors six key data sources — ",
         "query history, warehouse metering, storage usage, warehouse load history, login history, "
         "and database storage history. For each source, it computes the age of the most recent record."),
        ("Classifies freshness — ",
         "FRESH (under 30 minutes), STALE (30 to 60 minutes), CRITICAL (over 60 minutes). The overall "
         "status reflects the worst-case source."),
        ("Displays a banner — ",
         "the main dashboard page shows a green, yellow, or red banner with the data timestamp and age. "
         "The Alert Management page shows per-source freshness detail."),
    ]
    for bold_part, rest in freshness_items:
        add_bullet(doc, rest, bold_prefix=bold_part)

    # =====================================================================
    # 4. ARCHITECTURE
    # =====================================================================
    doc.add_page_break()
    doc.add_heading("4. Architecture — How It Works", level=1)
    add_para(doc, "The framework is built in four layers, each serving a distinct purpose:")

    # Layer 1
    doc.add_heading("Layer 1: Staging (15 View Models)", level=2)
    add_para(doc,
        "The staging layer connects to Snowflake's built-in ACCOUNT_USAGE metadata views — 14 views "
        "covering query history, warehouse metering, storage metrics, access history, login history, "
        "and serverless feature usage. These views contain 365 days of historical data at no additional "
        "cost to the customer."
    )
    add_para(doc,
        "The staging models clean and standardise this raw data: column names are normalised, "
        "timestamps are consistently typed, and each model mirrors exactly one source view. No "
        "business logic is applied here — the goal is a clean, reliable foundation."
    )

    # Layer 2
    doc.add_heading("Layer 2: Intermediate — Business Logic (41 Table Models)", level=2)
    add_para(doc,
        "The intermediate layer is where the intelligence lives. Over 40 models (31 intermediate "
        "and 10 alert models) perform cost calculations, utilisation analysis, anti-pattern detection, "
        "anomaly identification, alert state tracking, cost forecasting, seasonality baselines, "
        "recommendation lifecycle management, and data freshness monitoring."
    )
    add_para(doc, "Key computations include:")

    computations = [
        "Converting query execution time and warehouse size into estimated dollar cost per query",
        "Calculating warehouse utilisation as the ratio of active compute to total available capacity",
        "Detecting idle periods where warehouses are running but executing nothing",
        "Identifying query anti-patterns by analysing partition scan ratios, spill volumes, result set sizes, and query frequency",
        "Computing rolling averages and flagging anomalies where current costs exceed historical baselines",
        "Tracking alert episodes to prevent duplicate notifications",
        "Projecting future costs using linear regression on historical trends",
        "Computing seasonality-aware baselines for smarter anomaly detection",
        "Tracking recommendation lifecycle and calculating realised ROI",
        "Monitoring data source freshness to surface staleness to end users",
    ]
    for comp in computations:
        add_bullet(doc, comp)

    # Layer 3
    doc.add_heading("Layer 3: Publication — Dashboard-Ready (16 Table Models)", level=2)
    add_para(doc,
        "The publication layer aggregates and shapes data for consumption by the Streamlit dashboard "
        "and the recommendations engine. These models are pre-computed, optimised for fast dashboard "
        "queries, and designed to answer specific business questions without requiring the dashboard "
        "to perform complex joins or calculations at query time."
    )

    # Layer 4
    doc.add_heading("Layer 4: Presentation — Streamlit Dashboard (12 Pages)", level=2)
    add_para(doc,
        "The dashboard is a twelve-page interactive application running natively inside Snowflake "
        "via Streamlit-in-Snowflake. It requires no external hosting, no separate authentication, "
        "and no data movement. Users access it directly within their Snowflake environment. A data "
        "freshness banner on the main page shows whether the underlying data is current or stale."
    )

    add_para(doc, "The twelve pages are:")

    pages = [
        ("1. Executive Summary", "Total spend, month-over-month trends, compute/storage/serverless split, top warehouses and users by cost, data freshness banner"),
        ("2. Warehouse Deep Dive", "Per-warehouse utilisation, idle time, queue contention, credit consumption, efficiency scoring"),
        ("3. Team Attribution", "Cost broken down by team, drillable to individual users and query types"),
        ("4. Storage Explorer", "Storage by database and table, unused table identification, Time Travel waste highlighting"),
        ("5. Trend Analysis", "90-day cost trends with anomaly flags, day-of-week and hour-of-day heatmaps"),
        ("6. Alert Management", "Active alerts, alert history, configuration status, pipeline health, seasonality-aware detection info, data source freshness"),
        ("7. Warehouse Optimiser", "Right-sizing recommendations with before/after comparisons and SQL commands to apply"),
        ("8. Query Optimiser", "Anti-pattern summary, top optimisation candidates ranked by waste, trend analysis"),
        ("9. Storage Optimiser", "Unused tables, Time Travel waste, transient table candidates with savings estimates"),
        ("10. Recommendations Hub", "Unified savings report with ROI tracking dashboard, conversion funnel, and CSV export"),
        ("11. Cost Forecast", "Linear trend projections with confidence intervals, team forecasts, and projected annual spend"),
        ("12. Report Settings", "Weekly executive report preview, recipient configuration, and manual trigger"),
    ]

    add_styled_table(doc,
        ["Page", "Description"],
        [[p[0], p[1]] for p in pages],
        col_widths=[4.5, 12.5],
    )

    # =====================================================================
    # 5. DEMO APPROACH
    # =====================================================================
    doc.add_page_break()
    doc.add_heading("5. Demo Approach — How We Prove It Works", level=1)
    add_para(doc,
        "A framework that only analyses historical data cannot demonstrate its detection capabilities "
        "convincingly. If we tell a client \"this will catch full table scans,\" they will ask "
        "\"show me.\" To make that possible, we take an active approach:"
    )

    doc.add_heading("Environment Discovery", level=2)
    add_para(doc,
        "Before generating any workload, we scan the actual Snowflake environment to understand what "
        "exists: which databases and tables are available, how large they are, which warehouses are "
        "provisioned, and what roles and users are active. This scan produces a complete inventory "
        "that informs which tables to target and which warehouses to use."
    )

    doc.add_heading("Intentional Workload Generation", level=2)
    add_para(doc,
        "We created a workload generator that executes ten carefully designed scenarios, each "
        "targeting a specific detection capability of the framework:"
    )

    doc.add_heading("Anti-Pattern Scenarios", level=3)
    anti_scenarios = [
        "Scanning 1.3 billion rows of inventory data without filters to trigger the full table scan detector",
        "Running SELECT * against a 65-million-row customer table to trigger the select-star detector",
        "Executing memory-intensive window functions on a small warehouse to force spill-to-storage",
        "Running the same join query 25 times with minor variations to trigger the repeated query detector",
        "Joining two dimension tables without a proper ON condition to create a cartesian product",
        "Sorting a million-row table with ORDER BY but no LIMIT to trigger the large sort detector",
    ]
    for s in anti_scenarios:
        add_bullet(doc, s)

    doc.add_heading("Cost and Attribution Scenarios", level=3)
    cost_scenarios = [
        "Running the same expensive join on three different warehouses (X-Small, Small, Medium) to demonstrate that the framework correctly attributes different costs to different teams based on warehouse assignment",
        "Executing four expensive queries in rapid succession on a single warehouse to create a daily cost spike that triggers the anomaly alert",
        "Resuming a warehouse and running no queries on it to trigger the idle warehouse alert",
    ]
    for s in cost_scenarios:
        add_bullet(doc, s)

    add_para(doc,
        "Every workload query is tagged with a structured label identifying the team, the scenario "
        "name, and a unique run identifier. The dbt models parse these tags to perform team-level "
        "attribution, creating a complete chain from \"this team ran this query\" to \"this query "
        "cost this much\" to \"this anti-pattern was detected.\""
    )

    doc.add_heading("Demo Flow", level=2)
    demo_steps = [
        ("\"Let me run a query\" — ", "execute one of the workload scenarios while the client watches"),
        ("\"Now let's see what the framework caught\" — ", "open the Query Optimiser page and show the detected anti-pattern"),
        ("\"Here's the cost impact\" — ", "switch to the Executive Summary to show how the query affected overall cost"),
        ("\"And here's the recommendation\" — ", "switch to the Recommendations Hub to show the prioritised fix"),
        ("\"This works across teams\" — ", "run the multi-warehouse scenario to show cost attribution across teams"),
        ("\"And it alerts automatically\" — ", "show the Alert Management page where the cost spike triggered a Teams notification"),
    ]
    for i, (bold_part, rest) in enumerate(demo_steps, 1):
        add_bullet(doc, rest, bold_prefix=f"{i}. {bold_part}")

    # =====================================================================
    # 6. WHAT MAKES THIS DIFFERENT
    # =====================================================================
    doc.add_heading("6. What Makes This Different", level=1)

    doc.add_heading("vs. Snowflake's Built-In Cost Tools", level=2)
    add_para(doc,
        "Snowflake provides resource monitors and budgets that can alert when spending crosses a "
        "threshold. But they only tell you that costs are high — not why they are high or what to do "
        "about it. Our framework adds the analytical intelligence layer: root cause analysis, "
        "anti-pattern detection, team attribution, and prioritised recommendations with dollar estimates."
    )

    doc.add_heading("vs. Third-Party SaaS Tools", level=2)
    add_para(doc, "Solutions like Select.dev, Keebo, and Sundeck provide similar functionality, but with significant trade-offs:")

    add_styled_table(doc,
        ["Factor", "Our Framework", "SaaS Tools"],
        [
            ["Data residency", "Stays in customer's Snowflake account", "Data sent to third-party cloud"],
            ["Ongoing cost", "No licence fees (dbt + Streamlit included with Snowflake)", "$500–$5,000+/month"],
            ["Customisation", "Fully tailored to customer's org structure", "Generic, one-size-fits-all"],
            ["Transparency", "All logic visible in dbt SQL models — auditable, extensible", "Black-box recommendations"],
            ["Business context", "Integrates team mappings, query tags, org structure", "Limited to Snowflake metadata"],
            ["Extensibility", "Customer can add new models, metrics, dashboards at any time", "Limited to vendor roadmap"],
        ],
        col_widths=[3.5, 6.5, 6.5],
    )

    doc.add_paragraph()
    doc.add_heading("vs. Manual Analysis", level=2)
    add_para(doc,
        "Many organisations rely on ad hoc SQL queries against ACCOUNT_USAGE views, run by a platform "
        "engineer when someone asks \"why is the bill so high this month?\" This approach is reactive, "
        "inconsistent, and not scalable. Our framework automates the analysis, runs it continuously, "
        "and presents the results in a format that non-technical stakeholders can understand and act on."
    )

    # =====================================================================
    # 7. DELIVERY TIMELINE
    # =====================================================================
    doc.add_heading("7. Delivery Timeline", level=1)

    doc.add_heading("Phase 1: Cost Visibility and Attribution (Weeks 1–4)", level=2)

    add_styled_table(doc,
        ["Week", "Focus", "Key Deliverables"],
        [
            ["Week 1", "Foundation — dbt project setup, 14 source connections, staging layer, configuration seeds", "Clean, standardised data foundation"],
            ["Week 2", "Intelligence — intermediate models for cost attribution, utilisation analysis, anomaly detection; publication models for dashboard", "Working cost models at every level of detail"],
            ["Week 3", "Dashboard — five Streamlit pages: Executive Summary, Warehouse Deep Dive, Team Attribution, Storage Explorer, Trend Analysis", "Interactive dashboard v1 deployed in Snowflake"],
            ["Week 4", "Alerting — six alert types with episode-based deduplication, Teams webhook integration, Snowflake Task automation, Alert Management page", "Fully operational cost visibility with automated alerting"],
        ],
        col_widths=[2.0, 7.0, 7.5],
    )

    doc.add_paragraph()
    doc.add_heading("Phase 2: Query Optimisation and Recommendations (Weeks 5–8)", level=2)

    add_styled_table(doc,
        ["Week", "Focus", "Key Deliverables"],
        [
            ["Week 5", "Warehouse right-sizing — execution time analysis, queue wait analysis, spill rate analysis, hourly utilisation patterns", "Resize, auto-suspend, and scheduling recommendations with dollar estimates"],
            ["Week 6", "Query anti-patterns — six detection models for full scans, SELECT *, spill, repeated queries, cartesian joins, large sorts", "Query-level optimisation candidates with cost impact"],
            ["Week 7", "Storage optimisation — unused tables, Time Travel waste, transient candidates; unified recommendations engine with priority scoring", "Complete prioritised savings report (warehouse + query + storage)"],
            ["Week 8", "Polish — full test suite, cost forecasting, recommendation tracking, executive reports, dashboard refinement, documentation, stakeholder walkthrough", "Production-ready framework with knowledge transfer"],
        ],
        col_widths=[2.0, 7.0, 7.5],
    )

    # =====================================================================
    # 8. WHAT THE CUSTOMER GETS
    # =====================================================================
    doc.add_heading("8. What the Customer Gets", level=1)
    add_para(doc, "At the conclusion of the eight-week engagement, the customer receives:")

    deliverables = [
        ("1. A production-ready dbt project ",
         "with 72 tested, documented data models and 6 configuration seeds that transform Snowflake's "
         "raw metadata into actionable cost intelligence, including cost forecasting, seasonality "
         "baselines, recommendation lifecycle tracking, and data freshness monitoring."),
        ("2. A twelve-page interactive dashboard ",
         "running natively in Snowflake, providing cost visibility from executive summary down to "
         "individual query level, plus cost forecasting with confidence intervals and an ROI tracking "
         "dashboard."),
        ("3. An automated alerting pipeline ",
         "that monitors costs every 15 minutes using seasonality-aware anomaly detection and sends "
         "notifications to Microsoft Teams when anomalies or threshold breaches occur."),
        ("4. A prioritised savings report ",
         "that quantifies every identified optimisation opportunity in dollars, ranks them by effort "
         "and impact, provides the exact SQL to apply each fix, and tracks implementation status with "
         "realised ROI verification."),
        ("5. Cost forecasting ",
         "that projects next month, next quarter, and annual spend using linear regression with 95% "
         "confidence intervals, broken down by cost category and by team."),
        ("6. Weekly executive email reports ",
         "delivered automatically via Snowflake's native email integration, summarising costs, trends, "
         "top drivers, savings opportunities, and alert activity."),
        ("7. A workload generation toolkit ",
         "that can reproduce any anti-pattern or cost scenario on demand, enabling the team to validate "
         "the framework's detection capabilities at any time."),
        ("8. Complete documentation ",
         "including a user guide for the dashboard, a technical guide for extending the framework, "
         "and an operational runbook for maintaining the pipeline."),
        ("9. A reusable accelerator ",
         "that can be deployed to any Snowflake customer environment. The framework dynamically reads "
         "the target environment's metadata and adapts its analysis accordingly, with no manual "
         "configuration beyond credit pricing and alert preferences."),
    ]
    for bold_part, rest in deliverables:
        add_bullet(doc, rest, bold_prefix=bold_part)

    # =====================================================================
    # 9. TECHNICAL FOUNDATION
    # =====================================================================
    doc.add_heading("9. Technical Foundation", level=1)
    add_para(doc,
        "The framework is built on technologies that are already part of the Snowflake ecosystem:"
    )

    add_styled_table(doc,
        ["Technology", "Role", "Why"],
        [
            ["dbt (data build tool)", "Data transformation", "Industry-standard; version control, automated testing, documentation, incremental processing. All business logic expressed as transparent, auditable SQL."],
            ["Streamlit in Snowflake", "Interactive dashboard", "Native to Snowflake; no separate hosting, no additional authentication, no data extraction."],
            ["Snowflake Tasks", "Automation", "Native scheduling — alert pipeline every 15 minutes; full model refresh every 6 hours."],
            ["Snowflake External Access", "Teams integration", "Enables webhook calls to Microsoft Teams, controlled by network rules and secret management."],
            ["SNOWFLAKE.ACCOUNT_USAGE", "Data source", "14 metadata views with 365 days of query history, warehouse metering, storage metrics — available at no additional cost."],
        ],
        col_widths=[4.5, 3.5, 9.0],
    )

    doc.add_paragraph()
    add_callout_box(doc,
        "No additional software licences, external infrastructure, or ongoing subscription fees are required."
    )

    # =====================================================================
    # SUMMARY
    # =====================================================================
    doc.add_page_break()
    doc.add_heading("Summary", level=1)

    add_para(doc,
        "The Snowflake Cost Optimisation Framework transforms Snowflake's built-in metadata from raw "
        "operational data into a strategic asset for cost management. It answers the questions that "
        "matter — who is spending how much, why costs are growing, what is being wasted, and what to "
        "do about it — and delivers those answers through an interactive dashboard, automated alerts, "
        "and a prioritised savings report."
    )
    add_para(doc,
        "It runs entirely inside the customer's Snowflake account, requires no external tools or "
        "ongoing fees, and can be deployed to any Snowflake environment as a reusable accelerator."
    )
    add_para(doc,
        "The framework does not just show what happened. It forecasts what will happen next, detects "
        "problems as they occur using seasonality-aware intelligence, explains why they matter, tells "
        "you exactly how to fix them, tracks whether the fixes worked, and delivers executive "
        "summaries without anyone needing to log in.",
        bold=True
    )

    # ---- Footer ----
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— End of Document —")
    run.italic = True
    run.font.color.rgb = GREY
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Prepared by Bilvantis — Snowflake Cost Optimisation Practice")
    run.font.size = Pt(9)
    run.font.color.rgb = GREY
    run.font.name = "Calibri"

    # Save
    output_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "Snowflake_Cost_Optimisation_Framework_Overview.docx",
    )
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_framework_overview()
