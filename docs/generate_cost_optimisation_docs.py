"""Generate two Word documents for Snowflake Cost Optimisation Framework.

Document 1: Executive Solution Overview (~8 pages)
Document 2: Full Technical Specification (~35-40 pages)

Usage:
    python generate_cost_optimisation_docs.py              # Both documents
    python generate_cost_optimisation_docs.py --exec-only  # Executive overview only
    python generate_cost_optimisation_docs.py --tech-only  # Technical spec only
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import argparse
import os

# ---------------------------------------------------------------------------
# Color palette
# ---------------------------------------------------------------------------
DARK_BLUE = "2E4057"
ACCENT_BLUE = "5A7D9A"
LIGHT_BLUE = "D6E4F0"
GREEN = "27AE60"
RED = "E74C3C"
ORANGE = "F39C12"
LIGHT_GREY = "F2F2F2"
WHITE = "FFFFFF"
DARK_GREY = "4A4A4A"

RGB_DARK_BLUE = RGBColor(0x2E, 0x40, 0x57)
RGB_ACCENT_BLUE = RGBColor(0x5A, 0x7D, 0x9A)
RGB_GREEN = RGBColor(0x27, 0xAE, 0x60)
RGB_RED = RGBColor(0xE7, 0x4C, 0x3C)
RGB_ORANGE = RGBColor(0xF3, 0x9C, 0x12)
RGB_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
RGB_MUTED = RGBColor(0x99, 0x99, 0x99)


# ---------------------------------------------------------------------------
# Shared styling helpers
# ---------------------------------------------------------------------------
def set_cell_shading(cell, color):
    """Set background color for a table cell."""
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), color)
    shading_elm.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_table_borders(table, color="BFBFBF"):
    """Set borders for the entire table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        borders.append(border)
    tblPr.append(borders)


def add_styled_table(doc, headers, rows, col_widths=None, header_color=DARK_BLUE):
    """Add a formatted table with header styling and alternating row shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table)

    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGB_WHITE
        set_cell_shading(cell, header_color)

    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            if row_idx % 2 == 1:
                set_cell_shading(cell, LIGHT_GREY)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def setup_doc_styles(doc):
    """Configure default font and heading styles."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for i in range(1, 4):
        hs = doc.styles[f"Heading {i}"]
        hs.font.color.rgb = RGB_DARK_BLUE


def create_title_page(doc, title_text, subtitle_text, version="2.0"):
    """Create a professional title page."""
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(title_text)
    run.font.size = Pt(32)
    run.bold = True
    run.font.color.rgb = RGB_DARK_BLUE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(subtitle_text)
    run.font.size = Pt(18)
    run.font.color.rgb = RGB_ACCENT_BLUE

    doc.add_paragraph()
    doc.add_paragraph()

    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Prepared by", "Srinivas Addanki"),
        ("Date", "12 April 2026"),
        ("Version", version),
        ("Classification", "Confidential"),
    ]
    for i, (label, value) in enumerate(meta_data):
        meta_table.rows[i].cells[0].text = label
        meta_table.rows[i].cells[1].text = value
        for cell in meta_table.rows[i].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11)
            if cell == meta_table.rows[i].cells[0]:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True

    doc.add_page_break()


def add_callout_box(doc, text, bg_color=LIGHT_BLUE):
    """Add a highlighted callout paragraph with background shading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Apply shading to the paragraph
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), bg_color)
    shd.set(qn("w:val"), "clear")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.bold = True
    return p


def add_ascii_diagram(doc, lines):
    """Add an ASCII art diagram in monospace font."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_kpi_cards(doc, cards):
    """Add KPI cards as a single-row table. cards = [(label, value, color), ...]"""
    table = doc.add_table(rows=2, cols=len(cards))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, "BFBFBF")
    for i, (label, value, color) in enumerate(cards):
        # Value row
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(value)
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor.from_string(color) if isinstance(color, str) else color
        set_cell_shading(cell, LIGHT_GREY)
        # Label row
        cell2 = table.rows[1].cells[i]
        cell2.text = ""
        p2 = cell2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(label)
        run2.font.size = Pt(9)
        run2.bold = True
    return table


def add_phase_timeline(doc, phases):
    """Add a visual phase timeline as a colored table row.
    phases = [(label, weeks, color), ...]
    """
    total_cols = sum(w for _, w, _ in phases)
    table = doc.add_table(rows=2, cols=total_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table, "FFFFFF")

    col = 0
    for label, weeks, color in phases:
        # Merge cells for this phase
        if weeks > 1:
            start_cell = table.rows[0].cells[col]
            end_cell = table.rows[0].cells[col + weeks - 1]
            start_cell.merge(end_cell)
            start_cell2 = table.rows[1].cells[col]
            end_cell2 = table.rows[1].cells[col + weeks - 1]
            start_cell2.merge(end_cell2)

        cell = table.rows[0].cells[col]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGB_WHITE
        set_cell_shading(cell, color)

        cell2 = table.rows[1].cells[col]
        cell2.text = ""
        p2 = cell2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(f"Weeks {col + 1}-{col + weeks}")
        run2.font.size = Pt(8)

        col += weeks

    return table


def add_bold_paragraph(doc, text):
    """Add a bold paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p


def add_doc_footer(doc):
    """Add standard footer to document."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— End of Document —")
    run.italic = True
    run.font.color.rgb = RGB_MUTED

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(
        "This document is confidential and intended for internal use and customer presentation."
    )
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGB_MUTED


def add_sql_block(doc, sql_text):
    """Add a SQL code block in monospace."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F5F5F5")
    shd.set(qn("w:val"), "clear")
    pPr.append(shd)
    run = p.add_run(sql_text)
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


# ---------------------------------------------------------------------------
# PLACEHOLDER: Document functions will be appended below
# ---------------------------------------------------------------------------
def create_executive_overview():
    """Create Document 1: Executive Solution Overview (~8 pages)."""
    doc = Document()
    setup_doc_styles(doc)

    # --- PAGE 1: TITLE ---
    create_title_page(doc, "Snowflake Cost Optimisation\nFramework", "Executive Solution Overview\n\nPrepared for Thomson Reuters")

    # --- PAGE 2: EXECUTIVE SUMMARY ---
    doc.add_heading("Executive Summary", level=1)

    add_callout_box(
        doc,
        "With 350+ source systems feeding billions of records into Snowflake daily, and "
        "multiple business units building their own data products and dbt contracts within "
        "isolated MDS workspaces, Thomson Reuters faces a familiar but acute challenge: "
        "compute costs are growing because users across verticals — with varying levels "
        "of Snowflake expertise — write and execute queries and contracts that are not "
        "optimised for cost efficiency.",
    )

    doc.add_paragraph(
        "Our Snowflake Cost Optimisation Framework transforms your existing Snowflake metadata "
        "into actionable cost intelligence — delivered in three phases across 12 weeks, with "
        "zero additional licence costs. The framework identifies exactly which data products, "
        "dbt contracts, stored procedures, and ad-hoc queries are consuming the most warehouse "
        "resources, pinpoints the root causes, and delivers concrete recommendations — so your "
        "platform team can inform business units precisely what to fix, with the evidence to "
        "drive action."
    )

    doc.add_paragraph()
    add_kpi_cards(doc, [
        ("Typical Waste Identified", "40-60%", RED),
        ("Average Savings Achieved", "20-40%", GREEN),
        ("Additional Licence Cost", "$0", RGB_DARK_BLUE),
    ])

    doc.add_paragraph()
    doc.add_paragraph(
        "The framework follows a proven three-phase approach:"
    )
    phases_overview = [
        ("Phase 1 — Identify & Solve", "Diagnose exactly where money goes and why across all 350+ sources and MDS workspaces. Identify which data products and dbt contracts are eating your warehouses. Deliver a cost dashboard, root cause analysis, and a prioritised solution with recommendations."),
        ("Phase 2 — Optimise", "Implement the optimisations identified in Phase 1 — warehouse right-sizing, query tuning, storage cleanup, and data product contract improvements. Every change backed by before/after evidence."),
        ("Phase 3 — Govern", "Embed ongoing cost governance with automated anomaly alerts, chargeback by business unit, and FinOps review cadence — so savings persist and new contracts are written in an optimised way from day one."),
    ]
    for label, desc in phases_overview:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(label + ": ")
        run.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # --- PAGE 3: THE CHALLENGE ---
    doc.add_heading("The Challenge", level=1)

    doc.add_paragraph(
        "Thomson Reuters operates a sophisticated data platform with data flowing from "
        "350+ source systems through S3, AWS Glue, Fivetran, and Snowpipe into Snowflake. "
        "Across the four data layers — Raw/Role, TRIM (TR Enterprise Model), MDS (business "
        "workspaces), and the Broad/Data Products layer — costs accumulate at every stage. "
        "Based on our introductory discussion, six interconnected challenges stand out:"
    )

    add_styled_table(doc,
        ["#", "Challenge", "Thomson Reuters Context"],
        [
            ["1", "No pinpointed cost attribution", "Platform team knows costs are high but cannot tell each BU exactly which of their dbt jobs, contracts, or queries are driving spend"],
            ["2", "Oversized warehouses", "Warehouses configured for peak loads run at low utilisation during off-hours; suspend/resume tuning is manual and inconsistent"],
            ["3", "Unoptimised dbt contracts & queries", "Business teams writing data products in MDS prioritise delivery over performance — contracts with full scans, no partition awareness, inefficient joins"],
            ["4", "Storage sprawl across layers", "Data retained across Raw, TRIM, MDS, and Broad layers with excessive Time Travel, stale clones, and unused tables"],
            ["5", "Serverless cost growth", "Snowpipe ingestion, auto-clustering, and task execution costs growing with each new source system added"],
            ["6", "Reactive cost management", "Year-on-year platform costs increase; optimisation 'takes a back step when other commitments are priority' (your words)"],
        ],
    )

    doc.add_paragraph()
    add_callout_box(
        doc,
        "Your team has already proven optimisation works: one business unit that prioritised "
        "code improvements showed 'considerable improvement' in cost savings. The challenge "
        "is scaling this across all 350+ sources — you need a tool that pinpoints problems "
        "automatically and gives business teams the exact evidence they need to act.",
    )

    doc.add_paragraph()
    doc.add_heading("Savings Benchmarks by Category", level=2)
    add_styled_table(doc,
        ["Cost Category", "Typical Waste", "Achievable Savings", "Common Root Cause"],
        [
            ["Warehouse Compute", "30-50%", "20-35%", "Oversized, idle, poor auto-suspend"],
            ["Query Efficiency", "15-40%", "10-25%", "Full scans, spills, no caching"],
            ["Storage", "20-30%", "15-25%", "Unused tables, excessive Time Travel"],
            ["Serverless", "10-25%", "10-20%", "Over-clustering, unnecessary pipes"],
            ["Data Transfer", "5-15%", "5-10%", "Cross-region replication, external stages"],
        ],
    )

    doc.add_page_break()

    # --- THE ROOT CAUSE ---
    doc.add_heading("The Root Cause — Uncontrolled Data Product Development", level=1)

    doc.add_paragraph(
        "The six challenges on the previous page are symptoms. The underlying root cause is "
        "structural: your MDS workspaces give business units the freedom to create data products "
        "and dbt contracts — but without guardrails on how those contracts are written. Business "
        "teams prioritise delivery over performance. Not everybody is aware of partition columns, "
        "filter strategies, or warehouse sizing best practices. The result: contracts that work "
        "correctly but consume far more compute than necessary."
    )

    add_callout_box(
        doc,
        "As discussed in our introductory call: 'We can't go through every single source data "
        "system and check their code to check the performance of their queries to analyse it.' "
        "You need a tool that does this automatically — pinpointing exactly which data products "
        "and contracts are expensive, why they're expensive, and how to fix them. That's what "
        "Phase 1 delivers.",
    )

    doc.add_paragraph()
    doc.add_heading("Thomson Reuters: How the Cost Multiplier Works", level=2)
    add_ascii_diagram(doc, [
        " Your Current State:                     After This Engagement:",
        "",
        " S3 -> Raw/Role Layer                    S3 -> Raw/Role Layer",
        "     |                                       |",
        "     v                                       v",
        " TRIM (TR Enterprise Model)              TRIM (TR Enterprise Model)",
        "     |                                       |",
        "     v                                       v",
        " MDS Workspace (per BU)                  MDS Workspace (per BU)",
        "  +-- Team A: complex queries ------+    +-- Optimised contracts ------+",
        "  +-- Team B: different approach ---+    +-- Best-practice templates --+",
        "  +-- Team C: full table scans -----+    +-- Cost-aware development --+",
        "  +-- Team D: no cost awareness ----+    |",
        "     |                                    v",
        "     v                               Data Products (Broad Layer)",
        " Data Products (Broad Layer)           +-- Reusable across BUs",
        "  +-- Inconsistent quality               +-- Quality assured",
        "  +-- High compute cost                  +-- Cost-efficient",
        "  +-- No visibility into impact          +-- Full cost attribution",
    ])

    doc.add_paragraph()
    doc.add_heading("The Impact Across Your User Types", level=2)
    add_styled_table(doc,
        ["User Type", "What They Do at TR Today", "The Problem", "What They Actually Need"],
        [
            [
                "Platform / Data Engineers\n(your team)",
                "Manage ingestion from 350+ sources, build TRIM models, configure warehouses, step in to help BUs with sizing and suspend times",
                "Cannot go through every source system's code to check performance — bandwidth limited, no automated detection",
                "A tool that automatically identifies expensive code and generates evidence-based recommendations to send to BU teams",
            ],
            [
                "BU Technical Users\n(MDS workspace users)",
                "Write dbt contracts and data products in MDS, create transforms, build models for their vertical",
                "Prioritise delivery over performance — may not know partition columns, filter strategies, or optimal join patterns. '$100 query that an expert writes for $10'",
                "Automated feedback on contract performance + optimised contract templates that follow TR guidelines",
            ],
            [
                "Business / Reporting Users",
                "Run reports, build dashboards, query data products in the Broad layer, request ad-hoc data pulls",
                "Depend on BU teams for data access — may write ad-hoc queries without understanding cost impact",
                "Well-curated data products in the Broad layer + self-service dashboards — no raw table access needed",
            ],
        ],
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "Our framework addresses the three needs identified in our discussion: "
        "(1) Identify which data products and contracts are eating your warehouses, "
        "(2) Optimise existing contracts with concrete before/after evidence, and "
        "(3) Establish a framework so new contracts are written in an optimised way from day "
        "one — breaking the cycle of optimise-then-regress."
    )

    doc.add_page_break()

    # --- PAGES: OUR SOLUTION ---
    doc.add_heading("Our Solution", level=1)

    doc.add_heading("How It Fits Into Your Architecture", level=2)
    add_ascii_diagram(doc, [
        "  YOUR EXISTING DATA FLOW:                 OUR FRAMEWORK:",
        "",
        "  350+ Sources                              SNOWFLAKE METADATA",
        "      |                                     +-------------------+",
        "  S3 + Glue/Fivetran/Snowpipe               | ACCOUNT_USAGE    |",
        "      |                                      | METERING_HISTORY |",
        "      v                                      | QUERY_ATTRIBUTION|",
        "  +--------+  +------+  +-----+  +------+   +-------------------+",
        "  |Raw/Role|->| TRIM |->| MDS |->| Broad|          |",
        "  | Layer  |  | (TR  |  |(per |  |(Data |          v",
        "  |        |  | Ent. |  | BU) |  | Prod)|   +-------------------+",
        "  +--------+  |Model)|  +-----+  +------+   | dbt Cost Models  |",
        "               +------+                      | (20+ models)     |",
        "                                             +-------------------+",
        "  Governance: Immuta + Alation + SailPoint          |",
        "                                                    v",
        "                                             +-------------------+",
        "                                             | Streamlit         |",
        "                                             | Cost Dashboard    |",
        "                                             | + Alert System    |",
        "                                             +-------------------+",
        "",
        "  No data leaves your Snowflake account.  $0 licence cost.",
    ])

    doc.add_paragraph()
    doc.add_heading("Three-Phase Approach", level=2)
    add_styled_table(doc,
        ["Phase", "Focus", "Duration", "Key Outcome"],
        [
            ["Phase 1: Identify & Solve", "Diagnose cost drivers, map user query patterns,\ndeliver cost dashboard + findings report\nwith root causes and solution recommendations", "Weeks 1-4", "You know exactly what's wrong and have a concrete plan to fix it"],
            ["Phase 2: Optimise", "Implement deeper optimisations from Phase 1:\nwarehouse right-sizing, query tuning,\nstorage cleanup, data product POCs", "Weeks 5-8", "Measurable savings delivered with before/after proof"],
            ["Phase 3: Govern", "Anomaly detection, chargeback model,\nFinOps review cadence, drift prevention", "Weeks 9-12", "Self-sustaining governance so savings persist long-term"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("Six Key Capabilities", level=2)
    capabilities = [
        "Unified Cost View — All cost categories in one dashboard: compute, storage, serverless, data transfer, and more",
        "Per-Query Cost Attribution — Exact dollar cost per query, per job, per data product — not just warehouse-level totals",
        "Root Cause Analysis — Understand why each expensive query or contract costs what it does, with actionable fix recommendations",
        "Automated Anomaly Detection — Cost spikes detected and alerted before they hit the invoice",
        "Tag-Based Chargeback — Cost ownership by team, vertical, and data product — each BU manages their own budget",
        "FinOps Maturity Assessment — Structured progression from reactive to proactive cost governance",
    ]
    for cap in capabilities:
        doc.add_paragraph(cap, style="List Bullet")

    doc.add_page_break()

    # --- PAGES 5-7: WHAT YOU GET (Phase Deep Dives) ---
    doc.add_heading("What You Get", level=1)
    doc.add_paragraph(
        "Each phase delivers tangible business outcomes — not just technical artifacts. "
        "Here is what changes for your organisation after each phase."
    )

    # ---- PHASE 1 ----
    doc.add_heading("Phase 1: Identify & Solve — Your Immediate Priority", level=2)
    doc.add_paragraph(
        "This is what you asked for in our introductory call: 'Identify what are the list "
        "of data products that have complex queries which are eating my warehouse.' Phase 1 "
        "does exactly that — and goes further. We scan every warehouse, every dbt job, every "
        "stored procedure, and every ad-hoc query across all your MDS workspaces. We deliver "
        "a complete diagnosis with evidence-based recommendations that your platform team can "
        "send directly to business unit owners — showing them exactly what's expensive, why, "
        "and how to fix it."
    )

    doc.add_heading("What We Deliver and Why It Matters", level=3)
    add_styled_table(doc,
        ["What We Deliver", "Business Questions You Can Now Answer", "Business Advantage"],
        [
            [
                "Interactive cost dashboard\n(compute, storage, serverless,\ncost by BU / MDS workspace)",
                "What is our total spend across all layers?\nWhich BU / MDS workspace drives the most cost?",
                "Single source of truth — shareable across teams. The 'matrix' your team asked to see.",
            ],
            [
                "Data product & contract\ncost ranking\n(the list you asked for)",
                "Which data products and dbt contracts\nare eating my warehouse?\nHow much does each one cost?",
                "Exactly what you need: a ranked list of expensive data products with $ attribution — evidence to drive action with BU owners",
            ],
            [
                "Root cause analysis per\nexpensive contract\n(why, not just what)",
                "Why is this contract expensive?\nIs it full scans, bad joins, missing filters,\nor oversized warehouse?",
                "Actionable detail — your team can email BU owners with: 'Here is the problem, here is the fix, here is the potential saving'",
            ],
            [
                "Automated alert framework\n(cost spike notifications)",
                "Did something unusual happen today?\nWhich warehouse or job spiked?",
                "The notification system you asked for — auto-alerts when costs spike, sent via email / Slack / Teams so BU teams are informed immediately",
            ],
            [
                "Tiered Optimisation\nRecommendations\n(ranked by ROI)",
                "What are users actually trying to get?\nWhat's the fastest fix for each\nexpensive pattern?",
                "Every expensive pattern is assigned the right-sized fix — from quick contract patches to curated data products. 80% of savings come from fixes that deploy in minutes. Each recommendation includes cost-benefit analysis.",
            ],
            [
                "Prioritised findings report\nwith quick-win recommendations",
                "What should we fix this week?\nWhat requires deeper Phase 2 work?",
                "Immediate action — warehouse auto-suspend, sizing quick wins can be applied within days",
            ],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("Our Approach: Speed of Fix Matters", level=3)
    doc.add_paragraph(
        "Not every problem needs a complex solution. Our framework assigns the right-sized fix "
        "to each problem — from quick contract patches to full curated data products. The key "
        "insight: 80% of savings come from fixes that deploy in minutes, not days."
    )
    add_styled_table(doc,
        ["Approach", "Speed", "Impact"],
        [
            ["Quick contract fixes\n(targeted improvements to existing code)", "Minutes", "Addresses the top-20 most expensive contracts immediately — send evidence-based fix to BU owner"],
            ["Platform configuration\n(leverage built-in Snowflake capabilities)", "Minutes", "Zero-maintenance improvements — Snowflake handles the optimisation natively"],
            ["Shared data access layers\n(reduce duplicate effort across BUs)", "Hours", "Multiple users querying the same data in different ways → one optimised access path"],
            ["Full curated data products\n(only where ROI clearly justifies)", "Days", "Reserved for top 5-10 high-value cross-BU patterns — we prove ROI before investing build time"],
        ],
        header_color=DARK_BLUE,
    )
    doc.add_paragraph(
        "Phase 1 identifies which approach applies to each problem. Phase 2 implements them "
        "in priority order — quick wins first, bigger investments only where proven."
    )

    doc.add_paragraph()
    add_callout_box(
        doc,
        "Phase 1 Outcome: You don't just see your costs — you understand exactly why they're "
        "high and have a concrete, prioritised plan to fix them. Quick wins can be applied "
        "within days for 5-15% savings. Larger investments are recommended only where ROI "
        "is proven. Phase 2 implements the full plan with before/after evidence.",
    )

    doc.add_paragraph()
    # ---- PHASE 2 ----
    doc.add_heading("Phase 2: Optimise — Implement the Savings", level=2)
    doc.add_paragraph(
        "Phase 1 identified the problems and ranked them by cost impact. Phase 2 implements "
        "the deeper optimisations — reverse-engineering expensive dbt contracts to create "
        "optimised versions, right-sizing warehouses, cleaning up storage, and demonstrating "
        "before/after savings. For each optimisation, we provide the evidence: 'This contract "
        "used to cost $X/month, now it costs $Y/month — here's what changed.'"
    )

    doc.add_heading("What We Deliver and Why It Matters", level=3)
    add_styled_table(doc,
        ["What We Deliver", "Business Questions You Can Now Answer", "Business Advantage"],
        [
            [
                "Warehouse right-sizing\nexecution with before/after testing",
                "What size should each warehouse actually be?\nWhat's the verified saving after resize?",
                "Largest single saving — 20-35% of compute cost. We test in sandbox before applying to production.",
            ],
            [
                "Query-level optimisation\n(anti-pattern fixes, QAS evaluation)",
                "Which specific queries should be rewritten?\nWould Query Acceleration Service help?",
                "Targeted fixes — the top 10 expensive queries often account for 15-25% of compute spend",
            ],
            [
                "Storage cleanup execution\n(unused tables, TT reduction, clone removal)",
                "Which tables can we safely drop or archive?\nWhat's the verified storage saving?",
                "Immediate savings — unused tables and excessive Time Travel are pure waste",
            ],
            [
                "Per-query cost attribution\n(exact $ per query via QUERY_ATTRIBUTION)",
                "What does each scheduled job actually cost us?\nWhich pipelines have the worst ROI?",
                "Data-driven decisions — retire or optimise jobs based on actual cost, not guesswork",
            ],
            [
                "Prioritised fix implementation\n(ranked by speed and savings)",
                "Which fixes deliver the\nbiggest savings fastest?\nWhat's the verified before/after?",
                "Quick wins deployed first for immediate impact. Larger investments only where ROI is proven — every change backed by before/after evidence.",
            ],
            [
                "Optimised contract templates\nfor new data products",
                "Can new dbt contracts be\noptimised from day one?\nHow do we prevent future waste?",
                "New data products are built using optimised templates that follow best practices — so every new contract is efficient by default, without requiring deep Snowflake expertise from BU teams.",
            ],
        ],
    )

    doc.add_paragraph()
    add_callout_box(
        doc,
        "Phase 2 Outcome: Measurable, verified savings with before/after proof. Warehouse "
        "resizing, query fixes, and storage cleanup are implemented and validated. Quick wins "
        "delivered first, bigger investments only where ROI is proven. Optimised contract "
        "templates ensure every new data product is cost-efficient from day one.",
    )

    doc.add_page_break()

    # ---- HOW NEW DATA PRODUCTS CHANGE ----
    doc.add_heading("How New Data Products Enter the System — Before and After", level=2)
    doc.add_paragraph(
        "The biggest long-term impact of this framework is not fixing today's expensive queries — "
        "it's changing how new data products are created going forward. Today, there is no gate "
        "between 'BU team writes a dbt contract' and 'contract runs in production consuming credits.'"
    )

    add_styled_table(doc,
        ["", "Today", "After This Engagement"],
        [
            [
                "Before creating\na data product",
                "No check: does this data already exist?\nNo cost estimate. No quality gate.",
                "Automatic overlap check against existing contracts.\nCost estimate before deployment.\nOptimised template provided.",
            ],
            [
                "Writing the\ncontract / query",
                "Each BU writes their own SQL.\nVarying expertise levels.\nNo best-practice enforcement.",
                "Optimised contract templates follow TR guidelines.\nBU teams get efficient code without needing\ndeep Snowflake expertise.",
            ],
            [
                "After deployment",
                "No visibility until invoice.\nNo alert if costs spike.\nAnother BU may build the same thing.",
                "Continuous cost tracking per contract.\nAnomalies detected within hours.\nDuplicate effort flagged automatically.",
            ],
            [
                "Cross-BU sharing",
                "Each BU works in isolation.\nSame data built multiple times\nacross MDS workspaces.",
                "Shared data products in Broad layer.\nOne contract serves multiple BUs.\nCost shared, quality assured.",
            ],
        ],
        header_color=DARK_BLUE,
    )

    doc.add_paragraph()
    add_callout_box(
        doc,
        "The key shift: users still have the freedom to create new data products — but the "
        "system now guides them toward efficient, governed, reusable outcomes. Every new "
        "contract is cost-estimated before deployment and checked for overlap. This is how "
        "you stop the cycle of 'optimise, then someone writes something unoptimised again' "
        "— without slowing anyone down.",
    )

    doc.add_page_break()

    # ---- PHASE 3 ----
    doc.add_heading("Phase 3: Govern — Stop the Cycle of Regress", level=2)
    doc.add_paragraph(
        "This addresses your third concern: 'How do I give users the luxury of still writing "
        "contracts and products, but with auto-optimisation behind the screen — sticking to "
        "TR guidelines and policies?' Phase 3 embeds cost-aware governance so that new data "
        "products are written in an optimised way from day one, anomalies are caught "
        "automatically, and each BU owns their cost budget."
    )

    doc.add_heading("What We Deliver and Why It Matters", level=3)
    add_styled_table(doc,
        ["What We Deliver", "Business Questions You Can Now Answer", "Business Advantage"],
        [
            [
                "Automated anomaly detection\nand alerting",
                "Did something unusual happen with our costs today?\nIs this spike expected or a problem?",
                "Cost incidents caught in hours, not weeks — auto-alerts via email/Slack/Teams before runaway spend hits the invoice",
            ],
            [
                "Chargeback by team\nand vertical",
                "How much is each team/vertical spending vs their budget?\nWhich cost centre is over/under?",
                "True cost ownership — each vertical manages their spend like any other budget line item",
            ],
            [
                "FinOps review cadence\n(weekly/monthly/quarterly)",
                "Are we maintaining our savings?\nWhat new optimisation opportunities have appeared?",
                "Structured reviews prevent cost drift and catch new waste early — your team runs this independently after handover",
            ],
            [
                "Cost gate for new\ndata products",
                "What will this new contract cost?\nDoes something similar already exist?",
                "The gate you're missing today — new contracts are cost-estimated and checked for overlap before production deployment",
            ],
            [
                "Duplicate effort prevention",
                "Are BUs building the same data products\nin isolation? Can we consolidate?",
                "Automatic detection of overlapping contracts across BUs — prevents wasted compute from 'N teams, N copies' problem",
            ],
            [
                "Cost drift detection",
                "Has a warehouse been resized without approval?\nAre costs creeping back up?",
                "Configuration changes that cause cost regression are flagged automatically — policy enforcement without manual review",
            ],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("Query Governance by User Type", level=3)
    doc.add_paragraph(
        "Different user groups at Thomson Reuters need different governance approaches — "
        "aligned with how Immuta already enforces policy-based access control."
    )
    add_styled_table(doc,
        ["User Type", "Governance Approach", "What We Deliver"],
        [
            [
                "Platform Team\n(your engineers)",
                "Automated detection, cost dashboards, alert management",
                "Framework that automatically flags expensive code across all MDS workspaces + evidence reports to send to BU owners",
            ],
            [
                "BU Technical Users\n(MDS workspace users)",
                "Optimised dbt contract templates, cost feedback loops",
                "TR-specific contract guidelines + automated cost scoring for new data products before they go to production",
            ],
            [
                "Business / Reporting Users",
                "Self-service via curated data products in Broad layer",
                "Governed data products with Alation cataloguing + Streamlit dashboards for common analyses",
            ],
        ],
    )

    doc.add_paragraph()
    add_callout_box(
        doc,
        "Phase 3 Outcome: Cost governance runs on autopilot. Anomalies are caught within hours, "
        "each vertical owns their budget, and regular reviews ensure savings persist. Every user "
        "type — from engineers to business users — has an appropriate, cost-efficient way to access data. "
        "The framework becomes a self-sustaining capability your team owns entirely.",
    )

    # ---- CUMULATIVE VALUE ----
    doc.add_paragraph()
    doc.add_heading("Cumulative Value: How the Phases Build on Each Other", level=2)
    add_styled_table(doc,
        ["", "Before Engagement", "After Phase 1\n(Identify & Solve)", "After Phase 2\n(Optimise)", "After Phase 3\n(Govern)"],
        [
            [
                "Problem Understanding",
                "Don't know where money goes or why",
                "Full diagnosis: cost drivers, root causes, user patterns, data product opportunities",
                "Validated with before/after proof",
                "Continuously monitored with anomaly detection",
            ],
            [
                "Savings",
                "None",
                "Quick wins applied (5-15%)\nStrategic roadmap for long-term savings",
                "Deeper optimisations implemented (20-40% verified)",
                "Savings sustained via governance + drift prevention",
            ],
            [
                "User Access",
                "Everyone queries raw tables",
                "User patterns mapped.\nTiered fix recommendations ranked by ROI.",
                "Tiered fixes deployed + optimised contract templates in use",
                "Each user type has cost-efficient access path",
            ],
            [
                "Decision Making",
                "Gut feel and invoice review",
                "Data-driven: dashboard + findings + solution plan",
                "ROI-proven: before/after evidence for every change",
                "Budget-managed with automated governance per vertical",
            ],
            [
                "Sustainability",
                "N/A",
                "Solution plan in hand",
                "Optimisations implemented and validated",
                "Self-sustaining — continuous improvement loop",
            ],
        ],
        header_color=DARK_BLUE,
    )

    doc.add_page_break()

    doc.add_heading("Executive Dashboard Preview", level=2)
    doc.add_paragraph(
        "This is an example of the cost dashboard we deliver in Phase 1 — the 'matrix' "
        "your team asked to see, shareable across BU owners:"
    )
    add_ascii_diagram(doc, [
        "+-------------------------------+-------------------------------+",
        "| TOTAL MONTHLY SPEND           | COST TREND (90 DAYS)          |",
        "|  $124,500 (+8% MoM)           |  ___                          |",
        "|                               | /   \\___    ___/              |",
        "|  Compute:  $89,200 (72%)      |/        \\__/                  |",
        "|  Storage:  $22,100 (18%)      |                               |",
        "|  Svrless:  $13,200 (10%)      | Jan  Feb  Mar  Apr            |",
        "+-------------------------------+-------------------------------+",
        "| TOP 5 COST DRIVERS (data      | SAVINGS OPPORTUNITIES         |",
        "| products + dbt jobs)           |                               |",
        "|  1. [BU-A] MDS contract  $31K |  Warehouse sizing:  $18,500   |",
        "|  2. [BU-B] ETL pipeline  $25K |  Contract optim.:   $12,200   |",
        "|  3. [BU-C] Stored proc   $16K |  Storage cleanup:    $6,800   |",
        "|  4. [BU-A] Ad-hoc queries $11K|  Serverless tuning:  $3,100   |",
        "|  5. [BU-D] Reporting      $6K |                               |",
        "+-------------------------------+-------------------------------+",
    ])

    doc.add_page_break()

    # --- PAGES 6-7: WHY THIS APPROACH ---
    doc.add_heading("Why This Approach", level=1)

    doc.add_heading("Comparison: Three Approaches to Snowflake Cost Management", level=2)
    add_styled_table(doc,
        ["Capability", "Our Framework", "Native Snowflake", "SaaS Tools"],
        [
            ["Cost attribution", "Full: compute + storage + serverless + per-query", "Basic: budgets and resource monitors", "Good: varies by vendor"],
            ["Query-level costing", "Yes — QUERY_ATTRIBUTION_HISTORY", "No", "Partial"],
            ["Query plan analysis", "Yes — GET_QUERY_OPERATOR_STATS()", "Manual via EXPLAIN", "Limited"],
            ["Anomaly detection", "Custom + Snowflake ANOMALY functions", "Threshold alerts only", "Yes"],
            ["Chargeback / tags", "Full tag-based allocation", "Manual tag setup only", "Partial"],
            ["Customisation", "Fully tailored to your org", "One-size-fits-all", "Limited to vendor UI"],
            ["Data residency", "Stays in your Snowflake", "In your Snowflake", "Sent to third-party"],
            ["Ongoing licence cost", "$0 (dbt + Streamlit included)", "$0", "$500-$5,000+/month"],
            ["Transparency", "All logic visible in dbt SQL", "Black-box", "Black-box"],
            ["Extensibility", "Add models, metrics, dashboards", "Limited", "Vendor roadmap only"],
        ],
    )

    doc.add_paragraph()
    add_callout_box(
        doc,
        "Key differentiator: Our framework combines the depth of custom engineering with "
        "the speed of open-source accelerators (dbt-snowflake-monitoring, Sundeck OpsCenter) "
        "— delivering enterprise-grade cost governance at zero ongoing licence cost.",
    )

    doc.add_paragraph()
    doc.add_heading("Proven at Scale: Our Production Data Product Platform", level=2)
    doc.add_paragraph(
        "We don't just recommend data products — we build and operate them in production. "
        "Our team runs a production-grade dbt + Snowflake + Streamlit platform that already "
        "demonstrates the exact capabilities we're proposing:"
    )
    add_styled_table(doc,
        ["What We Operate", "Scale", "Relevance to This Engagement"],
        [
            ["Data Products in Production", "10+ across multiple domains", "Proven three-layer architecture (staging → intermediate → publication)"],
            ["Streamlit Dashboards", "13 production dashboards", "Same approach for cost attribution and governance dashboards"],
            ["Automated Alert Framework", "30+ alert types with episode tracking", "Foundation for cost anomaly detection and drift alerting"],
            ["External Integrations", "Datadog, Incident.io, Jira — pull-based", "Zero-maintenance integration with your monitoring stack"],
            ["Data Quality Framework", "7 dimensions, RAG scoring, SLOs", "Directly applicable to cost data quality and governance"],
            ["Terraform Governance", "4 environments, per-product roles", "Enterprise-grade access control and schema management"],
            ["Data Contracts", "Versioned public models, 30-day notice", "Proven contract management for multi-team consumption"],
            ["Refresh Cadence", "15-minute cycle, <10 min execution", "Near-real-time cost monitoring capability"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("Your Data Maturity Journey", level=2)
    doc.add_paragraph(
        "Thomson Reuters is already transitioning: your data product concept (MDS → Broad layer) "
        "is in the 'nascent stage' as your team described. This engagement accelerates the "
        "journey by adding cost governance and optimisation guardrails:"
    )
    add_styled_table(doc,
        ["", "Era 1: Pipelines\n(Parts of Current State)", "Era 2: Data Products\n(Where You're Heading)", "Era 3: Intelligent Products\n(Future Vision)"],
        [
            ["What you build", "ETL jobs, batch loads, raw tables", "Curated datasets with owners, SLAs, contracts, quality rules", "Self-adapting products that detect drift and auto-remediate"],
            ["How users get data", "'Can you pull this data?' — tickets and ad-hoc SQL", "Self-service via catalogues, dashboards, governed access", "Data finds the user — personalised, proactive delivery"],
            ["Cost behaviour", "Uncontrolled — every user queries raw tables", "Managed — curated layers reduce compute 10-50x", "Optimised — intelligent caching, auto-sizing, anomaly prevention"],
            ["Quality approach", "Manual checks, customer complaints", "Automated testing, SLOs, RAG dashboards", "ML-powered anomaly detection, self-healing pipelines"],
        ],
        header_color=DARK_BLUE,
    )
    doc.add_paragraph(
        "This engagement delivers the Era 2 foundation: cost visibility, optimised access patterns, "
        "and data product recommendations. The governance framework in Phase 3 sets the stage for "
        "the Era 3 capabilities as your platform matures."
    )

    doc.add_page_break()

    # --- DELIVERY TIMELINE ---
    doc.add_heading("Delivery Timeline", level=1)

    add_phase_timeline(doc, [
        ("Phase 1: Identify & Solve", 4, DARK_BLUE),
        ("Phase 2: Optimise", 4, ACCENT_BLUE),
        ("Phase 3: Govern", 4, GREEN),
    ])

    doc.add_paragraph()
    doc.add_heading("Key Milestones", level=2)
    add_styled_table(doc,
        ["Milestone", "Week", "Deliverable"],
        [
            ["Discovery & Quick Wins", "Week 1", "Account profile, user landscape, initial findings, quick-win recommendations applied"],
            ["Diagnosis & Dashboard", "Week 3", "Interactive cost dashboard + root cause analysis + Data Product Opportunity Report"],
            ["Solution Delivered", "Week 4", "Prioritised findings report with concrete solution plan for Phases 2 & 3"],
            ["Savings Verified", "Week 8", "Before/after proof of optimisations: warehouse resizing, query fixes, storage cleanup"],
            ["Governance Active", "Week 11", "Anomaly detection, chargeback, FinOps cadence, drift prevention in place"],
        ],
    )

    doc.add_page_break()

    # --- PAGES 7-8: INVESTMENT & NEXT STEPS ---
    doc.add_heading("Investment & Next Steps", level=1)

    doc.add_heading("Team Composition", level=2)
    add_styled_table(doc,
        ["Role", "Allocation", "Phases"],
        [
            ["Lead Engineer (Snowflake + dbt)", "Full-time", "All phases"],
            ["Dashboard Developer (Streamlit)", "Part-time", "Weeks 3-4, 7-8, 11-12"],
            ["Project Lead", "Part-time", "Throughout"],
            ["FinOps Consultant", "Part-time", "Phase 3 (Weeks 9-12)"],
        ],
    )

    doc.add_paragraph()
    add_callout_box(
        doc,
        "$0 ongoing licence cost. The framework runs entirely on Snowflake-native capabilities "
        "(ACCOUNT_USAGE, dbt, Streamlit) — no third-party SaaS subscriptions required.",
        LIGHT_BLUE,
    )

    doc.add_paragraph()
    doc.add_heading("Next Steps", level=2)
    doc.add_paragraph(
        "Based on our introductory call of 9 April 2026, both teams agreed on the following path:"
    )
    add_styled_table(doc,
        ["#", "Action", "Owner", "Timeline"],
        [
            ["1", "Share observability demo / dashboard screenshots\n(showing what the cost visibility framework looks like)", "Bilvantis", "This week"],
            ["2", "Review this proposal and confirm engagement scope", "Thomson Reuters", "By 18 April 2026"],
            ["3", "Arrange access to Snowflake instance for POC/pilot\n(or share non-sensitive metadata export)", "Thomson Reuters", "Week of 21 April"],
            ["4", "Begin Phase 1: Identify & Solve\n(deploy framework on TR's Snowflake instance)", "Bilvantis delivery team", "Upon access"],
        ],
    )

    doc.add_paragraph()
    add_callout_box(
        doc,
        "As agreed in the call: we will demonstrate our observability capabilities with "
        "screenshots and a focused demo. Once Thomson Reuters is satisfied with the approach, "
        "we begin with Phase 1 on your Snowflake instance — delivering the data product cost "
        "ranking and automated alerting you asked for within the first 4 weeks.",
        LIGHT_BLUE,
    )

    add_doc_footer(doc)

    output_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "Snowflake_Cost_Optimisation_Executive_Overview.docx",
    )
    doc.save(output_path)
    print(f"Executive Overview saved to: {output_path}")
    return output_path


def _build_tech_spec_sections_8_to_15(doc):
    """Build sections 8-15 of the technical specification."""

    # === SECTION 8: EDITION CONSIDERATIONS ===
    doc.add_heading("8. Edition Considerations", level=1)
    doc.add_paragraph(
        "Feature availability varies by Snowflake edition. Our framework adapts to the "
        "customer's edition with fallback strategies for Standard Edition."
    )
    add_styled_table(doc,
        ["Feature", "Standard", "Enterprise", "Business Critical", "Used In"],
        [
            ["ACCOUNT_USAGE views", "Yes", "Yes", "Yes", "All phases"],
            ["QUERY_ATTRIBUTION_HISTORY", "No", "Yes", "Yes", "Phase 1-2 (per-query $)"],
            ["ACCESS_HISTORY", "No", "Yes", "Yes", "Phase 1 (unused tables)"],
            ["WAREHOUSE_EVENTS_HISTORY", "No", "Yes", "Yes", "Phase 2 (WH analysis)"],
            ["Object tags", "Yes", "Yes", "Yes", "Phase 3 (chargeback)"],
            ["Budgets", "No", "Yes", "Yes", "Phase 3 (governance)"],
            ["ANOMALY_DETECTION ML", "No", "Enterprise+", "Yes", "Phase 3 (anomalies)"],
            ["Query Acceleration Service", "No", "Yes", "Yes", "Phase 2 (QAS)"],
            ["Multi-cluster warehouses", "No", "Yes", "Yes", "Phase 2 (sizing)"],
            ["SYSTEM$CLUSTERING_*", "Yes", "Yes", "Yes", "Phase 2 (clustering)"],
            ["GET_QUERY_OPERATOR_STATS", "Yes", "Yes", "Yes", "Phase 2 (query plans)"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("Fallback Strategies for Standard Edition", level=2)
    add_styled_table(doc,
        ["Missing Feature", "Fallback Approach"],
        [
            ["QUERY_ATTRIBUTION_HISTORY", "Estimate per-query cost: (execution_time / 3600) x WH_credits_per_hour x credit_price"],
            ["ACCESS_HISTORY", "Use TABLE_STORAGE_METRICS + QUERY_HISTORY to approximate table usage patterns"],
            ["WAREHOUSE_EVENTS_HISTORY", "Monitor WAREHOUSE_METERING_HISTORY for usage pattern changes"],
            ["Budgets", "Implement custom budget tracking via dbt models + alerting"],
            ["ANOMALY_DETECTION ML", "Use custom z-score based detection (Section 7.1)"],
        ],
    )

    doc.add_page_break()

    # === SECTION 9: OPEN SOURCE ACCELERATORS ===
    doc.add_heading("9. Open Source Accelerators", level=1)
    doc.add_paragraph(
        "We leverage proven open-source projects to accelerate delivery and avoid "
        "reinventing common patterns."
    )

    doc.add_heading("9.1 dbt-snowflake-monitoring", level=2)
    doc.add_paragraph(
        "A dbt package by Select.dev that provides pre-built models for Snowflake cost "
        "and performance monitoring. We use it as a foundation and extend with custom models."
    )
    for item in [
        "Pre-built staging models for all ACCOUNT_USAGE views",
        "Warehouse cost attribution models",
        "Query performance analysis models",
        "Storage cost breakdown models",
        "Active community with regular updates",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()
    doc.add_heading("9.2 Sundeck OpsCenter", level=2)
    doc.add_paragraph(
        "An open-source Snowflake cost management tool that provides a Streamlit-based "
        "dashboard and pre-built cost analysis queries."
    )
    for item in [
        "Ready-to-deploy Streamlit dashboard",
        "Cost analysis and warehouse sizing recommendations",
        "Query profiling and anti-pattern detection",
        "Open-source (Apache 2.0 licence)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()
    doc.add_heading("9.3 snowflake_spend", level=2)
    doc.add_paragraph(
        "A lightweight dbt package focused specifically on Snowflake spend tracking "
        "and cost attribution."
    )

    doc.add_paragraph()
    doc.add_heading("9.4 Build vs Buy Analysis", level=2)
    add_styled_table(doc,
        ["Factor", "Build Custom", "Use Accelerators", "Our Approach (Hybrid)"],
        [
            ["Time to value", "6-8 weeks", "1-2 weeks", "3-4 weeks (accelerated with customisation)"],
            ["Customisation", "Full control", "Limited to package design", "Full — extend accelerator models"],
            ["Maintenance", "Team-owned", "Community-maintained", "Accelerators for base, custom for extensions"],
            ["Business context", "Fully integrated", "Generic", "Accelerators + custom team/product models"],
            ["Cost", "Engineering time only", "Free (open-source)", "Optimised engineering time"],
        ],
    )

    doc.add_page_break()

    # === SECTION 10: DELIVERY PLAN ===
    doc.add_heading("10. Delivery Plan", level=1)

    doc.add_heading("Phase 1: Identify & Solve (Weeks 1-4)", level=2)
    add_styled_table(doc,
        ["Week", "Activities", "Deliverables"],
        [
            ["Week 1", "Environment access, discovery, account profiling.\nGrant IMPORTED PRIVILEGES.\nProfile warehouses, databases, users, query volume.\nMap user groups across verticals.\nSet up dbt project, install accelerator packages.", "Discovery report: account profile, user landscape, initial findings, quick-win recommendations"],
            ["Week 2", "Build cost models + root cause analysis.\nCompute costs (WAREHOUSE_METERING + QUERY_ATTRIBUTION),\nstorage (TABLE_STORAGE_METRICS + COPY_HISTORY),\nserverless (METERING_HISTORY by service type).\nRun query pattern analysis:\nmine 90 days of QUERY_HISTORY, group by PARAMETERIZED_HASH,\nextract table/column/filter patterns, assign optimisation tiers.", "Working dbt models + query pattern analysis + root cause findings"],
            ["Week 3", "Build Streamlit dashboard: executive summary, warehouse deep-dive,\nBU attribution, storage explorer, trend analysis.\nGenerate tiered recommendation report\n(Tier 1-4 fixes ranked by ROI).\nApply quick-win config changes (auto-suspend, sizing).", "Interactive dashboard + tiered recommendations + quick wins applied"],
            ["Week 4", "Compile prioritised findings report with solutions.\nStakeholder walkthrough: present diagnosis + solution plan.\nDefine Phase 2 roadmap based on findings.\nKnowledge transfer session 1.", "Final findings report with solution recommendations, Phase 2 roadmap, dashboard"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("Phase 2: Optimise (Weeks 5-8)", level=2)
    add_styled_table(doc,
        ["Week", "Activities", "Deliverables"],
        [
            ["Week 5", "Implement warehouse right-sizing from Phase 1 findings.\nSandbox testing of resize recommendations.\nProfile WAREHOUSE_EVENTS_HISTORY for suspend/resume patterns.", "Warehouse resizing executed with before/after validation"],
            ["Week 6", "Implement query-level optimisations.\nFix top anti-patterns identified in Phase 1.\nImplement GET_QUERY_OPERATOR_STATS for deep analysis.\nEvaluate QAS candidates.", "Query optimisations applied with measured savings"],
            ["Week 7", "Execute storage cleanup from Phase 1 candidates.\nDrop unused tables, reduce TT, remove stale clones.\nRun SYSTEM$CLUSTERING_* analysis.\nDeploy Tier 1-3 fixes from Phase 1 recommendations.\nBuild metadata-aware contract generator.", "Storage savings realised + tiered fixes deployed + contract generator"],
            ["Week 8", "Deploy contract generator for new data products.\nEnhance cost dashboard with optimisation results.\nCompile verified savings report (before/after).\nStakeholder walkthrough.\nKnowledge transfer session 2.", "Verified savings report, contract generator deployed, Phase 3 readiness"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("Phase 3: Governance & FinOps (Weeks 9-12)", level=2)
    add_styled_table(doc,
        ["Week", "Activities", "Deliverables"],
        [
            ["Week 9", "Implement anomaly detection (ML-powered + custom z-score).\nConfigure alert thresholds.\nSet up drift detection monitors.", "Anomaly detection model, alerting framework"],
            ["Week 10", "Build chargeback model (3-tier allocation).\nIntegrate Snowflake Budgets.\nImplement tag-based cost allocation.", "Chargeback framework, budget configuration"],
            ["Week 11", "Establish FinOps review cadence.\nCreate review playbooks (weekly/monthly/quarterly).\nRun FinOps maturity assessment.", "Review cadence documents, maturity scorecard"],
            ["Week 12", "Final dashboard integration (governance + chargeback views).\nComprehensive testing.\nFinal stakeholder walkthrough.\nKnowledge transfer session 3 + handover.", "Complete governance framework, full handover package"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("Visual Timeline", level=2)
    add_phase_timeline(doc, [
        ("Phase 1: Identify & Solve", 4, DARK_BLUE),
        ("Phase 2: Optimise", 4, ACCENT_BLUE),
        ("Phase 3: Govern", 4, GREEN),
    ])

    doc.add_page_break()

    # === SECTION 11: PREREQUISITES & ASSUMPTIONS ===
    doc.add_heading("11. Prerequisites & Assumptions", level=1)

    doc.add_heading("Prerequisites", level=2)
    add_styled_table(doc,
        ["#", "Requirement", "Purpose", "Phase"],
        [
            ["1", "Grant IMPORTED PRIVILEGES on the SNOWFLAKE database", "Access to ACCOUNT_USAGE metadata views", "All"],
            ["2", "Provide a dedicated Snowflake warehouse (Small or Medium)", "Compute for dbt models and dashboard queries", "All"],
            ["3", "Provide a dedicated database/schema for framework objects", "Storage for cost models and dashboard", "All"],
            ["4", "Share Snowflake contract details (credit price, edition)", "Accurate dollar-cost conversion", "All"],
            ["5", "Provide mapping of warehouses/roles to BUs and MDS workspaces", "Cost attribution to specific business units", "Phase 1"],
            ["6", "Nominate technical point of contact (platform team)", "Collaboration during discovery and validation", "All"],
            ["7", "Confirm Snowflake edition (Standard/Enterprise/Business Critical)", "Feature availability scoping", "All"],
            ["8", "Confirm QUERY_ATTRIBUTION_HISTORY availability", "Per-query cost attribution approach", "Phase 1"],
            ["9", "Identify tag administration roles (for Phase 3)", "Tag-based chargeback implementation", "Phase 3"],
            ["10", "Provide list of known data products and owning BU teams", "Prioritise analysis on highest-cost products", "Phase 1"],
            ["11", "Confirm Immuta policy configuration access (read-only)", "Align cost governance with existing access policies", "Phase 3"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("Assumptions", level=2)
    for item in [
        "Customer is on Snowflake Enterprise Edition or higher (required for ACCESS_HISTORY, QUERY_ATTRIBUTION_HISTORY)",
        "ACCOUNT_USAGE views have data for at least 30 days (ideally 90+ days for trend analysis)",
        "The framework warehouse will not be used for other workloads (to avoid cost contamination)",
        "Query tags are either already in use or can be introduced for attribution",
        "Customer has dbt Cloud or is open to using it (alternatively, Snowflake Tasks can schedule models)",
        "Tag administration privileges can be granted for Phase 3 chargeback implementation",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # === SECTION 12: RISK & MITIGATION ===
    doc.add_heading("12. Risk & Mitigation", level=1)
    add_styled_table(doc,
        ["Risk", "Likelihood", "Impact", "Mitigation"],
        [
            ["IMPORTED PRIVILEGES not granted promptly", "Medium", "Blocks all work", "Raise as Day 1 action; provide exact GRANT statement"],
            ["Customer on Standard Edition", "Low", "No ACCESS_HISTORY, QUERY_ATTRIBUTION", "Fall back to estimated costs and TABLE_STORAGE_METRICS"],
            ["No query tags in use", "Medium", "Limits attribution granularity", "Attribute by warehouse/role; recommend tagging roadmap"],
            ["Very high query volume (>1M/day)", "Low", "Slow model builds", "Incremental materialisation with 7-day lookback"],
            ["Stakeholder availability", "Medium", "Delays sign-off", "Schedule walkthroughs at project start; async review"],
            ["QUERY_ATTRIBUTION_HISTORY not available", "Low", "No per-query costing", "Use estimated cost formula; flag for future enablement"],
            ["Governance adoption resistance", "Medium", "Phase 3 underutilised", "Start with lightweight weekly reviews; demonstrate value with anomaly catches"],
            ["Tag complexity overwhelms team", "Low", "Chargeback delayed", "Start with L1 (cost_centre) only; expand progressively"],
            ["Anomaly false positives", "Medium", "Alert fatigue", "Tune thresholds iteratively; start with high-confidence (z > 3.0) only"],
        ],
    )

    doc.add_page_break()

    # === SECTION 13: COMPETITIVE POSITIONING ===
    doc.add_heading("13. Competitive Positioning", level=1)

    doc.add_heading("13.1 vs Native Snowflake Cost Management", level=2)
    add_styled_table(doc,
        ["Capability", "Native Snowflake", "Our Framework", "Gap Filled"],
        [
            ["Cost dashboards", "Resource monitors (threshold alerts only)", "Interactive Streamlit with drill-down", "Analytical layer"],
            ["Attribution", "Warehouse-level only", "Query + team + product + tag-level", "Granular attribution"],
            ["Recommendations", "None", "Prioritised by estimated $ savings", "Actionable insights"],
            ["Anomaly detection", "Budget threshold alerts", "ML-powered + statistical detection", "Proactive monitoring"],
            ["Query analysis", "EXPLAIN (manual)", "Automated anti-pattern detection", "Scale and automation"],
            ["Chargeback", "Tags (manual setup)", "Automated 3-tier allocation", "End-to-end chargeback"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("13.2 vs SaaS Tools (Select.dev / Keebo / Sundeck)", level=2)
    add_styled_table(doc,
        ["Factor", "Our Framework", "SaaS Tools"],
        [
            ["Data residency", "Stays in customer's Snowflake account", "Data sent to third-party"],
            ["Customisation", "Fully tailored to customer's org structure", "Generic, one-size-fits-all"],
            ["Ongoing cost", "$0 (dbt + Streamlit are free/included)", "$500-$5,000+/month"],
            ["Business context", "Integrates team mappings, tags, products", "Limited to Snowflake metadata"],
            ["Extensibility", "Add models, metrics, dashboards freely", "Limited to vendor roadmap"],
            ["Transparency", "All logic visible in dbt SQL", "Black-box recommendations"],
            ["Implementation", "Requires engineering effort upfront", "Quick setup, limited customisation"],
            ["Support", "Ongoing relationship with delivery team", "Vendor support desk"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("13.3 vs DIY / Internal Build", level=2)
    add_styled_table(doc,
        ["Factor", "Our Framework", "DIY Build"],
        [
            ["Time to value", "4 weeks to first dashboard", "3-6 months typically"],
            ["Expertise required", "Snowflake cost domain knowledge included", "Team must learn ACCOUNT_USAGE intricacies"],
            ["Completeness", "8+ cost categories, 25+ views, 3 phases", "Usually covers compute only initially"],
            ["Best practices", "Built on open-source accelerators + proven patterns", "Discovered through trial and error"],
            ["Risk", "Proven approach, experienced team", "High risk of scope creep and gaps"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("13.4 TCO Comparison (12-Month)", level=2)
    add_styled_table(doc,
        ["Cost Component", "Our Framework", "SaaS Tool (Mid-tier)", "DIY Build"],
        [
            ["Implementation", "12 weeks of delivery", "1-2 weeks setup", "3-6 months engineering"],
            ["Ongoing licence", "$0", "$2,000-$5,000/month", "$0"],
            ["Maintenance", "Minimal (dbt models + dashboard)", "Vendor-managed", "Significant (team-owned)"],
            ["12-month licence total", "$0", "$24,000-$60,000", "$0"],
            ["Customisation", "Included in implementation", "Limited or extra cost", "Ongoing engineering cost"],
        ],
    )

    doc.add_page_break()

    # === SECTION 14: TEAM & EXPERTISE ===
    doc.add_heading("14. Team & Expertise", level=1)

    doc.add_heading("Team Composition", level=2)
    add_styled_table(doc,
        ["Role", "Responsibility", "Allocation", "Phases"],
        [
            ["Lead Engineer", "Architecture, dbt model development, Snowflake expertise, query analysis", "Full-time", "All phases (12 weeks)"],
            ["Dashboard Developer", "Streamlit dashboard development, UX design, wireframe implementation", "Part-time", "Weeks 3-4, 7-8, 11-12"],
            ["Project Lead", "Stakeholder management, delivery oversight, customer communication", "Part-time", "Throughout"],
            ["FinOps Consultant", "Governance framework, chargeback design, maturity assessment, review cadence", "Part-time", "Phase 3 (Weeks 9-12)"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("Key Competencies", level=2)
    for item in [
        "Snowflake architecture, ACCOUNT_USAGE schema, and performance tuning",
        "dbt (data build tool) — modelling, testing, documentation, incremental patterns",
        "Streamlit dashboard development (native in Snowflake)",
        "Terraform infrastructure-as-code for Snowflake",
        "Data observability and monitoring frameworks",
        "FinOps practices and cost governance",
        "Open-source Snowflake monitoring ecosystem (dbt-snowflake-monitoring, OpsCenter)",
        "Data product architecture — three-layer design (staging/intermediate/publication) with versioned contracts",
        "Automated alert frameworks — episode tracking, multi-channel routing (Datadog, Incident.io, Jira)",
        "Data quality governance — 7-dimension framework with RAG scoring, SLOs, and automated testing",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()
    doc.add_heading("Production Track Record", level=2)
    doc.add_paragraph(
        "Our team operates a production-grade data product platform on Snowflake that "
        "demonstrates the exact capabilities proposed in this specification:"
    )
    add_styled_table(doc,
        ["Capability", "Production Evidence"],
        [
            ["Data Products", "10+ products in production with three-layer architecture, versioned contracts, and automated testing"],
            ["Streamlit Dashboards", "13 dashboards across executive, operational, and alert tiers — same approach for cost governance"],
            ["Alert Framework", "30+ alert types with stateful episode tracking, suppression rules, and <2-hour development time per alert"],
            ["Data Quality", "7-dimension quality framework (Accuracy, Completeness, Consistency, Timeliness, Validity, Uniqueness, Integrity) with RAG scoring"],
            ["External Integration", "Pull-based architecture: Datadog, Incident.io, and Jira query Snowflake directly — zero API maintenance"],
            ["Infrastructure as Code", "Terraform-managed Snowflake (4 environments, per-product consumer roles, principle of least privilege)"],
            ["Refresh Cadence", "15-minute cycle with all models completing in <10 minutes — near-real-time monitoring"],
            ["Data Contracts", "Versioned public models with 30-day change notice and 90-day deprecation periods"],
        ],
    )

    # === SECTION 15: NEXT STEPS ===
    doc.add_paragraph()
    doc.add_heading("15. Next Steps", level=1)
    doc.add_paragraph(
        "Based on the introductory call of 9 April 2026:"
    )
    add_styled_table(doc,
        ["#", "Action", "Owner", "Timeline"],
        [
            ["1", "Share observability demo / dashboard screenshots", "Bilvantis", "This week"],
            ["2", "Review and approve this specification", "Thomson Reuters", "By 18 April 2026"],
            ["3", "Arrange Snowflake instance access for POC/pilot\n(or share non-sensitive metadata export)", "Thomson Reuters", "Week of 21 April"],
            ["4", "Begin Phase 1: deploy cost framework on TR's Snowflake", "Bilvantis delivery team", "Upon access"],
        ],
    )

    doc.add_page_break()


def _build_tech_spec_appendices(doc):
    """Build appendices A-D of the technical specification."""

    # === APPENDIX A: ACCOUNT_USAGE VIEWS ===
    doc.add_heading("Appendix A: ACCOUNT_USAGE Views Reference", level=1)
    doc.add_paragraph(
        "Complete reference of SNOWFLAKE.ACCOUNT_USAGE views used by the framework."
    )

    doc.add_heading("Compute Views", level=2)
    add_styled_table(doc,
        ["View", "Description", "Retention", "Key Columns"],
        [
            ["WAREHOUSE_METERING_HISTORY", "Credit consumption per warehouse per hour", "365 days", "warehouse_name, credits_used, start_time"],
            ["QUERY_HISTORY", "Full detail of every query executed", "365 days", "query_id, warehouse_name, user_name, execution_time, bytes_scanned, query_tag"],
            ["WAREHOUSE_LOAD_HISTORY", "Warehouse utilisation: running, queued, blocked", "365 days", "warehouse_name, avg_running, avg_queued_load"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("Storage Views", level=2)
    add_styled_table(doc,
        ["View", "Description", "Retention", "Key Columns"],
        [
            ["TABLE_STORAGE_METRICS", "Active, Time Travel, Fail-safe bytes per table", "Current snapshot", "table_name, active_bytes, time_travel_bytes, failsafe_bytes"],
            ["STORAGE_USAGE", "Total account-level storage over time", "365 days", "usage_date, storage_bytes, stage_bytes, failsafe_bytes"],
            ["DATABASE_STORAGE_USAGE_HISTORY", "Storage per database over time", "365 days", "database_name, average_database_bytes"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("Access and Lineage Views", level=2)
    add_styled_table(doc,
        ["View", "Description", "Retention", "Key Columns"],
        [
            ["ACCESS_HISTORY", "Tables/columns read or written by each query", "365 days", "query_id, direct_objects_accessed, base_objects_accessed"],
            ["LOGIN_HISTORY", "User login events", "365 days", "user_name, client_ip, reported_client_type"],
            ["SESSIONS", "Session details", "365 days", "session_id, user_name, client_application_id"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("Serverless Views", level=2)
    add_styled_table(doc,
        ["View", "Description", "Retention", "Key Columns"],
        [
            ["AUTOMATIC_CLUSTERING_HISTORY", "Auto-clustering credits", "365 days", "table_name, credits_used, num_bytes_reclustered"],
            ["MATERIALIZED_VIEW_REFRESH_HISTORY", "MV refresh credits", "365 days", "table_name, credits_used"],
            ["PIPE_USAGE_HISTORY", "Snowpipe credits", "365 days", "pipe_name, credits_used, bytes_inserted"],
            ["SERVERLESS_TASK_HISTORY", "Serverless task credits", "365 days", "task_name, credits_used"],
            ["SEARCH_OPTIMIZATION_HISTORY", "Search optimisation credits", "365 days", "table_name, credits_used"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("Cost and Billing Views", level=2)
    add_styled_table(doc,
        ["View", "Description", "Retention", "Key Columns"],
        [
            ["METERING_HISTORY", "Unified credit consumption (all service types)", "365 days", "service_type, credits_used, usage_date"],
            ["METERING_DAILY_HISTORY", "Daily billed credits", "365 days", "usage_date, credits_billed"],
            ["QUERY_ATTRIBUTION_HISTORY", "Per-query cost attribution", "365 days", "query_id, credits_attributed_compute"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("Governance View", level=2)
    add_styled_table(doc,
        ["View", "Description", "Retention", "Key Columns"],
        [
            ["WAREHOUSE_EVENTS_HISTORY", "Warehouse lifecycle events", "365 days", "warehouse_name, event_name, timestamp"],
        ],
    )

    doc.add_page_break()

    # === APPENDIX B: OPTIMISATION FUNCTIONS ===
    doc.add_heading("Appendix B: Optimisation Functions Reference", level=1)

    add_styled_table(doc,
        ["Function", "Signature", "Returns", "Use Case"],
        [
            ["GET_QUERY_OPERATOR_STATS", "GET_QUERY_OPERATOR_STATS(query_id)", "Table: operator_id, operator_type, operator_statistics, execution_time_breakdown", "Identify expensive operators within a query's execution plan"],
            ["SYSTEM$CLUSTERING_INFORMATION", "SYSTEM$CLUSTERING_INFORMATION(table, columns)", "JSON: cluster_by_keys, total_partition_count, average_overlaps, average_depth", "Evaluate clustering effectiveness for a table"],
            ["SYSTEM$CLUSTERING_DEPTH", "SYSTEM$CLUSTERING_DEPTH(table, columns)", "Number: clustering depth (1.0 = perfect)", "Quick clustering quality check"],
            ["SYSTEM$ESTIMATE_QUERY_ACCELERATION", "SYSTEM$ESTIMATE_QUERY_ACCELERATION(query_id)", "JSON: originalQueryTime, acceleratedQueryTimes for scale factors 1x-24x", "Evaluate QAS benefit for a specific query"],
        ],
    )

    doc.add_page_break()

    # === APPENDIX C: DASHBOARD WIREFRAMES ===
    doc.add_heading("Appendix C: Dashboard Wireframes", level=1)

    doc.add_heading("1. Executive Summary Dashboard", level=2)
    add_ascii_diagram(doc, [
        "+-------------------------------+-------------------------------+",
        "| TOTAL MONTHLY SPEND           | COST TREND (90 DAYS)          |",
        "|  $124,500 (+8% MoM)           |  ___                          |",
        "|                               | /   \\___    ___/              |",
        "|  Compute:  $89,200 (72%)      |/        \\__/                  |",
        "|  Storage:  $22,100 (18%)      |                               |",
        "|  Svrless:  $13,200 (10%)      | Jan  Feb  Mar  Apr            |",
        "+-------------------------------+-------------------------------+",
        "| TOP 5 WAREHOUSES              | COST BY TEAM                  |",
        "|  1. ANALYTICS_WH    $31,200   |  Engineering    45%  ====     |",
        "|  2. ETL_WH          $24,800   |  Analytics      28%  ===      |",
        "|  3. REPORTING_WH    $15,600   |  Data Science   18%  ==       |",
        "|  4. DATA_SCIENCE_WH $11,400   |  Ad Hoc          9%  =        |",
        "|  5. AD_HOC_WH        $6,200   |                               |",
        "+-------------------------------+-------------------------------+",
    ])

    doc.add_paragraph()
    doc.add_heading("2. Warehouse Deep-Dive Dashboard", level=2)
    add_ascii_diagram(doc, [
        "+-------------------------------------------------------------------+",
        "| WAREHOUSE: [Dropdown]                          Cost: $31,200/mo   |",
        "+-------------------------------------------------------------------+",
        "| Utilisation     | Idle: 34%  Active: 58%  Queued: 8%              |",
        "| Size            | LARGE (8 credits/hr)                            |",
        "| Auto-suspend    | 300s (Rec: reduce to 60s)                       |",
        "+-------------------------------------------------------------------+",
        "| HOURLY USAGE                   | TOP USERS                        |",
        "|   __                           |  svc_etl        42% ====         |",
        "|  /  \\     __                   |  analyst_1      23% ==           |",
        "| /    \\___/  \\                  |  analyst_2      15% ==           |",
        "|/             \\___              |  data_eng_1     11% =            |",
        "| 6am   12pm  6pm  12am          |  other           9% =            |",
        "+--------------------------------+----------------------------------+",
    ])

    doc.add_paragraph()
    doc.add_heading("3. Team Attribution Dashboard", level=2)
    add_ascii_diagram(doc, [
        "+-------------------------------------------------------------------+",
        "| TEAM: [Dropdown / All]                   Period: [Last 30 Days]   |",
        "+-------------------------------------------------------------------+",
        "| TEAM COST BREAKDOWN             | TEAM TREND                      |",
        "|  Engineering   $56,200  ======   |    ___                          |",
        "|  Analytics     $34,900  ====     |   /   \\                         |",
        "|  Data Science  $22,400  ===      |  /     \\___                     |",
        "|  Ad Hoc        $11,000  =        | /          \\___                 |",
        "+-------------------------------------------------------------------+",
        "| TOP QUERIES BY COST (THIS TEAM)                                   |",
        "|  Query ID    | User       | Cost    | Duration | Recommendation   |",
        "|  abc123...   | analyst_1  | $142    | 45 min   | Add filter       |",
        "|  def456...   | svc_etl    | $98     | 32 min   | Reduce WH size   |",
        "+-------------------------------------------------------------------+",
    ])

    doc.add_paragraph()
    doc.add_heading("4. Storage Explorer Dashboard", level=2)
    add_ascii_diagram(doc, [
        "+-------------------------------------------------------------------+",
        "| TOTAL STORAGE: 45.2 TB          Monthly Cost: $22,100             |",
        "+-------------------------------------------------------------------+",
        "| BY TYPE                         | BY DATABASE                     |",
        "|  Active:      32.1 TB  (71%)    |  PROD_DW     18.4 TB  ======   |",
        "|  Time Travel:  9.8 TB  (22%)    |  RAW_VAULT   12.2 TB  ====     |",
        "|  Fail-safe:    3.3 TB  ( 7%)    |  STAGING      8.1 TB  ===      |",
        "+-------------------------------------------------------------------+",
        "| UNUSED TABLES (No reads 90+ days)          Total: 4.2 TB         |",
        "|  Table                  | Size   | Last Read  | Action            |",
        "|  raw.legacy_orders      | 1.8 TB | 180d ago   | Archive           |",
        "|  staging.tmp_migration  | 1.2 TB | 120d ago   | Drop              |",
        "+-------------------------------------------------------------------+",
    ])

    doc.add_paragraph()
    doc.add_heading("5. Trend Analysis Dashboard", level=2)
    add_ascii_diagram(doc, [
        "+-------------------------------------------------------------------+",
        "| COST TREND: Last 90 Days                    Granularity: [Daily]  |",
        "+-------------------------------------------------------------------+",
        "|  $6K |         *                                                  |",
        "|  $5K |    *   * *        *                   *                    |",
        "|  $4K |  ** * *   * *  * * * *  * *  *  **  * * *                  |",
        "|  $3K | *         *  **       **   **  *  **     *                 |",
        "|  $2K |                                                            |",
        "|      +---+---+---+---+---+---+---+---+---+---+---+---+           |",
        "|       Jan         Feb         Mar         Apr                     |",
        "+-------------------------------------------------------------------+",
        "| ANOMALIES DETECTED: 3                                             |",
        "|  Date       | Cost    | Expected | Z-Score | Cause               |",
        "|  Mar 15     | $5,800  | $4,100   | 2.8     | Ad-hoc WH spike     |",
        "|  Feb 28     | $5,200  | $4,000   | 2.5     | Month-end reporting  |",
        "+-------------------------------------------------------------------+",
    ])

    doc.add_paragraph()
    doc.add_heading("6. Recommendations Dashboard", level=2)
    add_ascii_diagram(doc, [
        "+-------------------------------------------------------------------+",
        "| TOTAL SAVINGS IDENTIFIED: $40,600/month                           |",
        "+-------------------------------------------------------------------+",
        "| BY CATEGORY               | BY EFFORT                            |",
        "|  WH Sizing   $18,500 45%  |  Low (config)     $22,300  55%       |",
        "|  Query Opt   $12,200 30%  |  Medium (SQL)     $14,100  35%       |",
        "|  Storage      $6,800 17%  |  High (arch)       $4,200  10%       |",
        "|  Serverless   $3,100  8%  |                                      |",
        "+-------------------------------------------------------------------+",
        "| TOP RECOMMENDATIONS                                               |",
        "|  # | Category  | Object        | Saving  | Effort | Action       |",
        "|  1 | WH Size   | ANALYTICS_WH  | $8,200  | Low    | L -> M       |",
        "|  2 | WH Size   | ETL_WH        | $5,100  | Low    | Suspend 60s  |",
        "|  3 | Query     | Top 10 queries| $4,800  | Med    | Add filters  |",
        "+-------------------------------------------------------------------+",
    ])

    doc.add_paragraph()
    doc.add_heading("7. Governance Dashboard", level=2)
    add_ascii_diagram(doc, [
        "+-------------------------------------------------------------------+",
        "| GOVERNANCE STATUS                          Period: Current Month  |",
        "+-------------------------------------------------------------------+",
        "| ANOMALIES              | DRIFT ALERTS         | BUDGET STATUS    |",
        "|  Open:    2            |  Open:    1           |  On track:  3    |",
        "|  Resolved: 5          |  Resolved: 4          |  Warning:   1    |",
        "|  This week: 1         |  This week: 0         |  Over:      0    |",
        "+-------------------------------------------------------------------+",
        "| FINOPS MATURITY                                                    |",
        "|  Visibility:    [=====>     ] Walk (3/5)                           |",
        "|  Optimisation:  [===>       ] Crawl (2/5)                          |",
        "|  Governance:    [==>        ] Crawl (1/5)                          |",
        "|  Accountability:[===>       ] Crawl (2/5)                          |",
        "+-------------------------------------------------------------------+",
    ])

    doc.add_paragraph()
    doc.add_heading("8. Chargeback Dashboard", level=2)
    add_ascii_diagram(doc, [
        "+-------------------------------------------------------------------+",
        "| CHARGEBACK REPORT                          Month: April 2026     |",
        "+-------------------------------------------------------------------+",
        "| COST BY COST CENTRE          | BUDGET vs ACTUAL                  |",
        "|  Engineering   $56,200 (45%) |  Eng:    $56K / $60K  [======> ]  |",
        "|  Analytics     $34,900 (28%) |  Ana:    $35K / $35K  [========]  |",
        "|  Data Science  $22,400 (18%) |  DS:     $22K / $25K  [======> ]  |",
        "|  Ad Hoc        $11,000 ( 9%) |  AdHoc:  $11K / $10K  [========!] |",
        "+-------------------------------------------------------------------+",
        "| ALLOCATION METHOD BREAKDOWN                                       |",
        "|  Tier 1 (Direct WH):     62%  High confidence                    |",
        "|  Tier 2 (Tag-based):     28%  High confidence                    |",
        "|  Tier 3 (Proportional):  10%  Medium confidence                  |",
        "+-------------------------------------------------------------------+",
    ])

    doc.add_paragraph()
    doc.add_heading("9. Tiered Optimisation Recommendations", level=2)
    add_ascii_diagram(doc, [
        "+-------------------------------------------------------------------+",
        "| OPTIMISATION RECOMMENDATIONS              [Filter by BU v]       |",
        "+-------------------------------------------------------------------+",
        "| TIER | PATTERN               | FIX                | SAVING      |",
        "+-------------------------------------------------------------------+",
        "| T1   | Finance: full-table   | Add partition      | $775/mo     |",
        "|      | scan on transactions  | filter (date>=90d) | Fix: 5 min  |",
        "|  ----------                  | ------             | ------      |",
        "| T1   | Ops: missing cluster  | ALTER TABLE ADD    | $420/mo     |",
        "|      | key on shipments      | CLUSTER BY (date)  | Fix: 5 min  |",
        "|  ----------                  | ------             | ------      |",
        "| T2   | 200 users: same       | CREATE MV          | $1,470/mo   |",
        "|      | daily revenue agg     | (1 refresh/day)    | Fix: 10 min |",
        "|  ----------                  | ------             | ------      |",
        "| T3   | 8 users x 2 BUs:     | CREATE VIEW with   | $1,170/mo   |",
        "|      | order+customer join   | optimised joins    | Fix: 1-2 hr |",
        "|  ----------                  | ------             | ------      |",
        "| T4   | 15 users x 4 BUs:    | Full dbt model +   | $2,800/mo   |",
        "|      | customer_360 pattern  | Alation + Immuta   | Fix: 3-5 day|",
        "+-------------------------------------------------------------------+",
        "| TOTAL IDENTIFIED: 47 patterns | Est. saving: $18,400/mo         |",
        "| Tier 1: 28 (60%) | Tier 2: 11 (23%) | Tier 3: 6 (13%) | T4: 2 |",
        "+-------------------------------------------------------------------+",
        "| SELECTED: T1 - Finance full-table scan                           |",
        "| Owner: @jane.smith (Finance BU) | Tables: raw.transactions      |",
        "| [Copy Fix SQL] [Send to Owner] [Mark Resolved]                   |",
        "+-------------------------------------------------------------------+",
    ])
    doc.add_paragraph(
        "The recommendation dashboard shows all identified patterns ranked by saving, "
        "grouped by tier. Platform team can send Tier 1 fixes directly to BU owners, "
        "deploy Tier 2 config changes themselves, and plan Tier 3-4 builds based on ROI. "
        "Progress tracking shows how many patterns are resolved and cumulative savings."
    )

    doc.add_page_break()

    # === APPENDIX D: SAMPLE SQL QUERIES ===
    doc.add_heading("Appendix D: Sample SQL Queries", level=1)
    doc.add_paragraph(
        "Ready-to-run SQL queries covering all analysis areas. Replace :credit_price "
        "with the customer's contract rate."
    )

    queries = [
        ("1. Monthly Cost Summary by Category", """\
SELECT
    DATE_TRUNC('month', usage_date) AS usage_month,
    service_type,
    SUM(credits_used) AS total_credits,
    SUM(credits_used) * :credit_price AS total_cost_usd
FROM snowflake.organization_usage.metering_history
WHERE usage_date >= DATEADD('month', -6, CURRENT_DATE())
GROUP BY 1, 2
ORDER BY usage_month DESC, total_credits DESC;"""),

        ("2. Top 20 Most Expensive Queries (Last 7 Days)", """\
SELECT
    qah.query_id,
    qh.user_name,
    qh.warehouse_name,
    qh.query_tag,
    qh.total_elapsed_time / 1000 AS elapsed_seconds,
    (qah.credits_attributed_compute
     + qah.credits_attributed_cloud_services
    ) * :credit_price AS query_cost_usd,
    LEFT(qh.query_text, 200) AS query_preview
FROM snowflake.account_usage.query_attribution_history qah
JOIN snowflake.account_usage.query_history qh
  ON qah.query_id = qh.query_id
WHERE qh.start_time >= DATEADD('day', -7, CURRENT_DATE())
ORDER BY query_cost_usd DESC
LIMIT 20;"""),

        ("3. Warehouse Utilisation and Idle Analysis", """\
WITH warehouse_activity AS (
    SELECT
        warehouse_name,
        DATE_TRUNC('day', start_time) AS usage_date,
        SUM(credits_used) AS daily_credits,
        COUNT(DISTINCT qh.query_id) AS daily_queries
    FROM snowflake.account_usage.warehouse_metering_history wmh
    LEFT JOIN snowflake.account_usage.query_history qh
      ON wmh.warehouse_name = qh.warehouse_name
      AND DATE_TRUNC('hour', qh.start_time) = wmh.start_time
    WHERE wmh.start_time >= DATEADD('day', -30, CURRENT_DATE())
    GROUP BY 1, 2
)
SELECT
    warehouse_name,
    AVG(daily_credits) AS avg_daily_credits,
    AVG(daily_queries) AS avg_daily_queries,
    AVG(daily_credits) / NULLIF(AVG(daily_queries), 0) AS credits_per_query
FROM warehouse_activity
GROUP BY warehouse_name
ORDER BY avg_daily_credits DESC;"""),

        ("4. Unused Tables (No Reads in 90+ Days)", """\
SELECT
    tsm.table_catalog AS database_name,
    tsm.table_schema AS schema_name,
    tsm.table_name,
    tsm.active_bytes / POWER(1024, 3) AS active_gb,
    tsm.time_travel_bytes / POWER(1024, 3) AS tt_gb,
    MAX(ah.query_start_time) AS last_read_time,
    DATEDIFF('day', MAX(ah.query_start_time), CURRENT_DATE()) AS days_since_read
FROM snowflake.account_usage.table_storage_metrics tsm
LEFT JOIN snowflake.account_usage.access_history ah
  ON ah.base_objects_accessed LIKE '%' || tsm.table_name || '%'
WHERE tsm.active_bytes > 0
GROUP BY 1, 2, 3, 4, 5
HAVING days_since_read > 90 OR last_read_time IS NULL
ORDER BY active_gb DESC;"""),

        ("5. Query Anti-Pattern: Full Table Scans", """\
SELECT
    query_id,
    user_name,
    warehouse_name,
    total_elapsed_time / 1000 AS elapsed_seconds,
    partitions_scanned,
    partitions_total,
    ROUND(partitions_scanned / NULLIF(partitions_total, 0), 2) AS scan_ratio,
    bytes_scanned / POWER(1024, 3) AS gb_scanned,
    LEFT(query_text, 200) AS query_preview
FROM snowflake.account_usage.query_history
WHERE start_time >= DATEADD('day', -7, CURRENT_DATE())
  AND partitions_total > 100
  AND partitions_scanned / NULLIF(partitions_total, 0) > 0.8
  AND query_type = 'SELECT'
ORDER BY gb_scanned DESC
LIMIT 20;"""),

        ("6. Serverless Cost Breakdown", """\
SELECT
    service_type,
    DATE_TRUNC('month', usage_date) AS usage_month,
    SUM(credits_used) AS total_credits,
    SUM(credits_used) * :credit_price AS cost_usd
FROM snowflake.organization_usage.metering_history
WHERE service_type NOT IN ('WAREHOUSE_METERING', 'CLOUD_SERVICES')
  AND usage_date >= DATEADD('month', -3, CURRENT_DATE())
GROUP BY 1, 2
ORDER BY usage_month DESC, total_credits DESC;"""),

        ("7. Storage Waste: Excessive Time Travel", """\
SELECT
    table_catalog AS database_name,
    table_schema AS schema_name,
    table_name,
    active_bytes / POWER(1024, 3) AS active_gb,
    time_travel_bytes / POWER(1024, 3) AS tt_gb,
    failsafe_bytes / POWER(1024, 3) AS fs_gb,
    ROUND(time_travel_bytes / NULLIF(active_bytes, 0), 2) AS tt_to_active_ratio
FROM snowflake.account_usage.table_storage_metrics
WHERE time_travel_bytes > 0
  AND time_travel_bytes > active_bytes  -- TT > active = likely excessive
ORDER BY tt_gb DESC
LIMIT 20;"""),

        ("8. Repeated Query Patterns (Materialisation Candidates)", """\
SELECT
    query_parameterized_hash,
    COUNT(*) AS execution_count,
    SUM(total_elapsed_time) / 1000 AS total_seconds,
    AVG(total_elapsed_time) / 1000 AS avg_seconds,
    SUM(bytes_scanned) / POWER(1024, 3) AS total_gb_scanned,
    ANY_VALUE(warehouse_name) AS sample_warehouse,
    ANY_VALUE(LEFT(query_text, 200)) AS sample_query
FROM snowflake.account_usage.query_history
WHERE start_time >= DATEADD('day', -7, CURRENT_DATE())
  AND query_type = 'SELECT'
GROUP BY query_parameterized_hash
HAVING execution_count > 50
ORDER BY total_seconds DESC
LIMIT 20;"""),

        ("9. Cost Anomaly Detection (Z-Score)", """\
WITH daily_costs AS (
    SELECT
        usage_date AS cost_date,
        SUM(credits_used) * :credit_price AS daily_cost
    FROM snowflake.organization_usage.metering_history
    WHERE usage_date >= DATEADD('day', -90, CURRENT_DATE())
    GROUP BY usage_date
),
stats AS (
    SELECT *,
        AVG(daily_cost) OVER (ORDER BY cost_date ROWS 30 PRECEDING) AS rolling_avg,
        STDDEV(daily_cost) OVER (ORDER BY cost_date ROWS 30 PRECEDING) AS rolling_std
    FROM daily_costs
)
SELECT
    cost_date,
    daily_cost,
    rolling_avg,
    ROUND((daily_cost - rolling_avg) / NULLIF(rolling_std, 0), 2) AS z_score,
    CASE
        WHEN z_score > 3.0 THEN 'CRITICAL'
        WHEN z_score > 2.5 THEN 'HIGH'
        WHEN z_score > 2.0 THEN 'MEDIUM'
        ELSE 'NORMAL'
    END AS anomaly_level
FROM stats
WHERE z_score > 2.0
ORDER BY cost_date DESC;"""),

        ("10. Chargeback Report by Tag", """\
WITH tagged_costs AS (
    SELECT
        tv.tag_value AS cost_centre,
        qah.credits_attributed_compute * :credit_price AS query_cost
    FROM snowflake.account_usage.query_attribution_history qah
    JOIN snowflake.account_usage.query_history qh
      ON qah.query_id = qh.query_id
    JOIN snowflake.account_usage.tag_references tr
      ON tr.object_name = qh.warehouse_name
      AND tr.tag_name = 'COST_CENTRE'
    JOIN snowflake.account_usage.tags tv
      ON tr.tag_id = tv.tag_id
    WHERE qh.start_time >= DATEADD('month', -1, CURRENT_DATE())
)
SELECT
    cost_centre,
    COUNT(*) AS query_count,
    SUM(query_cost) AS total_cost,
    AVG(query_cost) AS avg_query_cost
FROM tagged_costs
GROUP BY cost_centre
ORDER BY total_cost DESC;"""),
    ]

    for title, sql in queries:
        doc.add_heading(title, level=2)
        add_sql_block(doc, sql)
        doc.add_paragraph()


def create_technical_specification():
    """Create Document 2: Full Technical Specification (~35-40 pages)."""
    doc = Document()
    setup_doc_styles(doc)

    # === TITLE PAGE ===
    create_title_page(doc, "Snowflake Cost Optimisation\nFramework", "Full Technical Specification\n\nPrepared for Thomson Reuters", version="2.0")

    # === TABLE OF CONTENTS ===
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Problem Statement & Business Impact",
        "3. Solution Architecture",
        "4. Cost Analysis Framework",
        "    4.1 METERING_HISTORY (Unified Cost)",
        "    4.2 METERING_DAILY_HISTORY (Billed Credits)",
        "    4.3 QUERY_ATTRIBUTION_HISTORY (Per-Query $)",
        "    4.4 Complete Cost Category Breakdown",
        "5. Phase 1: Identify & Solve",
        "    5.1 Compute Analysis",
        "    5.2 Storage Analysis",
        "    5.3 Serverless Analysis",
        "    5.4 Tag-Based Attribution",
        "    5.5 Dashboard Wireframes",
        "6. Phase 2: Query Optimisation",
        "    6.1 Query Plan Analysis",
        "    6.2 Clustering Analysis",
        "    6.3 Query Acceleration Service",
        "    6.4 Anti-Pattern Detection",
        "    6.5 Warehouse Right-Sizing",
        "    6.6 Storage Optimisation",
        "    6.7 Query Pattern Analysis & Tiered Recommendation Engine",
        "    6.8 Governed Data Product Lifecycle (Contract Generation)",
        "7. Phase 3: Governance & FinOps",
        "    7.1 Anomaly Detection",
        "    7.2 Chargeback Model",
        "    7.3 FinOps Review Cadence",
        "    7.4 Drift Detection",
        "    7.5 FinOps Maturity Assessment",
        "    7.6 Query Governance by User Type",
        "8. Edition Considerations",
        "9. Open Source Accelerators",
        "10. Delivery Plan",
        "11. Prerequisites & Assumptions",
        "12. Risk & Mitigation",
        "13. Competitive Positioning",
        "14. Team & Expertise",
        "15. Next Steps",
        "Appendix A: ACCOUNT_USAGE Views Reference",
        "Appendix B: Optimisation Functions Reference",
        "Appendix C: Dashboard Wireframes",
        "Appendix D: Sample SQL Queries",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        if item.startswith("    "):
            p.paragraph_format.left_indent = Cm(1)

    doc.add_page_break()

    # =======================================================================
    # SECTION 1: EXECUTIVE SUMMARY
    # =======================================================================
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "This document presents the full technical specification for a Snowflake Cost "
        "Optimisation Framework tailored to Thomson Reuters' data platform. The framework "
        "addresses the three priorities identified in the introductory call of 9 April 2026: "
        "(1) identify which data products and dbt contracts are consuming the most warehouse "
        "resources, (2) optimise existing code with before/after evidence, and (3) establish "
        "governance so new contracts follow TR guidelines from day one."
    )
    doc.add_paragraph(
        "The framework analyses 8+ cost categories across TR's four data layers (Raw/Role, "
        "TRIM, MDS, Broad/Data Products) using Snowflake's METERING_HISTORY for unified cost "
        "tracking, QUERY_ATTRIBUTION_HISTORY for per-query dollar attribution, and 25+ "
        "ACCOUNT_USAGE views — covering all 350+ source systems and all MDS workspaces."
    )
    doc.add_paragraph("The three phases deliver progressive value:")
    for item in [
        "Phase 1 (Weeks 1-4): Identify & Solve — diagnose cost drivers, map user patterns, deliver findings report with solution recommendations",
        "Phase 2 (Weeks 5-8): Query Optimisation — query plan analysis, warehouse right-sizing, anti-pattern detection, prioritised savings",
        "Phase 3 (Weeks 9-12): Governance & FinOps — anomaly detection, chargeback, review cadence, maturity assessment",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_callout_box(doc, "Zero additional licence cost. Everything runs inside the customer's Snowflake account using dbt (open-source), Streamlit (included with Snowflake), and native metadata views.")

    doc.add_page_break()

    # =======================================================================
    # SECTION 2: PROBLEM STATEMENT & BUSINESS IMPACT
    # =======================================================================
    doc.add_heading("2. Problem Statement & Business Impact", level=1)

    doc.add_heading("2.1 Thomson Reuters — Current Challenges", level=2)
    doc.add_paragraph(
        "Based on the introductory call of 9 April 2026, the following challenges were "
        "identified across Thomson Reuters' Snowflake platform:"
    )
    add_styled_table(doc,
        ["Challenge", "TR-Specific Context", "Business Impact"],
        [
            ["No pinpointed cost attribution", "Platform team knows costs are high but cannot tell each BU exactly which dbt jobs, contracts, or stored procedures drive spend", "Cannot inform BU owners with actionable evidence to drive code optimisation"],
            ["Unoptimised data products & contracts", "BU teams in MDS workspaces write contracts that work correctly but don't leverage partition columns, filter strategies, or efficient join patterns", "'$100 query that an expert writes for $10' — one team proved savings but approach doesn't scale manually to 350+ sources"],
            ["Oversized warehouses", "Warehouses configured for peak loads; platform team manually adjusts suspend/resume times but cannot reach all workloads", "Paying for idle compute; auto-suspend tuning is ad-hoc, not systematic"],
            ["Storage across 4 layers", "Data retained in Raw/Role, TRIM, MDS, and Broad layers with potential for unused tables, excessive TT, and stale clones", "Storage costs compound across layers without visibility into what's actually used"],
            ["Growing serverless costs", "350+ sources ingested via Snowpipe, Glue, Fivetran — each new source adds serverless compute", "Serverless costs grow silently with each new source system onboarded"],
            ["Reactive cost management", "Optimisation 'takes a back step when other commitments are priority'; year-on-year vendor costs increase", "No proactive monitoring — cost spikes discovered from invoices, not alerts"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("2.2 Cost-per-Query-Day (QPD) Metric", level=2)
    doc.add_paragraph(
        "We introduce the Query-Per-Dollar (QPD) metric as a normalised efficiency indicator. "
        "QPD measures the number of queries executed per dollar spent, enabling cross-warehouse "
        "and cross-team efficiency comparisons independent of absolute spend."
    )
    add_sql_block(doc, "QPD = total_queries_executed / total_cost_dollars\n\n-- Example: Warehouse A runs 50,000 queries at $10,000 = 5.0 QPD\n-- Example: Warehouse B runs 20,000 queries at $10,000 = 2.0 QPD\n-- Warehouse A is 2.5x more cost-efficient")

    doc.add_paragraph()
    doc.add_heading("2.3 Cost-of-Inaction Formula", level=2)
    doc.add_paragraph(
        "Without intervention, Snowflake costs grow predictably. The cost-of-inaction over N quarters:"
    )
    add_sql_block(doc, "Cost_of_Inaction = Current_Quarterly_Spend x ((1 + growth_rate)^N - 1)\n\n-- At 20% QoQ growth, $100K/quarter becomes:\n-- After 4 quarters: $186K/quarter (+86% cumulative)\n-- After 8 quarters: $430K/quarter (+330% cumulative)\n-- Total excess spend over 2 years: ~$740K")

    doc.add_paragraph()
    doc.add_heading("2.4 Savings Benchmarks by Category", level=2)
    add_styled_table(doc,
        ["Cost Category", "% of Typical Bill", "Waste Range", "Achievable Savings", "Primary Technique"],
        [
            ["Warehouse Compute", "60-70%", "30-50%", "20-35%", "Right-sizing, auto-suspend, scheduling"],
            ["Query Efficiency", "(within compute)", "15-40%", "10-25%", "Anti-pattern fix, caching, QAS"],
            ["Storage", "15-25%", "20-30%", "15-25%", "Unused table cleanup, TT reduction"],
            ["Serverless", "5-15%", "10-25%", "10-20%", "Clustering review, pipe consolidation"],
            ["Data Transfer", "2-5%", "5-15%", "5-10%", "Region alignment, stage cleanup"],
            ["Materialized Views", "1-5%", "10-30%", "10-20%", "Remove unused MVs, consolidate"],
        ],
    )

    doc.add_page_break()

    # =======================================================================
    # SECTION 3: SOLUTION ARCHITECTURE
    # =======================================================================
    doc.add_heading("3. Solution Architecture", level=1)

    doc.add_heading("3.1 Technology Stack", level=2)
    doc.add_paragraph(
        "The framework integrates with Thomson Reuters' existing stack — Snowflake on AWS, "
        "dbt for transformations, Immuta for governance, Alation for cataloguing, and "
        "SailPoint for access management."
    )
    add_styled_table(doc,
        ["Component", "Technology", "Rationale"],
        [
            ["Data Source", "SNOWFLAKE.ACCOUNT_USAGE\nSNOWFLAKE.ORGANIZATION_USAGE", "Native metadata — 365 days, zero cost. Covers all 4 layers (Raw, TRIM, MDS, Broad)"],
            ["Data Modelling", "dbt Core (open-source)", "Already used by TR BU teams — consistent tooling across platform and cost monitoring"],
            ["Dashboards", "Streamlit in Snowflake", "Native, interactive, shareable across BUs — no separate hosting"],
            ["Scheduling", "dbt Cloud / Snowflake Tasks", "Automated daily refresh using TR's existing orchestration patterns"],
            ["Infrastructure", "Terraform (optional)", "Reproducible deployment — aligned with TR's infrastructure practices"],
            ["Accelerators", "dbt-snowflake-monitoring\nSundeck OpsCenter", "Pre-built cost models, open-source — no licence cost (key concern from call)"],
            ["Governance Integration", "Immuta + Alation + SailPoint", "Cost insights tagged to governance policies; catalogue integration for data product discovery"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("3.2 Data Flow Architecture", level=2)
    add_ascii_diagram(doc, [
        "Layer 1: SOURCES                Layer 2: STAGING           Layer 3: INTERMEDIATE",
        "+-------------------------+     +-------------------+     +----------------------+",
        "| ACCOUNT_USAGE           |     | stg_warehouse_    |     | int_compute_cost     |",
        "|  - QUERY_HISTORY        |---->|   metering        |---->| int_storage_cost     |",
        "|  - WAREHOUSE_METERING   |     | stg_query_history |     | int_serverless_cost  |",
        "|  - TABLE_STORAGE_METRICS|     | stg_storage_      |     | int_query_attribution|",
        "|  - ACCESS_HISTORY       |     |   metrics         |     | int_anti_patterns    |",
        "|  - PIPE_USAGE_HISTORY   |     | stg_access_       |     | int_warehouse_util   |",
        "|  - AUTO_CLUSTERING_HIST |     |   history         |     | int_anomaly_scores   |",
        "|  - WAREHOUSE_EVENTS_HIST|     | stg_metering_     |     | int_chargeback       |",
        "|  - COPY_HISTORY         |     |   history         |     +----------------------+",
        "|  - WAREHOUSE_LOAD_HIST  |     | stg_pipe_usage    |              |",
        "+-------------------------+     | stg_copy_history  |              v",
        "                                +-------------------+     Layer 4: PUBLICATION",
        "+-------------------------+                               +----------------------+",
        "| ORGANIZATION_USAGE      |                               | pub_cost_summary     |",
        "|  - METERING_HISTORY     |----+                          | pub_cost_by_warehouse|",
        "|  - METERING_DAILY_HIST  |    |                          | pub_cost_by_team     |",
        "|  - RATE_SHEET_DAILY     |    |                          | pub_query_costs      |",
        "+-------------------------+    |                          | pub_savings_opps     |",
        "                               |                          | pub_anomalies        |",
        "+-------------------------+    |                          | pub_chargeback       |",
        "| QUERY_ATTRIBUTION       |    |                          | pub_recommendations  |",
        "|  - COST_PER_QUERY       |----+                          +----------------------+",
        "+-------------------------+                                        |",
        "                                                                   v",
        "                                                          +------------------+",
        "                                                          | STREAMLIT        |",
        "                                                          | DASHBOARD        |",
        "                                                          | (8 views)        |",
        "                                                          +------------------+",
    ])

    doc.add_paragraph()
    doc.add_heading("3.3 dbt Model Catalogue", level=2)
    add_styled_table(doc,
        ["Layer", "Model", "Materialisation", "Purpose"],
        [
            ["Staging", "stg_warehouse_metering", "view", "Clean WAREHOUSE_METERING_HISTORY"],
            ["Staging", "stg_query_history", "incremental", "Clean QUERY_HISTORY (incremental by end_time)"],
            ["Staging", "stg_storage_metrics", "view", "Clean TABLE_STORAGE_METRICS"],
            ["Staging", "stg_access_history", "incremental", "Clean ACCESS_HISTORY"],
            ["Staging", "stg_metering_history", "view", "Clean METERING_HISTORY (all service types)"],
            ["Staging", "stg_pipe_usage", "view", "Clean PIPE_USAGE_HISTORY"],
            ["Staging", "stg_copy_history", "incremental", "Clean COPY_HISTORY"],
            ["Staging", "stg_clustering_history", "view", "Clean AUTOMATIC_CLUSTERING_HISTORY"],
            ["Staging", "stg_warehouse_events", "incremental", "Clean WAREHOUSE_EVENTS_HISTORY"],
            ["Intermediate", "int_compute_cost", "table", "Compute cost by warehouse/hour with credit pricing"],
            ["Intermediate", "int_query_cost", "incremental", "Per-query cost attribution"],
            ["Intermediate", "int_storage_cost", "table", "Storage cost by database/schema/table"],
            ["Intermediate", "int_serverless_cost", "table", "Serverless credits by service type"],
            ["Intermediate", "int_anti_patterns", "incremental", "Query anti-pattern detection"],
            ["Intermediate", "int_warehouse_utilisation", "table", "Warehouse utilisation (idle, active, queued)"],
            ["Intermediate", "int_anomaly_scores", "incremental", "Cost anomaly scoring"],
            ["Intermediate", "int_chargeback", "table", "Tag-based cost allocation"],
            ["Publication", "pub_cost_summary", "table", "Executive cost overview"],
            ["Publication", "pub_cost_by_warehouse", "table", "Warehouse-level cost drill-down"],
            ["Publication", "pub_cost_by_team", "table", "Team-level cost attribution"],
            ["Publication", "pub_savings_opportunities", "table", "Ranked optimisation candidates"],
            ["Publication", "pub_anomalies", "table", "Cost anomalies for alerting"],
            ["Publication", "pub_recommendations", "table", "Prioritised action items"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("3.4 Design Principles", level=2)
    for item in [
        "Idempotent: All models can be re-run without side effects",
        "Incremental: Large tables (query history, access history) use incremental materialisation",
        "Tested: dbt tests for uniqueness, referential integrity, accepted values, and freshness",
        "Documented: Every model and column documented, browsable via dbt docs",
        "Configurable: Credit prices, team mappings, and thresholds maintained as dbt seed files",
        "Modular: Each cost category is independently analysable and composable",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # =======================================================================
    # SECTION 4: COST ANALYSIS FRAMEWORK
    # =======================================================================
    doc.add_heading("4. Cost Analysis Framework", level=1)
    doc.add_paragraph(
        "Snowflake provides three complementary views for cost analysis, each offering "
        "different granularity and coverage. Our framework unifies all three."
    )

    doc.add_heading("4.1 METERING_HISTORY — Unified Cost View", level=2)
    doc.add_paragraph(
        "SNOWFLAKE.ORGANIZATION_USAGE.METERING_HISTORY provides a single, unified view of "
        "ALL credit consumption across 28+ service types. This is the foundation of our "
        "cost analysis — it captures every billable activity in the account."
    )

    add_styled_table(doc,
        ["SERVICE_TYPE", "Category", "Description"],
        [
            ["WAREHOUSE_METERING", "Compute", "User-managed warehouse credits"],
            ["CLOUD_SERVICES", "Compute", "Cloud services layer (metadata ops, auth, etc.)"],
            ["QUERY_ACCELERATION", "Compute", "Query Acceleration Service credits"],
            ["AUTO_CLUSTERING", "Serverless", "Automatic reclustering credits"],
            ["MATERIALIZED_VIEW", "Serverless", "MV refresh credits"],
            ["PIPE", "Serverless", "Snowpipe continuous loading credits"],
            ["SERVERLESS_TASK", "Serverless", "Serverless task execution credits"],
            ["SEARCH_OPTIMIZATION", "Serverless", "Search optimisation service credits"],
            ["REPLICATION", "Data Transfer", "Database replication credits"],
            ["FAILOVER", "Data Transfer", "Failover group credits"],
            ["SNOWPIPE_STREAMING", "Serverless", "Snowpipe Streaming credits"],
            ["LOGGING", "Serverless", "Event table logging credits"],
            ["READER_WAREHOUSE_METERING", "Compute", "Reader account warehouse credits"],
            ["EXTERNAL_FUNCTIONS", "Serverless", "External function invocation credits"],
            ["WAREHOUSE_METERING_READER", "Compute", "Reader warehouse (legacy)"],
        ],
    )

    doc.add_paragraph()
    add_sql_block(doc,
        "-- Unified cost view from METERING_HISTORY\n"
        "SELECT\n"
        "    service_type,\n"
        "    DATE_TRUNC('month', usage_date) AS usage_month,\n"
        "    SUM(credits_used) AS total_credits,\n"
        "    SUM(credits_used) * :credit_price AS total_cost_usd\n"
        "FROM snowflake.organization_usage.metering_history\n"
        "WHERE usage_date >= DATEADD('day', -90, CURRENT_DATE())\n"
        "GROUP BY 1, 2\n"
        "ORDER BY total_credits DESC;"
    )

    doc.add_paragraph()
    doc.add_heading("4.2 METERING_DAILY_HISTORY — Billed Credits", level=2)
    doc.add_paragraph(
        "SNOWFLAKE.ORGANIZATION_USAGE.METERING_DAILY_HISTORY provides the actual billed "
        "credit amounts per service type per day. This is the definitive source for "
        "reconciling against Snowflake invoices."
    )
    add_sql_block(doc,
        "-- Daily billed credits with dollar conversion\n"
        "SELECT\n"
        "    usage_date,\n"
        "    service_type,\n"
        "    credits_used_compute,\n"
        "    credits_used_cloud_services,\n"
        "    credits_billed,  -- Net after cloud services adjustment\n"
        "    credits_billed * :credit_price AS billed_cost_usd\n"
        "FROM snowflake.organization_usage.metering_daily_history\n"
        "WHERE usage_date >= DATEADD('day', -30, CURRENT_DATE())\n"
        "ORDER BY usage_date DESC;"
    )

    doc.add_paragraph()
    doc.add_heading("4.3 QUERY_ATTRIBUTION_HISTORY — Per-Query Dollar Cost", level=2)
    doc.add_paragraph(
        "SNOWFLAKE.ACCOUNT_USAGE.QUERY_ATTRIBUTION_HISTORY (available on Enterprise Edition+) "
        "provides Snowflake's own per-query cost attribution. This replaces estimated cost "
        "calculations with precise dollar-per-query data."
    )
    add_callout_box(doc,
        "QUERY_ATTRIBUTION_HISTORY is the gold standard for per-query costing. "
        "It attributes compute, cloud services, and query acceleration credits directly "
        "to individual queries — no estimation required."
    )
    add_sql_block(doc,
        "-- Per-query cost from QUERY_ATTRIBUTION_HISTORY\n"
        "SELECT\n"
        "    qah.query_id,\n"
        "    qh.warehouse_name,\n"
        "    qh.user_name,\n"
        "    qh.query_tag,\n"
        "    qah.credits_attributed_compute,\n"
        "    qah.credits_attributed_cloud_services,\n"
        "    qah.credits_attributed_query_acceleration,\n"
        "    (qah.credits_attributed_compute\n"
        "     + qah.credits_attributed_cloud_services\n"
        "     + qah.credits_attributed_query_acceleration\n"
        "    ) * :credit_price AS total_query_cost_usd\n"
        "FROM snowflake.account_usage.query_attribution_history qah\n"
        "JOIN snowflake.account_usage.query_history qh\n"
        "  ON qah.query_id = qh.query_id\n"
        "WHERE qh.start_time >= DATEADD('day', -7, CURRENT_DATE())\n"
        "ORDER BY total_query_cost_usd DESC\n"
        "LIMIT 100;"
    )

    doc.add_paragraph()
    doc.add_heading("4.4 Complete Cost Category Breakdown", level=2)
    doc.add_paragraph(
        "Our framework analyses 8+ distinct cost categories, each with dedicated "
        "source views and SQL patterns:"
    )
    add_styled_table(doc,
        ["Category", "Source View(s)", "Key Metrics", "Analysis Approach"],
        [
            ["Compute", "WAREHOUSE_METERING_HISTORY\nQUERY_HISTORY", "Credits/hour, cost/query, idle %", "Warehouse-level + query-level attribution"],
            ["Storage", "TABLE_STORAGE_METRICS\nSTORAGE_USAGE\nDATABASE_STORAGE_USAGE_HISTORY", "Active, TT, fail-safe bytes", "Table-level breakdown + trend analysis"],
            ["Serverless", "METERING_HISTORY\n(filtered by service_type)", "Credits by service type", "Per-service cost tracking"],
            ["Data Transfer", "DATA_TRANSFER_HISTORY", "Bytes transferred, credits", "Cross-region and cross-cloud analysis"],
            ["Materialized Views", "MATERIALIZED_VIEW_REFRESH_HISTORY", "Refresh credits, frequency", "ROI analysis: refresh cost vs query savings"],
            ["Search Optimisation", "SEARCH_OPTIMIZATION_HISTORY", "Maintenance credits", "Cost vs search performance improvement"],
            ["Replication", "REPLICATION_USAGE_HISTORY\nMETERING_HISTORY", "Replication credits, bytes", "Replication necessity and frequency review"],
            ["Query Acceleration", "QUERY_ACCELERATION_HISTORY\nQUERY_ATTRIBUTION_HISTORY", "QAS credits attributed", "QAS ROI: cost of acceleration vs time saved"],
        ],
    )

    doc.add_page_break()

    # =======================================================================
    # SECTION 5: PHASE 1 — COST VISIBILITY
    # =======================================================================
    doc.add_heading("5. Phase 1: Identify & Solve", level=1)
    doc.add_paragraph(
        "Objective: Diagnose exactly where money goes and why, map user query patterns across "
        "verticals, and deliver a cost dashboard with a prioritised findings report including "
        "root cause analysis, quick-win recommendations, and data product opportunities."
    )

    doc.add_heading("5.1 Compute Cost Analysis", level=2)
    doc.add_paragraph("Data Sources: WAREHOUSE_METERING_HISTORY, QUERY_HISTORY, QUERY_ATTRIBUTION_HISTORY")

    doc.add_heading("Idle vs Active Compute", level=3)
    add_styled_table(doc,
        ["Metric", "Source", "Description"],
        [
            ["Active credits", "WAREHOUSE_METERING_HISTORY", "Credits consumed while executing queries"],
            ["Idle credits", "WH_METERING - QUERY_HISTORY join", "Credits consumed while warehouse running but no queries executing"],
            ["Idle ratio", "Calculated", "idle_credits / total_credits — target < 10%"],
            ["Queue wait time", "QUERY_HISTORY.queued_overload_time", "Time queries spent waiting for compute slots"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("Compilation vs Execution Time", level=3)
    doc.add_paragraph(
        "Snowflake query execution includes compilation (parsing, optimisation) and execution "
        "(actual data processing). High compilation-to-execution ratios indicate queries that "
        "could benefit from prepared statements or query caching."
    )
    add_sql_block(doc,
        "SELECT\n"
        "    warehouse_name,\n"
        "    AVG(compilation_time) AS avg_compilation_ms,\n"
        "    AVG(execution_time) AS avg_execution_ms,\n"
        "    AVG(compilation_time) / NULLIF(AVG(execution_time), 0) AS compile_exec_ratio,\n"
        "    -- Ratio > 1.0 = more time compiling than executing\n"
        "    COUNT(*) AS query_count\n"
        "FROM snowflake.account_usage.query_history\n"
        "WHERE start_time >= DATEADD('day', -30, CURRENT_DATE())\n"
        "GROUP BY warehouse_name\n"
        "ORDER BY compile_exec_ratio DESC;"
    )

    doc.add_paragraph()
    doc.add_heading("Cache Hit Rate", level=3)
    doc.add_paragraph(
        "Snowflake caches results at multiple levels. High cache hit rates indicate efficient "
        "repeated query patterns; low rates suggest opportunities for materialisation."
    )
    add_sql_block(doc,
        "SELECT\n"
        "    warehouse_name,\n"
        "    COUNT_IF(bytes_scanned = 0 AND rows_produced > 0) AS result_cache_hits,\n"
        "    COUNT(*) AS total_queries,\n"
        "    result_cache_hits / NULLIF(total_queries, 0) AS cache_hit_rate\n"
        "FROM snowflake.account_usage.query_history\n"
        "WHERE start_time >= DATEADD('day', -30, CURRENT_DATE())\n"
        "  AND query_type = 'SELECT'\n"
        "GROUP BY warehouse_name\n"
        "ORDER BY cache_hit_rate;"
    )

    doc.add_paragraph()
    doc.add_heading("5.2 Storage Cost Analysis", level=2)
    doc.add_paragraph("Data Sources: TABLE_STORAGE_METRICS, STORAGE_USAGE, COPY_HISTORY")

    add_styled_table(doc,
        ["Metric", "Source", "Description"],
        [
            ["Active storage", "TABLE_STORAGE_METRICS", "Data currently in use"],
            ["Time Travel storage", "TABLE_STORAGE_METRICS", "Historical data for Time Travel window"],
            ["Fail-safe storage", "TABLE_STORAGE_METRICS", "7-day regulatory recovery (non-configurable)"],
            ["Storage by database/schema", "DATABASE_STORAGE_USAGE_HISTORY", "Hierarchical storage breakdown"],
            ["Unused tables", "ACCESS_HISTORY (no reads 90+ days)", "Candidates for archival or deletion"],
            ["Clone overhead", "TABLE_STORAGE_METRICS", "Storage from diverged zero-copy clones"],
            ["Data loading volume", "COPY_HISTORY", "Volume of data ingested via COPY INTO"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("COPY_HISTORY Analysis", level=3)
    add_sql_block(doc,
        "-- Data loading volume and costs\n"
        "SELECT\n"
        "    table_name,\n"
        "    DATE_TRUNC('day', last_load_time) AS load_date,\n"
        "    COUNT(*) AS files_loaded,\n"
        "    SUM(row_count) AS rows_loaded,\n"
        "    SUM(file_size) / POWER(1024, 3) AS gb_loaded\n"
        "FROM snowflake.account_usage.copy_history\n"
        "WHERE last_load_time >= DATEADD('day', -30, CURRENT_DATE())\n"
        "GROUP BY 1, 2\n"
        "ORDER BY gb_loaded DESC;"
    )

    doc.add_paragraph()
    doc.add_heading("5.3 Serverless Cost Analysis", level=2)
    doc.add_paragraph(
        "Serverless features consume credits without explicit warehouse provisioning. "
        "We track all serverless service types from METERING_HISTORY:"
    )
    add_styled_table(doc,
        ["Service", "Source View", "Key Metric", "Optimisation Lever"],
        [
            ["Snowpipe", "PIPE_USAGE_HISTORY", "Credits/day per pipe", "Consolidate pipes, batch loading"],
            ["Auto-clustering", "AUTOMATIC_CLUSTERING_HISTORY", "Credits/table/day", "Review clustering keys, reduce over-clustering"],
            ["MV Refresh", "MATERIALIZED_VIEW_REFRESH_HISTORY", "Credits/refresh", "Drop unused MVs, consolidate"],
            ["Serverless Tasks", "SERVERLESS_TASK_HISTORY", "Credits/execution", "Optimise task SQL, reduce frequency"],
            ["Search Optimisation", "SEARCH_OPTIMIZATION_HISTORY", "Credits/table", "Review search access patterns"],
            ["Snowpipe Streaming", "METERING_HISTORY", "Credits/day", "Review streaming necessity"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("5.4 Tag-Based Attribution Strategy", level=2)
    doc.add_paragraph(
        "Snowflake object tags enable cost attribution to teams, products, and cost centres. "
        "Our framework implements a three-level tagging strategy:"
    )
    add_styled_table(doc,
        ["Tag Level", "Tag Name", "Applied To", "Example Values"],
        [
            ["L1: Cost Centre", "cost_centre", "Warehouses, Databases", "engineering, analytics, data-science"],
            ["L2: Team", "team", "Warehouses, Schemas", "platform, reporting, ml-ops"],
            ["L3: Data Product", "data_product", "Schemas, Tables", "customer-360, revenue-model, risk-engine"],
        ],
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "For accounts not yet using tags, attribution falls back to warehouse-to-team "
        "mapping (maintained as a dbt seed file) and role-based allocation."
    )

    doc.add_paragraph()
    doc.add_heading("5.5 Dashboard Wireframes", level=2)
    doc.add_paragraph("Phase 1 delivers five interactive dashboard views:")

    doc.add_heading("Executive Summary View", level=3)
    add_ascii_diagram(doc, [
        "+-------------------------------+-------------------------------+",
        "| TOTAL MONTHLY SPEND           | COST TREND (90 DAYS)          |",
        "|  $124,500 (+8% MoM)           |  ___                          |",
        "|                               | /   \\___    ___/              |",
        "|  Compute:  $89,200 (72%)      |/        \\__/                  |",
        "|  Storage:  $22,100 (18%)      |                               |",
        "|  Svrless:  $13,200 (10%)      | Jan  Feb  Mar  Apr            |",
        "+-------------------------------+-------------------------------+",
        "| TOP 5 WAREHOUSES              | COST BY TEAM                  |",
        "|  1. ANALYTICS_WH    $31,200   |  Engineering    45%  ===      |",
        "|  2. ETL_WH          $24,800   |  Analytics      28%  ==       |",
        "|  3. REPORTING_WH    $15,600   |  Data Science   18%  =        |",
        "|  4. DATA_SCIENCE_WH $11,400   |  Ad Hoc          9%           |",
        "|  5. AD_HOC_WH        $6,200   |                               |",
        "+-------------------------------+-------------------------------+",
    ])

    doc.add_paragraph()
    doc.add_heading("Warehouse Deep-Dive View", level=3)
    add_ascii_diagram(doc, [
        "+-------------------------------------------------------------------+",
        "| WAREHOUSE: ANALYTICS_WH                          Cost: $31,200/mo |",
        "+-------------------------------------------------------------------+",
        "| Utilisation     | Idle: 34%  Active: 58%  Queued: 8%              |",
        "| Size            | LARGE (8 credits/hr)                            |",
        "| Auto-suspend    | 300s (Recommendation: reduce to 60s)            |",
        "+-------------------------------------------------------------------+",
        "| HOURLY CREDIT USAGE            | TOP USERS                        |",
        "|   __                           |  svc_etl        42%              |",
        "|  /  \\     __                   |  analyst_1      23%              |",
        "| /    \\___/  \\                  |  analyst_2      15%              |",
        "|/             \\___              |  data_eng_1     11%              |",
        "| 6am   12pm  6pm  12am          |  other           9%              |",
        "+--------------------------------+----------------------------------+",
    ])

    doc.add_page_break()

    # =======================================================================
    # SECTION 6: PHASE 2 — QUERY OPTIMISATION
    # =======================================================================
    doc.add_heading("6. Phase 2: Query Optimisation", level=1)
    doc.add_paragraph(
        "Phase 2 identifies why costs are high and provides actionable, prioritised "
        "recommendations with estimated dollar savings."
    )

    doc.add_heading("6.1 Query Plan Analysis — GET_QUERY_OPERATOR_STATS()", level=2)
    doc.add_paragraph(
        "GET_QUERY_OPERATOR_STATS() returns the execution plan operators for a completed query, "
        "including row counts, bytes, and time per operator. This enables identification of "
        "the most expensive operations within a query."
    )
    add_sql_block(doc,
        "-- Analyse query plan operators for expensive queries\n"
        "SELECT *\n"
        "FROM TABLE(GET_QUERY_OPERATOR_STATS(LAST_QUERY_ID()));\n"
        "\n"
        "-- Analyse a specific query by ID\n"
        "SELECT\n"
        "    operator_id,\n"
        "    operator_type,\n"
        "    operator_statistics:output_rows::NUMBER AS output_rows,\n"
        "    operator_statistics:input_rows::NUMBER AS input_rows,\n"
        "    execution_time_breakdown:overall_percentage::FLOAT AS pct_of_total\n"
        "FROM TABLE(GET_QUERY_OPERATOR_STATS(:query_id))\n"
        "ORDER BY pct_of_total DESC;"
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "Key patterns to detect from operator stats:"
    )
    add_styled_table(doc,
        ["Operator Pattern", "Indicates", "Action"],
        [
            ["TableScan with high output_rows", "Full table scan", "Add clustering key or filter pushdown"],
            ["Sort with high input_rows", "Expensive sort (ORDER BY)", "Add LIMIT or remove unnecessary ORDER BY"],
            ["JoinFilter with row explosion", "Cartesian/cross join", "Fix join condition"],
            ["Aggregate with spill", "Insufficient memory", "Upsize warehouse or simplify aggregation"],
            ["WindowFunction with full partition", "Window over entire table", "Add PARTITION BY to window function"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("6.2 Clustering Analysis", level=2)
    doc.add_paragraph(
        "Snowflake provides two system functions for evaluating table clustering effectiveness:"
    )

    doc.add_heading("SYSTEM$CLUSTERING_INFORMATION()", level=3)
    add_sql_block(doc,
        "-- Check clustering quality for a table\n"
        "SELECT SYSTEM$CLUSTERING_INFORMATION('my_db.my_schema.my_table', '(date_col)');\n"
        "\n"
        "-- Returns JSON with:\n"
        "-- cluster_by_keys, total_partition_count, total_constant_partition_count,\n"
        "-- average_overlaps, average_depth, partition_depth_histogram")

    doc.add_paragraph()
    doc.add_heading("SYSTEM$CLUSTERING_DEPTH()", level=3)
    add_sql_block(doc,
        "-- Quick clustering depth check\n"
        "SELECT SYSTEM$CLUSTERING_DEPTH('my_db.my_schema.my_table', '(date_col)');\n"
        "\n"
        "-- Returns a single number:\n"
        "-- 1.0 = perfectly clustered\n"
        "-- > 3.0 = consider reclustering or changing cluster key")

    doc.add_paragraph(
        "Our framework profiles all tables with auto-clustering enabled and evaluates "
        "whether the clustering cost is justified by query performance improvements."
    )

    doc.add_paragraph()
    doc.add_heading("6.3 Query Acceleration Service (QAS) Estimation", level=2)
    doc.add_paragraph(
        "SYSTEM$ESTIMATE_QUERY_ACCELERATION() estimates whether a query would benefit from "
        "the Query Acceleration Service and the expected speedup."
    )
    add_sql_block(doc,
        "-- Estimate QAS benefit for a specific query\n"
        "SELECT SYSTEM$ESTIMATE_QUERY_ACCELERATION(:query_id);\n"
        "\n"
        "-- Returns JSON with:\n"
        "-- originalQueryTime, queryAccelerationUpperBound (max speedup),\n"
        "-- acceleratedQueryTimes for different scale factors (1x-24x)\n"
        "\n"
        "-- Batch analysis: find top QAS candidates\n"
        "SELECT\n"
        "    query_id,\n"
        "    query_text,\n"
        "    total_elapsed_time,\n"
        "    SYSTEM$ESTIMATE_QUERY_ACCELERATION(query_id) AS qas_estimate\n"
        "FROM snowflake.account_usage.query_history\n"
        "WHERE start_time >= DATEADD('day', -7, CURRENT_DATE())\n"
        "  AND total_elapsed_time > 60000  -- queries > 1 minute\n"
        "ORDER BY total_elapsed_time DESC\n"
        "LIMIT 50;"
    )

    doc.add_paragraph()
    doc.add_heading("6.4 Anti-Pattern Detection", level=2)
    add_styled_table(doc,
        ["Anti-Pattern", "Detection SQL", "Impact", "Recommendation"],
        [
            ["Full table scans", "partitions_scanned/partitions_total > 0.8", "Excessive compute", "Add clustering key or WHERE filters"],
            ["SELECT *", "Query text LIKE 'SELECT *%'", "Scans all columns", "Select only needed columns"],
            ["Missing filters", "bytes_scanned / rows_produced > threshold", "Reads excess data", "Add WHERE clause, partition pruning"],
            ["Spill to storage", "bytes_spilled_to_local/remote > 0", "Memory exhaustion", "Upsize warehouse or simplify query"],
            ["Repeated queries", "GROUP BY query_parameterized_hash", "Wasted compute", "Cache results or materialise"],
            ["Cartesian joins", "rows_produced >> rows_scanned", "Join explosion", "Fix join condition"],
            ["Large ORDER BY", "Large result + ORDER BY without LIMIT", "Unnecessary sorting", "Add LIMIT or remove ORDER BY"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("Parameterized Hash Grouping", level=3)
    doc.add_paragraph(
        "QUERY_PARAMETERIZED_HASH groups queries with identical structure but different "
        "literal values — essential for identifying repeated query patterns:"
    )
    add_sql_block(doc,
        "-- Find most frequently repeated query patterns\n"
        "SELECT\n"
        "    query_parameterized_hash,\n"
        "    COUNT(*) AS execution_count,\n"
        "    SUM(total_elapsed_time) / 1000 AS total_seconds,\n"
        "    SUM(credits_used_cloud_services) AS total_credits,\n"
        "    ANY_VALUE(query_text) AS sample_query\n"
        "FROM snowflake.account_usage.query_history\n"
        "WHERE start_time >= DATEADD('day', -7, CURRENT_DATE())\n"
        "GROUP BY query_parameterized_hash\n"
        "HAVING execution_count > 100\n"
        "ORDER BY total_credits DESC\n"
        "LIMIT 20;"
    )

    doc.add_paragraph()
    doc.add_heading("6.5 Warehouse Right-Sizing", level=2)
    add_styled_table(doc,
        ["Signal", "Detection Source", "Detection Method", "Recommendation"],
        [
            ["Consistent queuing", "QUERY_HISTORY", "queued_overload_time > 0 frequently", "Scale up or enable multi-cluster"],
            ["High idle time", "WAREHOUSE_METERING vs QUERY_HISTORY", "Warehouse running, no queries", "Reduce auto-suspend (300s -> 60s)"],
            ["Oversized", "QUERY_HISTORY", "Small queries on L/XL warehouses", "Downsize or route to smaller WH"],
            ["Uneven load", "WAREHOUSE_LOAD_HISTORY", "Spikes at certain hours", "Consolidate usage windows"],
            ["Suspend/resume churn", "WAREHOUSE_EVENTS_HISTORY", "Frequent suspend/resume cycles", "Increase auto-suspend interval"],
            ["Lock contention", "LOCK_WAIT_HISTORY", "Queries waiting for locks", "Separate read/write warehouses"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("WAREHOUSE_EVENTS_HISTORY Analysis", level=3)
    add_sql_block(doc,
        "-- Warehouse suspend/resume frequency\n"
        "SELECT\n"
        "    warehouse_name,\n"
        "    event_name,\n"
        "    DATE_TRUNC('day', timestamp) AS event_date,\n"
        "    COUNT(*) AS event_count\n"
        "FROM snowflake.account_usage.warehouse_events_history\n"
        "WHERE timestamp >= DATEADD('day', -30, CURRENT_DATE())\n"
        "  AND event_name IN ('SUSPEND_WAREHOUSE', 'RESUME_WAREHOUSE')\n"
        "GROUP BY 1, 2, 3\n"
        "ORDER BY event_count DESC;"
    )

    doc.add_paragraph()
    doc.add_heading("6.6 Storage Optimisation", level=2)
    add_styled_table(doc,
        ["Opportunity", "Detection", "Savings Estimate", "Recommendation"],
        [
            ["Unused tables", "No reads in ACCESS_HISTORY 90+ days", "100% of table storage", "Archive or drop"],
            ["Excessive Time Travel", "Retention > 1 day on non-critical tables", "Up to 99% of TT storage", "Reduce retention to 1 day"],
            ["Transient candidates", "Tables rebuilt daily", "Eliminates fail-safe", "Convert to TRANSIENT"],
            ["Stale clones", "Clones not accessed in 30+ days", "100% of diverged storage", "Drop stale clones"],
            ["Large staging tables", "Temp data persisted permanently", "100% of staging storage", "Add lifecycle/auto-drop"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("6.7 Query Pattern Analysis & Tiered Recommendation Engine", level=2)
    doc.add_paragraph(
        "This is the most impactful capability in the framework. Rather than building "
        "curated views for every problem (which takes days and doesn't scale across 350+ "
        "sources), we mine Snowflake's query history to understand what's expensive and why "
        "— then assign the right fix tier to each problem. 80% of fixes deploy in minutes. "
        "Only the top 5-10 patterns justify a full curated data product."
    )

    doc.add_paragraph()
    doc.add_heading("The Core Insight: Speed of Fix Matters More Than Elegance", level=3)
    doc.add_paragraph(
        "Building a curated view or materialised table takes 1-5 days. If we recommend "
        "building views for every expensive pattern, nothing ships fast enough to matter. "
        "The tiered approach matches each problem to the fastest fix that delivers the saving:"
    )
    add_styled_table(doc,
        ["Tier", "Fix Type", "Build Time", "Example", "When to Use"],
        [
            ["Tier 1", "Fix the contract in place", "Minutes", "Send BU owner exact ALTER/patch: add clustering key, fix join order, add partition filter", "Contract is 80% right but missing one optimisation"],
            ["Tier 2", "Enable Snowflake native feature", "Minutes", "Enable result cache, create Materialized View, turn on QAS, adjust auto-suspend", "Snowflake already has the answer — just configure it"],
            ["Tier 3", "Lightweight SQL view", "1-2 hours", "CREATE VIEW with optimised SQL, clustering, partition awareness", "5+ users run similar queries; view replaces all of them"],
            ["Tier 4", "Full curated data product", "2-5 days", "dbt model + tests + Alation entry + Immuta policy + refresh schedule", "High-value, cross-BU pattern with proven ROI (top 5-10 only)"],
        ],
    )

    doc.add_heading("User Query Pattern Profiling", level=3)
    add_sql_block(doc,
        "-- Profile user query patterns: who queries what tables, how often, at what cost\n"
        "SELECT\n"
        "    qh.user_name,\n"
        "    qh.role_name,\n"
        "    qh.warehouse_name,\n"
        "    f.value:objectName::STRING AS table_accessed,\n"
        "    COUNT(DISTINCT qh.query_id) AS query_count,\n"
        "    SUM(qh.total_elapsed_time) / 1000 AS total_seconds,\n"
        "    SUM(qh.credits_used_cloud_services) AS total_credits\n"
        "FROM snowflake.account_usage.query_history qh,\n"
        "     LATERAL FLATTEN(input => PARSE_JSON(qh.direct_objects_accessed)) f\n"
        "WHERE qh.start_time >= DATEADD('day', -30, CURRENT_DATE())\n"
        "GROUP BY 1, 2, 3, 4\n"
        "ORDER BY total_credits DESC;"
    )

    doc.add_paragraph()
    doc.add_heading("Identifying Repeated Patterns & Assigning Tiers", level=3)
    doc.add_paragraph(
        "By grouping queries by QUERY_PARAMETERIZED_HASH and the tables they access, we "
        "identify clusters of users asking the same business question in different ways. "
        "Each cluster is then automatically assigned an optimisation tier based on fix complexity."
    )
    add_sql_block(doc,
        "-- Find expensive query clusters and assign optimisation tiers\n"
        "WITH query_clusters AS (\n"
        "    SELECT\n"
        "        query_parameterized_hash,\n"
        "        COUNT(DISTINCT user_name) AS distinct_users,\n"
        "        COUNT(*) AS total_executions,\n"
        "        SUM(total_elapsed_time) / 1000 AS total_seconds,\n"
        "        SUM(bytes_scanned) / POWER(1024, 3) AS total_gb_scanned,\n"
        "        ANY_VALUE(warehouse_name) AS sample_warehouse,\n"
        "        ANY_VALUE(LEFT(query_text, 300)) AS sample_query\n"
        "    FROM snowflake.account_usage.query_history\n"
        "    WHERE start_time >= DATEADD('day', -30, CURRENT_DATE())\n"
        "      AND query_type = 'SELECT'\n"
        "    GROUP BY query_parameterized_hash\n"
        ")\n"
        "SELECT\n"
        "    *,\n"
        "    CASE\n"
        "        -- Tier 1: Contract fix (missing clustering/partition filter)\n"
        "        WHEN total_gb_scanned > 10 AND distinct_users = 1\n"
        "            THEN 'Tier 1: Fix contract in place'\n"
        "        -- Tier 2: Native feature (result cache, MV, QAS)\n"
        "        WHEN total_executions > 100 AND total_gb_scanned < 5\n"
        "            THEN 'Tier 2: Enable result cache / MV'\n"
        "        -- Tier 3: Lightweight view (multiple users, same pattern)\n"
        "        WHEN distinct_users >= 3 AND total_executions >= 20\n"
        "            THEN 'Tier 3: Create lightweight view'\n"
        "        -- Tier 4: Full data product (cross-BU, high value)\n"
        "        WHEN distinct_users >= 5 AND total_gb_scanned > 50\n"
        "            THEN 'Tier 4: Full curated data product'\n"
        "        ELSE 'Tier 1: Review and patch'\n"
        "    END AS recommended_tier\n"
        "FROM query_clusters\n"
        "WHERE total_executions >= 10\n"
        "ORDER BY total_gb_scanned DESC\n"
        "LIMIT 50;"
    )

    doc.add_paragraph()
    doc.add_heading("Opportunity Scoring by Tier", level=3)
    doc.add_paragraph(
        "Each query cluster is scored on four dimensions to prioritise the recommendation report "
        "and assign the correct optimisation tier:"
    )
    add_styled_table(doc,
        ["Dimension", "Metric", "Tier Assignment Impact"],
        [
            ["Cost impact", "Total compute credits consumed by this pattern", "Higher cost = higher priority regardless of tier"],
            ["User reach", "Number of distinct users running similar queries", "1 user = Tier 1 fix; 3+ users = Tier 3 view; 5+ cross-BU = Tier 4"],
            ["Fix complexity", "Missing clustering key vs full rewrite needed", "Simple fix = Tier 1; needs new SQL = Tier 3/4"],
            ["Data freshness need", "How frequently users re-run the same query", "High frequency = MV (Tier 2); low = view (Tier 3)"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("Tiered Fix Examples", level=3)
    doc.add_paragraph(
        "For each tier, here's what the recommendation looks like in practice:"
    )
    add_sql_block(doc,
        "-- TIER 1 EXAMPLE: Fix contract in place (send to BU owner)\n"
        "-- Problem: Finance BU contract scans full table (no partition filter)\n"
        "-- Current cost: $800/month\n"
        "-- Fix: Add date partition filter. Estimated new cost: $25/month\n"
        "\n"
        "-- Original (in BU's dbt contract):\n"
        "-- SELECT * FROM raw.transactions WHERE status = 'active'\n"
        "\n"
        "-- Recommended fix (sent as patch to BU owner):\n"
        "-- SELECT * FROM raw.transactions\n"
        "-- WHERE status = 'active'\n"
        "--   AND transaction_date >= DATEADD('day', -90, CURRENT_DATE())\n"
        "--   -- ^ Partition pruning: scans 3 months instead of full history\n"
        "\n"
        "-- TIER 2 EXAMPLE: Enable Materialized View (config change, no code)\n"
        "-- Problem: 200 users run same aggregation daily\n"
        "-- Current cost: $1,500/month (200 full scans)\n"
        "-- Fix: MV refreshes once, all 200 users read from cache\n"
        "\n"
        "CREATE MATERIALIZED VIEW IF NOT EXISTS mv_daily_revenue AS\n"
        "SELECT region, product_category, order_date,\n"
        "       SUM(amount) AS total_revenue, COUNT(*) AS order_count\n"
        "FROM raw.orders\n"
        "GROUP BY 1, 2, 3;\n"
        "-- Estimated new cost: $30/month (1 refresh/day)\n"
        "\n"
        "-- TIER 3 EXAMPLE: Lightweight SQL view (1-2 hours to build)\n"
        "-- Problem: 8 users across 2 BUs query orders+customers with different joins\n"
        "-- Current cost: $1,200/month (4,500 ad-hoc queries)\n"
        "-- Fix: Single optimised view replaces all variations\n"
        "\n"
        "CREATE OR REPLACE VIEW optimised_views.order_analytics_daily AS\n"
        "SELECT o.order_date, o.region, o.product_category,\n"
        "       c.customer_segment,\n"
        "       COUNT(DISTINCT o.order_id) AS order_count,\n"
        "       SUM(o.order_amount) AS total_revenue\n"
        "FROM raw.orders o\n"
        "JOIN raw.customers c ON o.customer_id = c.customer_id\n"
        "GROUP BY 1, 2, 3, 4;\n"
        "-- Estimated new cost: $30/month (1 scheduled refresh/day)"
    )

    doc.add_paragraph()
    doc.add_heading("Tiered Recommendation Report", level=3)
    doc.add_paragraph(
        "For each identified pattern, the recommendation report provides:"
    )
    add_styled_table(doc,
        ["Field", "Description"],
        [
            ["Pattern ID", "QUERY_PARAMETERIZED_HASH cluster identifier"],
            ["Business Intent", "What question users are trying to answer (extracted from query patterns)"],
            ["Source Tables", "Which Raw/TRIM tables the pattern queries"],
            ["Current Impact", "Number of users, BUs, query count, monthly cost"],
            ["Assigned Tier", "Tier 1 (contract fix) / Tier 2 (native feature) / Tier 3 (view) / Tier 4 (data product)"],
            ["Recommended Fix", "Exact SQL patch, config change, or view definition — ready to deploy"],
            ["Estimated Build Time", "Minutes (Tier 1-2) / Hours (Tier 3) / Days (Tier 4)"],
            ["Estimated Monthly Saving", "Current cost minus fix cost — typically 10-50x reduction"],
            ["Owner", "BU owner to receive the fix (Tier 1) or platform team (Tier 3-4)"],
            ["Priority", "Ranked by cost impact × ease of fix — high-value, fast fixes first"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("6.8 Governed Data Product Lifecycle (Contract Generation)", level=2)
    doc.add_paragraph(
        "The tiered recommendation engine solves today's expensive queries. But the deeper "
        "question is: how do we prevent the same problem from recurring when new data products "
        "are created tomorrow? The answer is NOT 'build a view first and hope people find it' "
        "— that takes too long and doesn't scale. Instead, we generate optimised contracts "
        "automatically from schema metadata, so every new data product is efficient by default."
    )

    doc.add_heading("Current State: No Gate Before Production", level=3)
    doc.add_paragraph(
        "Today, when a BU team needs data, they go to their MDS workspace, write a dbt "
        "contract against Raw/TRIM tables, and execute it. There is no step that checks:"
    )
    for item in [
        "Does a contract or data product already exist that answers this question?",
        "Is this contract written in an optimised way (partition-aware, clustered, efficient joins)?",
        "What will this cost per month to run? Is that acceptable?",
        "Could another BU benefit from the same data (should this be in the Broad layer)?",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph(
        "The result: duplicate effort across BUs, no cost visibility until the invoice arrives, "
        "and a growing library of unoptimised contracts that the platform team cannot review "
        "manually across 350+ sources."
    )

    doc.add_heading("Target State: Metadata-Aware Contract Generation", level=3)
    doc.add_paragraph(
        "The key insight: users shouldn't have to learn about partition columns, clustering "
        "keys, or cost estimation. The framework handles it. When a BU team needs a new "
        "contract, the system generates one that's already optimised — in minutes, not days."
    )
    add_ascii_diagram(doc, [
        "  NEW DATA PRODUCT LIFECYCLE",
        "",
        "  [1] OVERLAP CHECK            [2] CONTRACT GENERATOR      [3] COST PRE-SCORE",
        "  +------------------+         +------------------+        +------------------+",
        "  | Framework checks |  No     | Reads            |        | Before deploy:   |",
        "  | existing         | overlap | INFORMATION_     |        | 'This contract   |",
        "  | contracts:       |-------->| SCHEMA:          |------->|  will cost ~$45/ |",
        "  | 'Does this data  |         | - Table sizes    |        |  mo. Finance BU  |",
        "  |  already exist?' |         | - Partition cols |        |  has similar at  |",
        "  +------------------+         | - Clustering keys|        |  $12/mo. Reuse?' |",
        "       |                       | - Column stats   |        +------------------+",
        "       | Match found           |                  |               |",
        "       v                       | Generates:       |               v",
        "  [1a] REUSE/EXTEND           | - Optimised SQL  |        [4] DEPLOY + MONITOR",
        "  +------------------+         | - Partition-aware |        +------------------+",
        "  | Extend existing  |         |   filters        |        | - Alation entry  |",
        "  | contract or use  |         | - Clustering key |        | - Immuta policies|",
        "  | as-is. Zero new  |         | - Cost estimate  |        | - Cost tracking  |",
        "  | build time.      |         +------------------+        | - Drift detection|",
        "  +------------------+                                     +------------------+",
    ])

    doc.add_heading("Metadata-Aware Contract Generator: How It Works", level=3)
    doc.add_paragraph(
        "The contract generator reads Snowflake's INFORMATION_SCHEMA to understand the "
        "source tables, then produces an optimised dbt contract automatically:"
    )
    add_sql_block(doc,
        "-- Step 1: Read source table metadata\n"
        "SELECT\n"
        "    t.table_name,\n"
        "    t.row_count,\n"
        "    t.bytes AS table_size_bytes,\n"
        "    t.clustering_key,\n"
        "    t.auto_clustering_on,\n"
        "    t.retention_time\n"
        "FROM information_schema.tables t\n"
        "WHERE t.table_schema = 'RAW'\n"
        "  AND t.table_name = :source_table;\n"
        "\n"
        "-- Step 2: Identify partition columns from clustering + query patterns\n"
        "SELECT\n"
        "    column_name,\n"
        "    data_type,\n"
        "    -- Check if this column appears frequently in WHERE clauses\n"
        "    (SELECT COUNT(*)\n"
        "     FROM snowflake.account_usage.access_history ah,\n"
        "          LATERAL FLATTEN(ah.direct_objects_accessed) f\n"
        "     WHERE f.value:objectName ILIKE '%' || :source_table || '%'\n"
        "    ) AS query_frequency\n"
        "FROM information_schema.columns\n"
        "WHERE table_name = :source_table\n"
        "ORDER BY ordinal_position;\n"
        "\n"
        "-- Step 3: Generate optimised contract template\n"
        "-- Output: dbt model SQL with:\n"
        "--   - Partition-aware WHERE clause (e.g., date >= DATEADD(-90d))\n"
        "--   - Clustering key matching existing table clustering\n"
        "--   - Incremental materialisation where applicable\n"
        "--   - Cost estimate based on table size and query frequency\n"
        "--\n"
        "-- Example generated contract:\n"
        "-- {{ config(materialized='incremental', unique_key='id',\n"
        "--           cluster_by=['transaction_date', 'region']) }}\n"
        "-- SELECT ...\n"
        "-- FROM {{ source('raw', 'transactions') }}\n"
        "-- WHERE transaction_date >= DATEADD('day', -90, CURRENT_DATE())\n"
        "-- {% if is_incremental() %}\n"
        "--   AND transaction_date >= (SELECT MAX(transaction_date) FROM {{ this }})\n"
        "-- {% endif %}"
    )

    doc.add_heading("Cost Pre-Scoring: Estimate Before You Deploy", level=3)
    doc.add_paragraph(
        "Before a new dbt contract is deployed to production, the framework analyses it "
        "and generates a cost pre-score. This is the gate that's missing today:"
    )
    add_sql_block(doc,
        "-- Cost Pre-Score: estimate cost of a new contract before deployment\n"
        "-- Uses EXPLAIN + historical patterns to estimate credits\n"
        "\n"
        "-- Step 1: Analyse the proposed query plan\n"
        "EXPLAIN USING JSON\n"
        "  SELECT ... FROM raw.large_table JOIN raw.lookup ...\n"
        "  WHERE date_col >= DATEADD('day', -30, CURRENT_DATE());\n"
        "\n"
        "-- Step 2: Check for overlapping existing contracts\n"
        "SELECT\n"
        "    ec.contract_name,\n"
        "    ec.source_tables,\n"
        "    ec.columns_produced,\n"
        "    ec.monthly_cost,\n"
        "    ec.owner_bu,\n"
        "    -- Calculate column overlap %\n"
        "    ARRAY_SIZE(ARRAY_INTERSECTION(ec.columns_produced, :proposed_columns))\n"
        "      / ARRAY_SIZE(:proposed_columns) * 100 AS overlap_pct\n"
        "FROM contract_registry ec\n"
        "WHERE ARRAYS_OVERLAP(ec.source_tables, :proposed_source_tables)\n"
        "ORDER BY overlap_pct DESC;\n"
        "\n"
        "-- Step 3: Generate pre-score report\n"
        "-- Output: 'Proposed contract estimated at $450/mo.\n"
        "--          Finance BU contract fin_transactions covers 85% at $12/mo.\n"
        "--          Recommendation: Extend fin_transactions with 2 columns\n"
        "--          instead of creating new contract. Estimated saving: $435/mo.'"
    )

    doc.add_heading("Continuous Overlap Detection", level=3)
    doc.add_paragraph(
        "The framework doesn't just check at creation time. It continuously "
        "monitors the contract registry and query patterns to detect:"
    )
    add_styled_table(doc,
        ["Detection Pattern", "Detection Method", "Action"],
        [
            ["Two BU contracts query 80%+ of the same source columns",
             "Contract registry column overlap analysis",
             "Alert platform team: recommend merging into shared contract in Broad layer"],
            ["New contract deployed without passing cost pre-score",
             "CI/CD hook or scheduled audit",
             "Flag for review; auto-generate optimisation recommendations"],
            ["Contract cost drifts 20%+ above its original estimate",
             "Monthly cost tracking vs pre-score baseline",
             "Cost drift alert with root cause (data growth? query change? new users?)"],
            ["Query pattern scans 10x more data than contract metadata suggests",
             "QUERY_HISTORY scan analysis vs INFORMATION_SCHEMA table size",
             "Tier 1 fix: send partition filter patch to contract owner"],
        ],
    )

    doc.add_paragraph()
    add_callout_box(doc,
        "The governed lifecycle means every new contract is optimised by default — the "
        "framework generates efficient SQL from metadata in minutes, not days. Users don't "
        "need to learn about partition columns or clustering keys. Overlap detection prevents "
        "duplicate effort across BUs. Cost pre-scoring ensures no contract goes to production "
        "without a cost estimate. The platform team shifts from reactive firefighting to "
        "proactive governance.",
    )

    doc.add_page_break()

    # =======================================================================
    # SECTION 7: PHASE 3 — GOVERNANCE & FINOPS
    # =======================================================================
    doc.add_heading("7. Phase 3: Governance & FinOps", level=1)
    doc.add_paragraph(
        "Phase 3 establishes the ongoing governance framework that prevents cost regression "
        "and embeds FinOps practices into the organisation."
    )

    doc.add_heading("7.1 Anomaly Detection", level=2)
    doc.add_paragraph(
        "Two complementary approaches to cost anomaly detection:"
    )

    doc.add_heading("Snowflake ANOMALY_DETECTION (ML-Powered)", level=3)
    doc.add_paragraph(
        "Snowflake's built-in ML functions (SNOWFLAKE.ML.ANOMALY_DETECTION) can train on "
        "historical cost data and detect statistical anomalies automatically."
    )
    add_sql_block(doc,
        "-- Create anomaly detection model on daily costs\n"
        "CREATE OR REPLACE SNOWFLAKE.ML.ANOMALY_DETECTION cost_anomaly_model(\n"
        "    INPUT_DATA => SYSTEM$REFERENCE('TABLE', 'pub_daily_costs'),\n"
        "    TIMESTAMP_COLNAME => 'cost_date',\n"
        "    TARGET_COLNAME => 'daily_total_cost',\n"
        "    LABEL_COLNAME => ''\n"
        ");\n"
        "\n"
        "-- Detect anomalies in recent data\n"
        "CALL cost_anomaly_model!DETECT_ANOMALIES(\n"
        "    INPUT_DATA => SYSTEM$REFERENCE('TABLE', 'pub_daily_costs_recent'),\n"
        "    TIMESTAMP_COLNAME => 'cost_date',\n"
        "    TARGET_COLNAME => 'daily_total_cost'\n"
        ");"
    )

    doc.add_paragraph()
    doc.add_heading("Custom Statistical Detection", level=3)
    doc.add_paragraph(
        "For accounts without ML functions, we implement custom z-score based anomaly detection:"
    )
    add_sql_block(doc,
        "-- Custom anomaly detection using z-scores\n"
        "WITH daily_stats AS (\n"
        "    SELECT\n"
        "        cost_date,\n"
        "        daily_total_cost,\n"
        "        AVG(daily_total_cost) OVER (ORDER BY cost_date ROWS 30 PRECEDING) AS rolling_avg,\n"
        "        STDDEV(daily_total_cost) OVER (ORDER BY cost_date ROWS 30 PRECEDING) AS rolling_std\n"
        "    FROM pub_daily_costs\n"
        ")\n"
        "SELECT *,\n"
        "    (daily_total_cost - rolling_avg) / NULLIF(rolling_std, 0) AS z_score,\n"
        "    CASE WHEN z_score > 2.5 THEN 'HIGH'\n"
        "         WHEN z_score > 2.0 THEN 'MEDIUM'\n"
        "         ELSE 'NORMAL' END AS anomaly_level\n"
        "FROM daily_stats\n"
        "WHERE z_score > 2.0\n"
        "ORDER BY cost_date DESC;"
    )

    doc.add_paragraph()
    doc.add_heading("7.2 Chargeback Model", level=2)
    doc.add_paragraph(
        "The chargeback model allocates costs to teams and products using a three-tier approach:"
    )
    add_styled_table(doc,
        ["Tier", "Method", "Coverage", "Accuracy"],
        [
            ["Tier 1: Direct", "Dedicated warehouses per team", "~60% of compute", "High"],
            ["Tier 2: Tag-based", "Object tags (cost_centre, team, data_product)", "~30% of compute", "High"],
            ["Tier 3: Proportional", "Role/user-based allocation for shared warehouses", "~10% of compute", "Medium"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("Budget Integration", level=3)
    doc.add_paragraph(
        "Snowflake Budgets (Enterprise Edition+) can be integrated with the chargeback model "
        "to set and monitor per-team spending limits:"
    )
    add_sql_block(doc,
        "-- Create budget for a team\n"
        "CREATE BUDGET IF NOT EXISTS engineering_budget\n"
        "    WITH\n"
        "        CREDIT_QUOTA = 5000,  -- Monthly credit limit\n"
        "        FREQUENCY = 'MONTHLY',\n"
        "        START_TIMESTAMP = '2026-05-01';\n"
        "\n"
        "-- Add warehouses to the budget\n"
        "ALTER BUDGET engineering_budget ADD\n"
        "    WAREHOUSE etl_wh,\n"
        "    WAREHOUSE analytics_wh;"
    )

    doc.add_paragraph()
    doc.add_heading("7.3 FinOps Review Cadence", level=2)
    add_styled_table(doc,
        ["Cadence", "Audience", "Focus", "Key Actions"],
        [
            ["Weekly (30 min)", "Engineering leads", "Cost anomalies, top movers, new recommendations", "Triage anomalies, assign optimisation tickets"],
            ["Monthly (60 min)", "Engineering + Finance", "Budget vs actual, chargeback review, trend analysis", "Approve optimisations, adjust budgets"],
            ["Quarterly (90 min)", "Leadership + Finance", "Strategic cost review, maturity assessment, roadmap", "Set cost targets, approve governance changes"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("7.4 Drift Detection", level=2)
    doc.add_paragraph(
        "Automated detection of configuration drift that impacts costs:"
    )
    add_styled_table(doc,
        ["Drift Type", "Detection", "Alert Trigger"],
        [
            ["Warehouse size change", "Monitor SHOW WAREHOUSES daily", "Size increased without approval"],
            ["Auto-suspend change", "Compare current vs baseline config", "Auto-suspend increased or disabled"],
            ["New warehouse created", "Compare warehouse count daily", "Unplanned warehouse provisioning"],
            ["Clustering key added", "Monitor AUTO_CLUSTERING_HISTORY", "New tables with auto-clustering"],
            ["Time Travel increased", "Monitor DATA_RETENTION_TIME_IN_DAYS", "Retention increased beyond standard"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("7.5 FinOps Maturity Assessment", level=2)
    doc.add_paragraph(
        "Based on the FinOps Foundation's framework, we assess maturity across four pillars "
        "using a Crawl / Walk / Run progression:"
    )
    add_styled_table(doc,
        ["Pillar", "Crawl", "Walk", "Run"],
        [
            ["Visibility", "Basic cost dashboard, manual attribution", "Automated attribution, team-level views", "Real-time cost feeds, predictive trending"],
            ["Optimisation", "Ad-hoc query fixes, manual WH sizing", "Automated anti-pattern detection, scheduled reviews", "Continuous optimisation, auto-remediation"],
            ["Governance", "Monthly invoice review", "Weekly anomaly review, budgets in place", "Automated drift detection, policy-as-code"],
            ["Accountability", "Centralised cost ownership", "Team-level chargeback visible", "Team budgets with automated alerts, FinOps champions"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("7.6 Query Governance by User Type", level=2)
    doc.add_paragraph(
        "Effective governance requires different approaches for different user types. "
        "A single policy fails because data engineers, analysts, and business users "
        "interact with Snowflake in fundamentally different ways."
    )
    add_styled_table(doc,
        ["User Type", "Access Pattern", "Governance Approach", "Deliverables"],
        [
            [
                "Data Engineers",
                "Complex transforms, pipeline SQL, dbt models",
                "Code review standards, approved patterns, cost-aware development",
                "Query style guide, dbt model templates, CI/CD cost checks",
            ],
            [
                "Analysts",
                "Ad-hoc exploration, report building, BI dashboards",
                "Pre-approved query library, BI tool integration, curated data access",
                "Query catalogue, Power BI / Tableau connected to curated layers, training on efficient SQL patterns",
            ],
            [
                "Business Users",
                "Report consumption, ad-hoc data requests, self-service needs",
                "Self-service dashboards, guided interfaces, no raw table access",
                "Streamlit apps for common analyses, request process for new data needs, guided exploration tools",
            ],
            [
                "Service Accounts",
                "Scheduled pipelines, automated reports, API integrations",
                "Pipeline cost monitoring, execution frequency review, timeout policies",
                "Pipeline cost dashboard, schedule optimisation recommendations, resource governor configs",
            ],
        ],
    )

    doc.add_page_break()

    # =======================================================================
    # SECTION 8: EDITION CONSIDERATIONS
    # =======================================================================
    _build_tech_spec_sections_8_to_15(doc)
    _build_tech_spec_appendices(doc)

    add_doc_footer(doc)

    output_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "Snowflake_Cost_Optimisation_Technical_Specification.docx",
    )
    doc.save(output_path)
    print(f"Technical Specification saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate Snowflake Cost Optimisation documents")
    parser.add_argument("--exec-only", action="store_true", help="Generate executive overview only")
    parser.add_argument("--tech-only", action="store_true", help="Generate technical spec only")
    args = parser.parse_args()

    if args.exec_only:
        create_executive_overview()
    elif args.tech_only:
        create_technical_specification()
    else:
        create_executive_overview()
        create_technical_specification()
