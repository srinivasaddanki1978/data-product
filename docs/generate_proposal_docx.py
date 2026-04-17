"""Generate Word document for Snowflake Cost Optimisation Framework proposal."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def set_cell_shading(cell, color):
    """Set background color for a table cell."""
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), color)
    shading_elm.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_table_borders(table):
    """Set borders for the entire table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "BFBFBF")
        borders.append(border)
    tblPr.append(borders)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table)

    # Header row
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "2E4057")

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            if row_idx % 2 == 1:
                set_cell_shading(cell, "F2F2F2")

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def create_proposal():
    """Create the Snowflake Cost Optimisation Framework proposal document."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Configure heading styles
    for i in range(1, 4):
        heading_style = doc.styles[f"Heading {i}"]
        heading_style.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    # --- TITLE PAGE ---
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Snowflake Cost Optimisation\nFramework")
    run.font.size = Pt(32)
    run.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Solution Proposal")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x5A, 0x7D, 0x9A)

    doc.add_paragraph()
    doc.add_paragraph()

    # Metadata table on title page
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Prepared by", "Srinivas Addanki"),
        ("Date", "11 April 2026"),
        ("Version", "1.0"),
        ("Classification", "Confidential"),
    ]
    for i, (label, value) in enumerate(meta_data):
        meta_table.rows[i].cells[0].text = label
        meta_table.rows[i].cells[1].text = value
        meta_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True if meta_table.rows[i].cells[0].paragraphs[0].runs else None
        for cell in meta_table.rows[i].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11)

    doc.add_page_break()

    # --- TABLE OF CONTENTS ---
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Problem Statement",
        "3. Solution Overview",
        "4. Phase 1 — Cost Visibility and Attribution",
        "5. Phase 2 — Query Optimisation and Recommendations",
        "6. Technical Architecture",
        "7. Delivery Plan",
        "8. Prerequisites and Assumptions",
        "9. Risk and Mitigation",
        "10. Why This Approach",
        "11. Team and Expertise",
        "12. Next Steps",
        "Appendix — Key Snowflake Metadata Views",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # --- 1. EXECUTIVE SUMMARY ---
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "Organisations running multiple business teams on Snowflake often face a common challenge: "
        "without standardised governance and visibility, compute and storage costs grow unpredictably. "
        "Teams adopt their own patterns, warehouses are provisioned without sizing guidelines, and there "
        "is no single view of who is spending how much and why."
    )
    doc.add_paragraph(
        "This proposal presents a Snowflake Cost Optimisation Framework delivered in two phases:"
    )
    bullets = [
        "Phase 1 (4 weeks): Analyse the Snowflake environment, identify cost drivers, and deliver an "
        "interactive dashboard showing cost attribution by warehouse, team, data product, and user.",
        "Phase 2 (4 weeks): Evaluate query efficiency, detect anti-patterns, and provide prioritised "
        "optimisation recommendations with estimated savings.",
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_paragraph(
        "The framework is built entirely on Snowflake-native metadata — no external monitoring tools "
        "are required. All analysis runs inside the customer's own Snowflake account, ensuring data "
        "never leaves their environment."
    )

    # --- 2. PROBLEM STATEMENT ---
    doc.add_heading("2. Problem Statement", level=1)
    doc.add_heading("Current Challenges", level=2)

    add_styled_table(
        doc,
        ["Challenge", "Impact"],
        [
            ["No cost attribution", 'Cannot answer "which team or process is driving our Snowflake bill"'],
            ["Oversized warehouses", 'Teams provision large warehouses "just in case" — paying for idle compute'],
            ["No query governance", "Inefficient SQL patterns (full table scans, SELECT *, missing filters) run unchecked"],
            ["Storage sprawl", "Unused tables, excessive Time Travel retention, and forgotten clones accumulate silently"],
            ["Reactive cost management", "Costs are reviewed monthly from invoices, not proactively monitored"],
            ["No standardisation", "Each team adopts its own patterns — no shared best practices"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("Business Impact", level=2)
    impacts = [
        "Snowflake costs grow 15–30% quarter-over-quarter without intervention",
        "Finance teams cannot allocate cloud costs to business units accurately",
        "Engineering teams lack data to justify optimisation work vs. new features",
        "Potential savings of 20–40% remain unrealised",
    ]
    for impact in impacts:
        doc.add_paragraph(impact, style="List Bullet")

    # --- 3. SOLUTION OVERVIEW ---
    doc.add_heading("3. Solution Overview", level=1)
    doc.add_heading("Core Idea", level=2)
    doc.add_paragraph(
        "Snowflake captures rich metadata about every query, warehouse, and storage object in its "
        "SNOWFLAKE.ACCOUNT_USAGE schema. This data is available for the past 365 days at no additional "
        "cost. Our framework transforms this raw metadata into actionable cost intelligence."
    )

    doc.add_heading("Solution Components", level=2)
    doc.add_paragraph(
        "The solution consists of three layers, all running inside the customer's Snowflake account:"
    )
    components = [
        "Data Source: SNOWFLAKE.ACCOUNT_USAGE metadata views (captured automatically by Snowflake)",
        "Transform Layer: dbt models that calculate costs, detect patterns, and generate recommendations",
        "Presentation Layer: Streamlit dashboard (native in Snowflake) for interactive exploration",
    ]
    for comp in components:
        doc.add_paragraph(comp, style="List Bullet")

    p = doc.add_paragraph()
    run = p.add_run("Everything runs inside Snowflake. No data extraction. No third-party SaaS tools. No ongoing licence costs.")
    run.bold = True

    # --- 4. PHASE 1 ---
    doc.add_heading("4. Phase 1 — Cost Visibility and Attribution", level=1)
    doc.add_heading("Objective", level=2)
    doc.add_paragraph("Provide clear, drillable answers to:")
    questions = [
        "What is our total Snowflake spend (compute, storage, serverless)?",
        "Which warehouses are the most expensive?",
        "Which teams, users, and data products drive the most cost?",
        "Where is storage being wasted?",
        "What are the cost trends over the past 90 days?",
    ]
    for q in questions:
        doc.add_paragraph(q, style="List Bullet")

    # 4.1 Compute
    doc.add_heading("4.1 Compute Cost Analysis", level=2)
    doc.add_paragraph("Data Source: WAREHOUSE_METERING_HISTORY, QUERY_HISTORY")
    add_styled_table(
        doc,
        ["Metric", "Description"],
        [
            ["Credits consumed per warehouse", "Hourly/daily/monthly credit burn"],
            ["Cost per query (estimated)", "Query runtime weighted by warehouse size and credit price"],
            ["Warehouse idle time", "Periods where a warehouse is running but executing no queries"],
            ["Queue wait time", "Queries waiting for available compute slots"],
            ["Cost by user / role", "Which users or service accounts drive the most compute"],
            ["Cost by data product", "Attributed via query tags, warehouse assignment, or role mapping"],
            ["Peak vs. off-peak usage", "Identify scheduling opportunities to shift workloads"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("Credit-to-Cost Conversion", level=3)
    p = doc.add_paragraph()
    run = p.add_run(
        "Estimated Query Cost ($) = (query_execution_time / 3600) "
        "x warehouse_credit_per_hour x credit_price_dollars"
    )
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    doc.add_paragraph(
        "Where warehouse_credit_per_hour depends on warehouse size (XS=1, S=2, M=4, L=8, XL=16) "
        "and credit_price_dollars is based on the customer's Snowflake contract."
    )

    # 4.2 Storage
    doc.add_heading("4.2 Storage Cost Analysis", level=2)
    doc.add_paragraph("Data Source: TABLE_STORAGE_METRICS, STORAGE_USAGE")
    add_styled_table(
        doc,
        ["Metric", "Description"],
        [
            ["Active storage", "Data currently in use"],
            ["Time Travel storage", "Historical data retained for the Time Travel window"],
            ["Fail-safe storage", "7-day regulatory recovery storage (non-configurable)"],
            ["Storage by database / schema", "Identify which areas consume the most storage"],
            ["Unused tables", "Tables with zero reads in 30/60/90 days (via ACCESS_HISTORY)"],
            ["Clone overhead", "Storage consumed by zero-copy clones that have diverged"],
        ],
    )

    # 4.3 Serverless
    doc.add_paragraph()
    doc.add_heading("4.3 Serverless Cost Analysis", level=2)
    doc.add_paragraph(
        "Data Source: PIPE_USAGE_HISTORY, AUTOMATIC_CLUSTERING_HISTORY, "
        "MATERIALIZED_VIEW_REFRESH_HISTORY, SERVERLESS_TASK_HISTORY"
    )
    add_styled_table(
        doc,
        ["Metric", "Description"],
        [
            ["Snowpipe credits", "Cost of continuous data loading"],
            ["Auto-clustering credits", "Cost of maintaining clustering on tables"],
            ["Materialized view refresh credits", "Cost of keeping MVs up-to-date"],
            ["Task execution credits", "Cost of scheduled serverless tasks"],
        ],
    )

    # 4.4 Dashboard
    doc.add_paragraph()
    doc.add_heading("4.4 Dashboard Deliverable", level=2)
    doc.add_paragraph(
        "An interactive Streamlit dashboard (hosted natively in Snowflake) with the following views:"
    )
    views = [
        "Executive Summary — Total spend, compute/storage/serverless split, month-over-month trend",
        "Warehouse Deep Dive — Per-warehouse cost, utilisation, idle %, queue time",
        "Team Attribution — Cost allocated to business teams via role/warehouse/query tag mapping",
        "Storage Explorer — Table-level storage breakdown, unused table list, Time Travel waste",
        "Trend Analysis — 90-day cost trends with anomaly highlighting",
    ]
    for v in views:
        doc.add_paragraph(v, style="List Bullet")

    # --- 5. PHASE 2 ---
    doc.add_heading("5. Phase 2 — Query Optimisation and Recommendations", level=1)
    doc.add_heading("Objective", level=2)
    doc.add_paragraph(
        "Identify why costs are high and provide actionable recommendations to reduce them, "
        "prioritised by estimated savings."
    )

    # 5.1 Warehouse
    doc.add_heading("5.1 Warehouse Right-Sizing", level=2)
    add_styled_table(
        doc,
        ["Signal", "Detection Method", "Recommendation"],
        [
            ["Consistent queuing", "queued_overload_time > 0 frequently", "Scale up warehouse or enable multi-cluster"],
            ["Consistent idle time", "Warehouse running with no queries for extended periods", "Reduce auto-suspend interval (e.g., 300s to 60s)"],
            ["Oversized for workload", "Small queries running on L/XL warehouses", "Downsize warehouse or route to a smaller one"],
            ["Uneven load", "Spikes at certain hours, idle otherwise", "Schedule workloads to consolidate usage windows"],
        ],
    )

    # 5.2 Query Anti-Patterns
    doc.add_paragraph()
    doc.add_heading("5.2 Query Anti-Pattern Detection", level=2)
    add_styled_table(
        doc,
        ["Anti-Pattern", "Detection Method", "Impact"],
        [
            ["Full table scans", "PARTITIONS_SCANNED / PARTITIONS_TOTAL > 0.8", "Excessive compute for large tables"],
            ["SELECT *", "Query text pattern matching", "Scans all columns; wastes I/O"],
            ["Missing filters", "High BYTES_SCANNED relative to ROWS_PRODUCED", "Reads far more data than needed"],
            ["Spill to storage", "BYTES_SPILLED_TO_LOCAL/REMOTE_STORAGE > 0", "Query needs more memory than warehouse provides"],
            ["Repeated identical queries", "Same QUERY_PARAMETERIZED_HASH running frequently", "Results should be cached or materialised"],
            ["Cartesian joins", "ROWS_PRODUCED >> ROWS_SCANNED in join queries", "Missing or incorrect join conditions"],
            ["Excessive ORDER BY", "Large result sets with ORDER BY but no LIMIT", "Sorting millions of rows unnecessarily"],
        ],
    )

    # 5.3 Storage Optimisation
    doc.add_paragraph()
    doc.add_heading("5.3 Storage Optimisation Recommendations", level=2)
    add_styled_table(
        doc,
        ["Opportunity", "Detection", "Recommendation"],
        [
            ["Unused tables", "No reads in ACCESS_HISTORY for 90+ days", "Archive or drop"],
            ["Excessive Time Travel", "Retention set to 90 days on non-critical tables", "Reduce to 1 day (save ~99% of TT storage)"],
            ["Transient table candidates", "Tables rebuilt daily (ephemeral data)", "Convert to TRANSIENT (eliminates Fail-safe)"],
            ["Stale clones", "Clones created months ago, not accessed", "Drop"],
            ["Large staging tables", "Temporary/staging data persisted permanently", "Add lifecycle policies or auto-drop"],
        ],
    )

    # 5.4 Prioritised Recommendations
    doc.add_paragraph()
    doc.add_heading("5.4 Prioritised Recommendations Report", level=2)
    doc.add_paragraph("Each recommendation will include:")
    add_styled_table(
        doc,
        ["Field", "Description"],
        [
            ["Category", "Warehouse / Query / Storage"],
            ["Object", "Warehouse name, query ID, table name"],
            ["Current Cost", "Estimated monthly cost of the current pattern"],
            ["Estimated Savings", "Projected monthly saving if recommendation is applied"],
            ["Effort", "Low (config change) / Medium (query rewrite) / High (architecture change)"],
            ["Risk", "Impact assessment of making the change"],
            ["Action", "Specific steps to implement the recommendation"],
            ["Priority", "Ranked by savings / effort (ROI-based)"],
        ],
    )

    # --- 6. TECHNICAL ARCHITECTURE ---
    doc.add_paragraph()
    doc.add_heading("6. Technical Architecture", level=1)
    doc.add_heading("Technology Stack", level=2)
    add_styled_table(
        doc,
        ["Component", "Technology", "Rationale"],
        [
            ["Data Modelling", "dbt (data build tool)", "Industry-standard transformation framework; version-controlled, testable, documented"],
            ["Data Platform", "Snowflake", "Customer's existing platform; zero data movement"],
            ["Dashboards", "Streamlit in Snowflake", "Native to Snowflake; no separate hosting; interactive"],
            ["Infrastructure", "Terraform (optional)", "Reproducible deployment of schemas, roles, warehouses"],
            ["Scheduling", "dbt Cloud or Snowflake Tasks", "Automated daily refresh of cost models"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("Data Flow", level=2)
    flow_steps = [
        "1. SNOWFLAKE.ACCOUNT_USAGE (raw metadata, 365 days retention)",
        "2. dbt Staging Models — Source definitions + light cleaning",
        "3. dbt Intermediate Models — Business logic: cost calculations, utilisation metrics, pattern detection",
        "4. dbt Publication Models — Consumer-ready: cost_by_warehouse, cost_by_team, optimisation_candidates",
        "5. Streamlit Dashboard — Interactive visualisation and recommendation explorer",
    ]
    for step in flow_steps:
        doc.add_paragraph(step, style="List Bullet")

    doc.add_heading("Model Design Principles", level=2)
    principles = [
        "Idempotent: Models can be re-run at any time without side effects",
        "Incremental: Large tables (query history) use incremental materialisation for efficiency",
        "Tested: dbt tests validate data quality (uniqueness, referential integrity, accepted values)",
        "Documented: Every model and column is documented in dbt, browsable via dbt docs",
        "Configurable: Credit prices, warehouse mappings, and team allocations maintained as seed files — easy to update without code changes",
    ]
    for p_text in principles:
        doc.add_paragraph(p_text, style="List Bullet")

    # --- 7. DELIVERY PLAN ---
    doc.add_heading("7. Delivery Plan", level=1)
    doc.add_heading("Phase 1 — Cost Visibility and Attribution (4 Weeks)", level=2)
    add_styled_table(
        doc,
        ["Week", "Activities", "Deliverables"],
        [
            [
                "Week 1",
                "Environment access and discovery. Grant IMPORTED PRIVILEGES. Profile the account: warehouses, databases, users, query volume. Set up dbt project.",
                "Discovery report: account profile, initial findings",
            ],
            [
                "Week 2",
                "Build core dbt models: warehouse credit usage, query cost attribution, storage breakdown. Configure team/data product mapping.",
                "Working dbt models for compute, storage, and serverless costs",
            ],
            [
                "Week 3",
                "Build Streamlit dashboard: executive summary, warehouse deep-dive, team attribution, storage explorer, trend analysis.",
                "Interactive dashboard v1",
            ],
            [
                "Week 4",
                "Refinement, testing, and documentation. Walkthrough with stakeholders. Knowledge transfer.",
                "Final Phase 1 dashboard, user guide, dbt documentation",
            ],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("Phase 2 — Optimisation Recommendations (4 Weeks)", level=2)
    add_styled_table(
        doc,
        ["Week", "Activities", "Deliverables"],
        [
            [
                "Week 5",
                "Build warehouse right-sizing models. Analyse utilisation patterns (idle time, queuing, peak/off-peak).",
                "Warehouse sizing recommendations",
            ],
            [
                "Week 6",
                "Build query anti-pattern detection. Identify full scans, spill-to-storage, repeated queries, SELECT *.",
                "Query-level optimisation candidates",
            ],
            [
                "Week 7",
                "Build storage optimisation models. Identify unused tables, Time Travel waste, transient candidates. Prioritise all recommendations by ROI.",
                "Storage recommendations + prioritised action list",
            ],
            [
                "Week 8",
                "Enhance dashboard with recommendations tab. Final testing. Stakeholder walkthrough. Knowledge transfer and handover.",
                "Complete framework, final report with estimated savings",
            ],
        ],
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Total Duration: 8 weeks")
    run.bold = True
    doc.add_paragraph("Phase 1 and Phase 2 delivered sequentially. Phase 1 delivers visible value by Week 3.")

    # --- 8. PREREQUISITES ---
    doc.add_heading("8. Prerequisites and Assumptions", level=1)
    doc.add_heading("Customer Prerequisites", level=2)
    add_styled_table(
        doc,
        ["#", "Requirement", "Purpose"],
        [
            ["1", "Grant IMPORTED PRIVILEGES on the SNOWFLAKE database to the analytics role", "Access to ACCOUNT_USAGE metadata views"],
            ["2", "Provide a dedicated Snowflake warehouse (Small or Medium) for running the framework", "Compute for dbt models and dashboard queries"],
            ["3", "Provide a dedicated database/schema for the framework objects", "Storage for the cost models and dashboard"],
            ["4", "Share Snowflake contract details (credit price, edition)", "Accurate dollar-cost conversion"],
            ["5", "Provide organisational mapping (warehouses/roles to teams)", "Cost attribution to business units"],
            ["6", "Nominate a technical point of contact", "Collaboration during discovery and validation"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("Assumptions", level=2)
    assumptions = [
        "Customer is on Snowflake Enterprise Edition or higher (required for ACCESS_HISTORY)",
        "ACCOUNT_USAGE views have data for at least 30 days (ideally 90+ days for trend analysis)",
        "The framework warehouse will not be used for other workloads (to avoid cost contamination)",
        "Query tags are either already in use or can be introduced for future cost attribution",
        "Customer has dbt Cloud or is open to using it (alternatively, Snowflake Tasks can schedule the models)",
    ]
    for a in assumptions:
        doc.add_paragraph(a, style="List Bullet")

    # --- 9. RISK ---
    doc.add_heading("9. Risk and Mitigation", level=1)
    add_styled_table(
        doc,
        ["Risk", "Likelihood", "Impact", "Mitigation"],
        [
            ["IMPORTED PRIVILEGES not granted promptly", "Medium", "Blocks all work", "Raise as Day 1 action; provide exact GRANT statement"],
            ["Customer on Standard Edition (no ACCESS_HISTORY)", "Low", "Cannot detect unused tables", "Fall back to TABLE_STORAGE_METRICS + manual review"],
            ["No query tags in use", "Medium", "Limits cost attribution granularity", "Attribute by warehouse and role; recommend query tagging as future improvement"],
            ["Very high query volume (>1M queries/day)", "Low", "Slow model builds", "Use incremental materialisation with 7-day lookback windows"],
            ["Stakeholder availability for walkthroughs", "Medium", "Delays sign-off", "Schedule walkthroughs at project start; async review via shared dashboard"],
        ],
    )

    # --- 10. WHY THIS APPROACH ---
    doc.add_heading("10. Why This Approach", level=1)
    doc.add_heading("vs. Snowflake-Native Cost Management", level=2)
    doc.add_paragraph(
        "Snowflake provides basic budgets and resource monitors, but they only alert on thresholds — "
        "they do not explain why costs are high or what to do about it. Our framework provides the "
        "analytical layer that turns alerts into action."
    )

    doc.add_heading("vs. Third-Party SaaS Tools (Select.dev, Keebo, Sundeck)", level=2)
    add_styled_table(
        doc,
        ["Factor", "Our Framework", "SaaS Tools"],
        [
            ["Data residency", "Stays in customer's Snowflake account", "Data sent to third-party"],
            ["Customisation", "Fully tailored to customer's org structure", "Generic, one-size-fits-all"],
            ["Ongoing cost", "No licence fees (dbt + Streamlit are free/included)", "$500–$5,000+/month"],
            ["Business context", "Integrates team mappings, query tags, data product names", "Limited to Snowflake metadata only"],
            ["Extensibility", "Customer can add new models, metrics, dashboards", "Limited to vendor roadmap"],
            ["Transparency", "All logic is visible in dbt SQL models", "Black-box recommendations"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("Proven Expertise", level=2)
    doc.add_paragraph("Our team operates a production-grade dbt + Snowflake + Streamlit platform that already:")
    proven = [
        "Manages 10+ data products on Snowflake",
        "Implements query tagging for cost attribution",
        "Queries SNOWFLAKE.ACCOUNT_USAGE for lineage and monitoring",
        "Runs interactive Streamlit dashboards for operational observability",
        "Integrates with Datadog and Incident.io for alerting",
    ]
    for item in proven:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph(
        "This framework extends our proven patterns to a new use case — cost optimisation."
    )

    # --- 11. TEAM ---
    doc.add_heading("11. Team and Expertise", level=1)
    doc.add_heading("Proposed Team Composition", level=2)
    add_styled_table(
        doc,
        ["Role", "Responsibility", "Allocation"],
        [
            ["Lead Engineer", "Architecture, dbt model development, Snowflake expertise", "Full-time (8 weeks)"],
            ["Dashboard Developer", "Streamlit dashboard development, UX design", "Part-time (Weeks 3-4, 7-8)"],
            ["Project Lead", "Stakeholder management, delivery oversight, customer communication", "Part-time (throughout)"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("Key Competencies", level=2)
    competencies = [
        "Snowflake architecture and performance tuning",
        "dbt (data build tool) — modelling, testing, documentation",
        "Streamlit dashboard development",
        "Terraform infrastructure-as-code for Snowflake",
        "Data observability and monitoring frameworks",
    ]
    for c in competencies:
        doc.add_paragraph(c, style="List Bullet")

    # --- 12. NEXT STEPS ---
    doc.add_heading("12. Next Steps", level=1)
    add_styled_table(
        doc,
        ["#", "Action", "Owner", "Timeline"],
        [
            ["1", "Review and approve this proposal", "Customer", "By 15 April 2026"],
            ["2", "Schedule a kickoff call", "Both", "Week of 16 April"],
            ["3", "Provide environment access and prerequisites", "Customer", "Before kickoff"],
            ["4", "Begin Phase 1 discovery", "Delivery team", "Kickoff + 1 day"],
        ],
    )

    # --- APPENDIX ---
    doc.add_page_break()
    doc.add_heading("Appendix — Key Snowflake Metadata Views", level=1)
    doc.add_paragraph(
        "The framework leverages the following SNOWFLAKE.ACCOUNT_USAGE views. These are available "
        "to any Snowflake account with IMPORTED PRIVILEGES granted on the SNOWFLAKE database."
    )

    doc.add_heading("Compute", level=2)
    add_styled_table(
        doc,
        ["View", "Description", "Retention"],
        [
            ["WAREHOUSE_METERING_HISTORY", "Credit consumption per warehouse, per hour", "365 days"],
            ["QUERY_HISTORY", "Full detail of every query: runtime, bytes scanned, spill, partitions, user, warehouse, query tag", "365 days"],
            ["WAREHOUSE_LOAD_HISTORY", "Warehouse utilisation: running, queued, blocked queries per interval", "365 days"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("Storage", level=2)
    add_styled_table(
        doc,
        ["View", "Description", "Retention"],
        [
            ["TABLE_STORAGE_METRICS", "Active, Time Travel, and Fail-safe bytes per table", "Current snapshot"],
            ["STORAGE_USAGE", "Total account-level storage over time", "365 days"],
            ["DATABASE_STORAGE_USAGE_HISTORY", "Storage per database over time", "365 days"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("Access and Lineage", level=2)
    add_styled_table(
        doc,
        ["View", "Description", "Retention"],
        [
            ["ACCESS_HISTORY", "Which tables/columns were read or written by each query", "365 days"],
            ["LOGIN_HISTORY", "User login events (for activity analysis)", "365 days"],
            ["SESSIONS", "Session details including client application", "365 days"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("Serverless Features", level=2)
    add_styled_table(
        doc,
        ["View", "Description", "Retention"],
        [
            ["AUTOMATIC_CLUSTERING_HISTORY", "Credits consumed by auto-clustering", "365 days"],
            ["MATERIALIZED_VIEW_REFRESH_HISTORY", "Credits consumed by MV refreshes", "365 days"],
            ["PIPE_USAGE_HISTORY", "Credits consumed by Snowpipe", "365 days"],
            ["SERVERLESS_TASK_HISTORY", "Credits consumed by serverless tasks", "365 days"],
            ["SEARCH_OPTIMIZATION_HISTORY", "Credits consumed by search optimisation", "365 days"],
        ],
    )

    # --- FOOTER ---
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— End of Document —")
    run.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("This document is confidential and intended for internal use and customer presentation.")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Save
    output_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "Snowflake_Cost_Optimisation_Framework_Proposal.docx",
    )
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_proposal()
